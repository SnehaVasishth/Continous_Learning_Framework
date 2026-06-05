"""Twelve separate per-section DOCX builders for the Keysight RFP response.

Each section is a standalone document, written in business and vision language
for an audience that has not seen the demo. The wording avoids implementation
specifics (stage labels, library names, framework names, "27-rule book"
shorthand). Where Keysight has named systems or terms in their own RFP and
Q&A material (Docunet, Jitterbit, AIOA, FCNV, CCC Request, KSP, Purview),
those terms are used verbatim because they are part of Keysight's vocabulary.

The combined twelve documents are listed in SECTIONS at the bottom of the
file and served individually under /api/docs/rfp-reply/<slug>.docx.
"""
# === v1.1 RFP-REPLY START ===
from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

_INK = RGBColor(0x13, 0x14, 0x26)
_MUTED = RGBColor(0x6B, 0x72, 0x80)
_ACCENT = RGBColor(0x1A, 0x55, 0xF9)


# ──────────────────────────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────────────────────────

def _shade(cell, hex_rgb: str) -> None:
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_rgb)
    tcPr.append(shd)


def _new_doc() -> Document:
    doc = Document()
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
    return doc


def _title_block(doc: Document, section_label: str, section_title: str) -> None:
    label = doc.add_paragraph()
    run = label.add_run(section_label)
    run.font.size = Pt(11)
    run.font.color.rgb = _MUTED
    label.paragraph_format.space_after = Pt(0)

    title = doc.add_paragraph()
    run = title.add_run(section_title)
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = _ACCENT
    title.paragraph_format.space_after = Pt(2)

    meta = doc.add_paragraph()
    run = meta.add_run(
        "Keysight Technologies . Agentic AI for SalesOps RFP Response . "
        "The Hackett Group on ZBrain . Response date 2026-05-11"
    )
    run.font.size = Pt(9.5)
    run.italic = True
    run.font.color.rgb = _MUTED
    meta.paragraph_format.space_after = Pt(12)


def _h(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = _INK if level > 1 else _ACCENT
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(4)


def _p(doc: Document, text: str, *, bold: bool = False, italic: bool = False, muted: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.bold = bold
    run.italic = italic
    if muted:
        run.font.color.rgb = _MUTED
    p.paragraph_format.space_after = Pt(4)


def _bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.size = Pt(10.5)
        p.paragraph_format.space_after = Pt(2)


def _kv_table(doc: Document, rows: list[tuple[str, str]], *, key_header: str = "Topic", val_header: str = "Approach") -> None:
    t = doc.add_table(rows=1 + len(rows), cols=2)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    hdr[0].text = key_header
    hdr[1].text = val_header
    _shade(hdr[0], "1A55F9")
    _shade(hdr[1], "1A55F9")
    for cell in hdr:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
    for i, (k, v) in enumerate(rows, start=1):
        t.rows[i].cells[0].text = k
        t.rows[i].cells[1].text = v
        for cell in t.rows[i].cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    t.columns[0].width = Cm(5.5)
    t.columns[1].width = Cm(11)


def _three_col_table(doc: Document, headers: tuple[str, str, str], rows: list[tuple[str, str, str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=3)
    t.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = h
        _shade(cell, "1A55F9")
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            t.rows[i].cells[j].text = val
            t.rows[i].cells[j].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            for p in t.rows[i].cells[j].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)


def _spacer(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)


def _save(doc: Document) -> bytes:
    out = BytesIO()
    doc.save(out)
    return out.getvalue()


# ──────────────────────────────────────────────────────────────────
# 0. Cover Letter
# ──────────────────────────────────────────────────────────────────
def build_cover_letter() -> bytes:
    doc = _new_doc()

    # Eyebrow
    label = doc.add_paragraph()
    run = label.add_run("Cover note and document index")
    run.font.size = Pt(11)
    run.font.color.rgb = _MUTED
    label.paragraph_format.space_after = Pt(0)

    # Title
    title = doc.add_paragraph()
    run = title.add_run("Keysight RFP Response")
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = _ACCENT
    title.paragraph_format.space_after = Pt(2)

    # Meta line
    meta = doc.add_paragraph()
    run = meta.add_run(
        "Keysight Technologies . Agentic AI for SalesOps . "
        "The Hackett Group on ZBrain . Response date 12 May 2026"
    )
    run.font.size = Pt(9.5)
    run.italic = True
    run.font.color.rgb = _MUTED
    meta.paragraph_format.space_after = Pt(18)

    # Opening paragraph: the response in one paragraph
    _p(doc,
       "The Hackett Group submits this response to Keysight's Request for Proposal for an "
       "Agentic AI platform that automates the SalesOps front-office operation. The proposal "
       "is delivered through our enterprise services practice and our proprietary platform, "
       "ZBrain, under a single accountable team. The engagement follows a hybrid "
       "build-and-deploy model: solution design and build run inside ZBrain's Solution Builder "
       "on the Hackett-operated tenancy, and production deployment runs separately inside "
       "Keysight's own AWS infrastructure, so the runtime, the data path, the LLM endpoints, "
       "and the Knowledge Base content remain within Keysight's security and operational "
       "boundary. Delivery is a ten-week engagement to a working production pilot, followed by "
       "hypercare and a long-running managed service. The delivery team includes the Hackett "
       "engineers who carried out the prior proof-of-concept work in this domain.")

    _h(doc, "How this response is organised", 2)
    _p(doc,
       "The response comprises twelve numbered sections plus a role-by-provider RACI matrix. "
       "Each document is self-contained and may be distributed independently to Keysight's "
       "topic specialists. The high-level scope of each document is listed below.")

    index_rows: list[tuple[str, str]] = [
        ("1. Executive Summary",
         "The opportunity, the proposal at a glance, and the headline commitments."),
        ("2. Solution Overview",
         "How the solution works end to end, the ZBrain modules in scope, and how the solution sits inside Keysight's existing enterprise systems."),
        ("3. Scope",
         "In-scope and out-of-scope items by use case and sub-case, engagement assumptions, and the risk register."),
        ("4. Implementation Approach",
         "The ten-week phased delivery (Requirement Finalization, MVP Build, Testing and Tuning), the 60 / 20 / 20 data partition, the continuous-learning workflow, and the post-deployment lifecycle."),
        ("5. AI / ML Capabilities",
         "The platform capabilities Hackett brings to the engagement and the applied behaviour in the proposed solution."),
        ("6. Delivery Plan",
         "Phase milestones, cumulative acceptance criteria, governance cadence, and the exit gates that close each phase."),
        ("7. Team Composition",
         "Delivery team and responsibilities, week-by-week resource utilisation, the hypercare team, and the steady-state support model."),
        ("8. Pricing Model",
         "Modular and turnkey pricing structures, implementation cost separated from ongoing cost, the outcome-based variant, and worked examples at Keysight's stated volumes."),
        ("9. Infrastructure",
         "Deployment architecture inside Keysight's AWS account, the ZBrain platform modules, integration channels with Salesforce, Jitterbit, Oracle EBS, DocuNet, Outlook, Azure AD, and the Cloudflare edge."),
        ("10. Security and Governance",
         "Enterprise-level governance aligned to Keysight's Microsoft, AWS, and Cloudflare standards, application-level governance through the ZBrain Governance and Monitor modules, citizenship-based access, data security, and the audit trail."),
        ("11. References",
         "Comparable implementations matching Keysight's domain and use-case criteria, and the reference selection process."),
        ("12. Demo",
         "Demo format and slot, the reference implementation shown, the supplementary materials available, and the next steps."),
        ("13. RACI Matrix",
         "Role-by-provider responsibility assignment across the engagement and the post-deployment lifecycle."),
    ]

    for label_text, desc in index_rows:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(label_text + ". ")
        r.bold = True
        r.font.size = Pt(10.5)
        r2 = p.add_run(desc)
        r2.font.size = Pt(10.5)
        p.paragraph_format.space_after = Pt(2)

    _spacer(doc)

    _h(doc, "Engagement contact", 2)
    _p(doc,
       "For any clarification on this response, the evaluation team is invited to contact our "
       "named engagement lead. The Hackett Group is available to present the solution in "
       "person and to introduce the customer references aligned to the SalesOps and "
       "order-management use-case profile.")

    # Signature block
    _spacer(doc)
    for line, sz, bold in [
        ("[Signatory Name]",                        11,   True),
        ("Managing Director, AI and ML Practice",   10.5, False),
        ("The Hackett Group",                       10.5, False),
        ("[email]  /  [phone]",                     10.5, False),
    ]:
        sp = doc.add_paragraph()
        r = sp.add_run(line)
        r.font.size = Pt(sz)
        r.bold = bold
        sp.paragraph_format.space_after = Pt(0)

    return _save(doc)


# ──────────────────────────────────────────────────────────────────
# 1. Executive Summary
# ──────────────────────────────────────────────────────────────────
def build_executive_summary() -> bytes:
    doc = _new_doc()
    _title_block(doc, "Section 1 of 13", "Executive Summary")

    _h(doc, "The opportunity", 2)
    _p(doc,
       "Keysight's SalesOps front office handles approximately 880 thousand inbound customer "
       "emails per year (about 2,000 per day, per Keysight's Q&A clarification, item 11), spread "
       "across roughly 50 region-segregated mailboxes serving the Americas, Europe, Asia-Pac, "
       "and Japan. Eighty to ninety concurrent users staff this queue today out of a 600 to 700 "
       "person global sales operations population. Almost every step in that flow is manual: "
       "classifying the request, opening attachments (typical email sizes one to twenty MB), "
       "finding the right Salesforce record, creating or updating the CCC Request, routing to "
       "the right CSR, and drafting the reply to the customer in the right language. This is "
       "high-skill work, but it is also high-volume and repetitive, and the team's time is the "
       "scarce resource.")
    _p(doc,
       "The RFP names seven inbound-request flows in scope and asks for an Agentic AI platform "
       "that takes on this work end to end while keeping the team in control. The stated business "
       "outcome is a 60 to 70 percent reduction in manual CSR effort and a turnaround improvement "
       "from hours to minutes.")

    _h(doc, "What Keysight already has in place", 2)
    _p(doc,
       "Keysight has done a great deal of the upstream thinking already. The seven use cases in "
       "the RFP, the AIOA validation pattern on the Trade Order Entry, SOM WO Update, and "
       "Service Contract flows, the operational rules your CSRs apply during classification, the "
       "Outlook pre-filters, and Docunet (your document store) together with the practice of "
       "filing every email against its CCC Request with Doc type FCNV, are all in place. What is "
       "missing is an "
       "Agentic AI layer that connects them, runs across them at scale, and gets better at the "
       "work over time. That is the solution we are proposing.")

    _h(doc, "What we propose", 2)
    _p(doc,
       "The Hackett Group proposes the ZBrain Agentic AI platform as the foundation, delivered "
       "as a Combination Provider (Product plus Services) under a single accountable team. We "
       "will follow a hybrid build-and-deploy model. Design and build happen inside ZBrain's "
       "Solution Builder — the visual orchestration surface that wires the end-to-end agents "
       "against Keysight's rule book, Knowledge Bases, and integration channels — running on "
       "the ZBrain platform tenancy The Hackett Group operates. The production solution then "
       "deploys separately onto Keysight's own infrastructure (AWS, per the Keysight standard "
       "called out in Q&A item 55), so the runtime, the data path, the LLM endpoints, the "
       "Governance and Monitor signals, and the Knowledge Base content all sit inside Keysight's "
       "boundary. Our delivery team includes the engineers who carried out Keysight's earlier "
       "POC in this domain, so context is preserved from day one.")
    _p(doc,
       "The proposed solution operationalises an AI assistant that reads each email and its "
       "attachments, classifies the intent against your rule book, extracts the fields that "
       "matter for each use case, resolves the customer and product entities in Salesforce, "
       "validates the request against your existing data, and decides how confident it is. "
       "Confidence is scored using the four-gate model Keysight described during the Q&A "
       "clarification (Classification, Extraction, Entity Resolution, Action Feasibility), with "
       "the overall tier set by the lowest gate. Per Q&A item 22, the initial production "
       "posture is AI drafts, CSR reviews, CSR sends; autonomy expands use case by use case as "
       "confidence in each gate is validated. The assistant drafts the customer reply in the "
       "customer's detected language and files the email against the right CCC Request in "
       "DocuNet with Doc type FCNV.")

    _h(doc, "Headline commitments", 2)
    _bullets(doc, [
        "Classification accuracy at or above the 90 percent threshold the RFP names, measured against the regression corpus at every sprint cut. Tuning uses a 60 / 20 / 20 train / validation / test split.",
        "Initial production posture per Q&A item 22: AI drafts, CSR reviews, CSR sends. Autonomy expansion is a deliberate Keysight decision per use case, not an automatic ramp; every transition is logged and reversible.",
        "All three RAG types (structured database grounding, knowledge-graph traversal, unstructured document search) served by the ZBrain Knowledge Base module, per Q&A item 53.",
        "Customer-facing replies drafted in the customer's detected language, using Keysight's existing translation knowledge base extended for SalesOps-specific terminology, per Q&A items 26 and 27.",
        "Auditable trace per email: every reading, every extracted field, every decision, every external write recorded against a per-email reference and exported to Keysight's SIEM. Retention approximately ten years per Q&A item 45.",
        "Eight-week delivery to a working production pilot, structured as four two-week sprints, with Knowledge Bases the Keysight team can tune without code changes. Implementation is fixed-fee; post-deployment T&M for new use cases and enhancements (Implementation Approach 4.10).",
        "Continuity from the POC engineering team, so the institutional knowledge is preserved across the engagement.",
    ])

    _h(doc, "Differentiators", 2)
    _bullets(doc, [
        "Single team accountable for product and services. One MSA, one delivery lead, one escalation path. No vendor finger-pointing between platform issues and configuration issues.",
        "Hybrid build-and-deploy model. Solution Builder, Governance, Monitor, and Knowledge Base modules are used inside the ZBrain platform tenancy during design and build; the deployed solution then runs on Keysight's own infrastructure, so Keysight retains data residency, runtime control, and LLM-endpoint ownership.",
        "Engineering continuity from the prior POC engagement. Context on Keysight's CSR rule book, the AIOA validation pattern, and the multi-format attachment handling does not have to be rebuilt.",
        "Model-agnostic platform. Keysight keeps the choice of LLM provider (commercial, open-source) and can change that choice without rebuilding the solution.",
        "Configured rather than custom-coded. Knowledge bases hold the intents, the routing rules, the entity graphs, and the operational rules. Routine changes are operator edits, not developer releases.",
        "Application-level governance from day one through the ZBrain Governance and Monitor modules, with integration points to a future enterprise governance fabric as Keysight adopts it (per Q&A item 55).",
        "Hackett's AI Development Methodology, compressed to eight weeks because the upstream requirements work is already in the RFP.",
        "Real integrations with the systems Keysight already runs: Salesforce, Oracle EBS and DocuNet through Jitterbit (per Q&A item 56), Logik.io, KSP, Outlook through Microsoft Graph, Azure AD for SSO, Cloudflare for internet-exposed surfaces. ServiceNow-side integration stays with Keysight's dev team.",
    ])

    return _save(doc)


# ──────────────────────────────────────────────────────────────────
# 2. Solution Overview
# ──────────────────────────────────────────────────────────────────
def build_solution_overview() -> bytes:
    doc = _new_doc()
    _title_block(doc, "Section 2 of 13", "Solution Overview")

    _h(doc, "The vision in one paragraph", 2)
    _p(doc,
       "An AI assistant working alongside your CSR team. It reads every inbound email at the "
       "moment it arrives, understands what the customer is asking for, finds the right records "
       "in Salesforce, validates the request against your existing data, and either acts on it, "
       "asks a CSR to approve, or hands it over for full review. It drafts the customer reply in "
       "the right language. It keeps an audit trail of everything it did. And it gets better at "
       "its job over time as your team corrects it.")

    _h(doc, "How an email moves through the solution", 2)
    _p(doc,
       "Think of an inbound email as a small package of information. The solution carries that "
       "package through seven natural steps, each one performed automatically where it can be and "
       "by a CSR where the request asks for judgement. The seven steps are framed below in plain "
       "language, not engineering vocabulary.")

    _three_col_table(doc,
                     ("Step", "What the solution does", "What the CSR sees"),
                     [
                         ("1. Triage at the mailbox door",
                          "Before the AI does anything, deterministic filters handle the obviously non-actionable mail: bounces, out-of-office replies, Brazil tax filings, portal verification codes, collection notices, and government-customer (KSO) traffic that must be redirected for compliance. These are the rules your team has already documented; the solution preserves them.",
                          "The CSR sees less noise. Filtered mail is logged but never lands in the main queue."),
                         ("2. Understand the request",
                          "The AI reads the email body, walks through every attachment (PDF, Word, Excel, scanned images, forwarded Outlook items), detects the language, and identifies the customer's intent against your operational rule book. It distinguishes a PO from a status enquiry from a service work-order request, even when the same email touches multiple topics.",
                          "Each request shows a clear intent label and a confidence score the CSR can trust."),
                         ("3. Pull out the fields that matter",
                          "For every intent, there is a defined set of fields that the downstream action needs: PO number, model, serial, ship-to, dollar amount, partner name, work-order number, and so on. The solution extracts those fields from the email and its attachments with industry-leading OCR and document-intelligence services, falling back to vision-capable AI when the document is image-only.",
                          "The extracted fields appear next to the email, ready for CSR review if the case is held for approval."),
                         ("4. Find the right records",
                          "The solution looks up the customer in Salesforce by email and by name. It checks for an existing CCC Request against the PO or WO number. It matches the product to the catalogue. Where Salesforce does not yet hold the customer, it triggers the standard CMD activation request your team uses today.",
                          "The CSR sees the linked Salesforce records directly inside the case view."),
                         ("5. Decide how confident it is",
                          "Every request gets four independent confidence scores: did the AI classify it correctly, did it extract every required field, did it find the matching Salesforce record, and can the downstream action actually execute? This is the four-gate model your team described to us during the Q&A. The lowest gate decides the path: high overall confidence means act, mid means one-click approve, low means full CSR review.",
                          "The CSR sees the four scores broken out, so they know exactly why the AI is asking for help."),
                         ("6. Act and reply",
                          "On high confidence, the solution creates or updates the CCC Request, attaches the email to Docunet with Doc type FCNV, assigns the right owner per your routing rules, and drafts a customer-facing reply in the customer's language. On mid confidence it does all of that as a proposal and waits for one CSR click. On low confidence the full case sits in the review queue.",
                          "The customer hears back faster; the CSR works only on the cases that need a human."),
                         ("7. Learn",
                          "Every CSR correction is captured and replayed as a candidate update to your rule book or your translation glossary. A drift detector watches the classification mix over time and alerts the team if something is shifting. Your operations dashboard shows the automation rate, accuracy, and exception trends.",
                          "Your team's corrections compound. The solution gets better at its job week over week."),
                     ])

    _h(doc, "How the solution is built and where it runs (hybrid model)", 2)
    _p(doc,
       "We will follow a hybrid build-and-deploy model. The design and build phase runs inside "
       "ZBrain's Solution Builder — the visual orchestration surface that the Hackett team uses "
       "to wire the end-to-end agents (Intake, Extract, Reconcile, Decide, Execute, "
       "Communicate) against the knowledge bases, rule books, and integration channels "
       "described in this section. Solution Builder runs on the ZBrain platform tenancy The "
       "Hackett Group operates. Configuration work is driven jointly with Keysight's rule "
       "owners, not a code-from-scratch build.")
    _p(doc,
       "Once built, the solution deploys separately onto Keysight's own infrastructure: AWS, "
       "per the Keysight cloud standard called out in Q&A item 55. The deployed footprint "
       "carries the runtime versions of the same modules used during build — the Knowledge "
       "Base content, the Governance policies, the Monitor telemetry, and the agent runtime — "
       "all inside Keysight's VPC, behind Keysight's IAM, with the LLM endpoint bound to a "
       "private network path. Prompts, customer data, and audit trails never leave Keysight's "
       "boundary. Section 9 (Infrastructure) describes the deployment topology; Section 10 "
       "(Security and Governance) describes the runtime governance posture.")

    _h(doc, "ZBrain modules that this solution uses", 2)
    _kv_table(doc,
              [
                  ("Solution Builder",
                   "The orchestration surface used to build the end-to-end agents for the seven RFP use cases. Hackett's engineers configure the agents in Solution Builder against Keysight's rule book; promotion through dev, UAT, and production is handled by the same Solution Builder release flow."),
                  ("Knowledge Base module",
                   "Holds the intent definitions, routing rules, per-region operational overlays, magic-SKU and distributor tables, per-language glossaries, and reply templates. Editable by Keysight rule owners; re-read on every classification, so updates take effect without a developer release. The platform supports retrieval over structured grounding data, knowledge-graph traversal, and unstructured document search, per Keysight's Q&A item 53 (all three RAG types required)."),
                  ("Governance module",
                   "Application-level governance for this solution: role-based access, allow-listed tool calls per agent, prompt-injection defences, citizenship-based access controls per Q&A item 25, and policy enforcement at the agent boundary. Integration with Keysight's enterprise governance posture as it matures (today: Microsoft for SSO and RBAC, no enterprise governance fabric yet, per Q&A item 55)."),
                  ("Monitor module",
                   "Real-time observability of every agent: classification mix, confidence distribution, per-gate pass rate, HITL queue depth, per-intent SLA adherence. The Monitor module is where the continuous-learning workflow surfaces correction opportunities, drift signals, and A/B candidates for promotion. The detail is in Implementation Approach 4.6 and AI/ML Capabilities 5.5."),
              ])

    _h(doc, "How the solution fits into Keysight's existing systems", 2)
    _p(doc,
       "Per Keysight's Q&A items 2, 3, 4, 5, 6, 48, and 56, the integration surface for this "
       "engagement is the set of systems below. The solution is a well-behaved client to each; "
       "it does not replace any of them.")
    _kv_table(doc,
              [
                  ("Salesforce (single global instance)",
                   "Primary service platform per Q&A item 2. System of record for Account, Contact, Quote, Opportunity, CCC Request, custom Work Order application, and Order data. The solution reads and writes through standard Salesforce APIs."),
                  ("Oracle EBS 12.2",
                   "Order Management and Supply Chain per Q&A item 5. Reached exclusively through Jitterbit middleware per Q&A item 56. Direct DB access and Informatica are not used for vendor integrations to Oracle EBS. We use Jitterbit the same way the existing Q2O flow (Salesforce to Oracle, per Q&A item 10) already does."),
                  ("Jitterbit middleware",
                   "Keysight's existing integration platform for Oracle EBS and DocuNet. The solution uses Jitterbit as a service; it does not introduce a parallel middleware."),
                  ("Logik.io",
                   "CPQ engine for new-code sales per Q&A item 3. Managed package in Salesforce but runs on its own platform with its own APIs (does not use the Salesforce data model). Renewals stay in Salesforce; this solution treats renewal-quote POs as out of scope per Q&A item 7."),
                  ("KSP (Keysight Support Portal)",
                   "Custom Keysight portal for customers to manage assets, cases, and work orders. Read-only self-service surface. WO status replies include the standard KSP pointer Keysight's CSRs use today."),
                  ("KSM / DocuNet",
                   "Custom Keysight document store for POs, SOAs, and attachments. Reached via Jitterbit. Every email attached to a CCC Request is filed in DocuNet with Doc type FCNV per Keysight's existing convention."),
                  ("Inspire",
                   "Keysight's custom pricing engine, deployed in AWS. Referenced for context; not in the integration scope for the seven RFP use cases."),
                  ("ServiceNow",
                   "Keysight-owned reminder and follow-up workflow engine per Q&A item 4. Keysight's dev team owns the ServiceNow side. The solution integrates as a consumer only and does not replace ServiceNow."),
                  ("Microsoft Outlook",
                   "Inbound channel. Reads from every mailbox in scope (50 today per Q&A item 12, with consolidation in progress but not complete by project end), applies Keysight's existing Outlook rules first, and back-stamps each handled email."),
                  ("Microsoft Azure AD",
                   "Enterprise SSO and RBAC per Q&A item 55. Operator login to the HITL portal and Knowledge Base authoring screens federates through Azure AD."),
                  ("Cloudflare",
                   "Used by Keysight for any internet-exposed agents per Q&A item 55. The solution's external surfaces inherit this posture."),
                  ("Snowflake and BigQuery",
                   "Scattered enterprise data landscape per Q&A item 54. No formal data fabric in place today. The solution does not assume one; integration with whichever data-fabric direction Keysight chooses is described in our Product Features 4.3 response."),
                  ("Translation knowledge base",
                   "Keysight's existing translation asset is merged into the per-language glossary in the ZBrain Knowledge Base module at deployment, with SalesOps-specific term overrides added as needed per Q&A item 26."),
              ])

    _h(doc, "Why ZBrain is the right platform for this", 2)
    _p(doc,
       "ZBrain is the Hackett-owned enterprise Agentic AI platform, built for exactly this class "
       "of problem: high-volume, multi-system, judgement-driven workflows where the AI must work "
       "alongside human experts rather than replace them. Under the hybrid model proposed here, "
       "design and build use ZBrain Solution Builder on the Hackett-operated platform tenancy, "
       "and the production solution deploys onto Keysight's own AWS infrastructure (the Keysight "
       "cloud standard per Q&A item 55). The split addresses Keysight's Q&A item 48 preference "
       "for SaaS-grade build ergonomics together with Q&A item 31's question on deployment "
       "location: Keysight gets ZBrain's productivity at build time and full runtime ownership "
       "in production.")
    _p(doc,
       "The platform is model-agnostic (Keysight can choose and switch between leading "
       "commercial or open-source models without rebuilding the solution, per Q&A item 40), "
       "and configured rather than custom-coded for each new rule or intent. Knowledge Base "
       "content, Governance policies, and Monitor configurations authored in Solution Builder "
       "promote into Keysight's deployment through a single release pipeline. Operator-tunable "
       "Knowledge Bases mean Keysight's rule owners can update intents, routing rules, "
       "glossaries, and templates without a developer release cycle once the deployment is "
       "live in Keysight's infra.")

    return _save(doc)


# ──────────────────────────────────────────────────────────────────
# 3. Scope
# ──────────────────────────────────────────────────────────────────
def build_scope() -> bytes:
    doc = _new_doc()
    _title_block(doc, "Section 3 of 13", "Scope")

    _h(doc, "3.1 In scope: the seven use cases", 2)
    _p(doc,
       "The seven use case diagrams in the RFP define the inbound-request flows the solution "
       "handles end to end. Each one is a complete business outcome, not just a classification "
       "label. The solution covers the happy path plus the documented fallout and variant "
       "sub-paths drawn into each diagram.")

    _three_col_table(doc,
                     ("Use case", "What it covers today", "What the solution will do"),
                     [
                         ("Trade Order Entry (PO Received)",
                          "Manual classify, manual Salesforce entry, AIOA partial PO validation, Trade CSR drives Q2O, Oracle EBS booking through Jitterbit, SOA generated by hand and emailed to the customer.",
                          "The whole flow end to end. SOA generated automatically and attached against the CCC; high-confidence cases publish on their own, mid-confidence cases get a one-click CSR approval, low-confidence cases sit in HITL."),
                         ("Trade Sales Change Order",
                          "Manual classify, existing-CCC lookup, clone-for-Change-Order, delta-amount math, CSR updates Oracle through Jitterbit, manual customer notify.",
                          "Existing-CCC matrix automated across all ten statuses. Delta amount computed. Customer notification auto-drafted; CSR confirms with one click."),
                         ("SOM Work Order Automation (single and multi-asset)",
                          "Manual classify, CCC shell, multi-asset clone loop (one CCC per asset), SOM CSR creates each WO and attaches the email by hand.",
                          "Multi-asset fan-out automated for the consistent-info case; ambiguous multi-asset cases route to an SOM CSR. WO creation, owner assignment, and email-attach all proceed automatically. CSR sees only the flagged cases."),
                         ("SOM WO Update and Change Order (single and multi-asset)",
                          "Manual classify, match existing WO, AIOA validates the inbound PO, manual WO update and customer reply.",
                          "The AIOA validation pattern is preserved on the inbound PO. WO updates happen automatically. The customer reply is auto-drafted in the customer's language; CSR confirms when judgement is needed."),
                         ("SOM WO Status and Inquiry",
                          "Manual Salesforce lookup, manual reply with the standard KSP pointer, optional fallout CCC when the request cannot be answered cleanly.",
                          "Status lookup and customer-friendly reply auto-drafted. When the inquiry cannot be auto-answered (for example a cal-cert retrieval), the case is handed to a CSR with the context pre-populated."),
                         ("Service Contracts (CCC Request Creation)",
                          "Manual classify, CCC shell, AIOA validates any attached PO, S+R CSR works the queue.",
                          "Quote vs Order request routing handled at intent classification. AIOA pass moves the case forward; AIOA fail routes the case to the S+R CSR queue with the failure reason pre-populated."),
                         ("SSD Change Request",
                          "Manual classify, Post-Order-Booking CSR, factory-and-Oracle dashboard loop, manual customer notify on completion.",
                          "Trade Order Modification CCC auto-created with the right type and sub-type. Dashboard hooks integrate with the factory and Oracle steps. Customer notification drafted automatically when the CCC closes."),
                     ])

    _h(doc, "3.2 Sub-case coverage detail (per use case)", 2)
    _p(doc,
       "Each of the seven diagrams in the RFP bundles a happy path plus several fallout and "
       "variant sub-paths. The solution covers every sub-path. The list below is the catalogue, "
       "with a short note on how each is handled.")

    _p(doc, "Use case 1. Trade Order Entry (PO Received).", bold=True)
    _bullets(doc, [
        "Standard PO. The end-to-end happy path. SOA generated and published per the autonomy tier.",
        "PO accompanied by a quote (888 series). The existing Salesforce Opportunity is linked to the CCC during entity resolution.",
        "Standard-PO subtypes: Stock Rotation (quarterly partner), Rebates (monthly partner, negative order amount), eBiz (Keysight Used Equipment Store), SOW (Z-prefix SKU, EID, Custom Solutions), Prebuild, Amendment, Cancellation, Change Quantity, Duplicate PO, Confirm Orders. Each subtype carries its own field rules and routing convention; the solution applies them where the inbound mail matches.",
        "FCNV Fallout when extracted fields are incomplete. The case stays with FCNV until the missing fields are filled.",
        "Optional CSR review checkpoint after FCNV assignment. Configurable per region.",
        "AIOA Fallout when PO validation flags the PO data. The case enters the AI OA Fallout queue and returns to the pipeline after correction.",
        "Quote Update Fallout when the matched quote and the PO disagree on price, quantity, or terms. The CSR sees the field-level diff.",
        "Q2O Conversion Fallout when the Salesforce Quote-to-Order step itself errors. The CSR re-attempts with the failure logged for support.",
        "SOA review by the CSR before publish. The default tier for outbound SOAs is one-click CSR approval.",
    ])

    _p(doc, "Use case 2. Trade Sales Change Order.", bold=True)
    _bullets(doc, [
        "Change Order against an existing CLOSED CCC. The CCC is cloned with a delta amount (new PO total minus old PO total), the Type field is set to Change Order, currency is matched, the Final Destination country is checked.",
        "FCNV Fallout when automation cannot reach a required field. The CSR completes the entry; the AI prefills the rest.",
        "Bill-to and Ship-to as different entities. The Account on the CCC uses the Bill-to; Assignment Lookup uses the Ship-to City and State.",
        "CSR updates the Existing Order in Oracle through Jitterbit. Retained as a CSR action because the Oracle write requires judgement; the AI prepares the proposed update for one-click confirmation.",
        "Customer update. The outbound message is auto-drafted in the customer's detected language.",
        "Closure. Once the customer confirms, the CSR moves the CCC to Closed.",
    ])

    _p(doc, "Use case 3. SOM Work Order Automation (single and multi-asset).", bold=True)
    _bullets(doc, [
        "Single asset. One CCC, one WO created and assigned to the right SOM owner.",
        "Multi-asset, consistent information (same address, same customer). One CCC fans out to multiple WOs created from the same record.",
        "Multi-asset, inconsistent information (different addresses or different customer details). One CCC is assigned to an SOM CSR for manual fan-out, because automatic split is not safe.",
        "CMD lookup. Account and asset (Model and Serial) confirmed against CMD; missing master data triggers the standard CMD activation request your team uses today.",
        "System error during WO creation. The CCC reassigns to the SOM CSR queue with the failure cause logged.",
        "Manual SOM email-from-WO-back-to-email loop preserved as the exception path for routing corrections.",
        "Email filed in Docunet against the WO record with Doc type FCNV.",
        "CCC closed without a reply on successful automation, per Keysight's published convention.",
    ])

    _p(doc, "Use case 4. SOM WO Update, Change Order, and multi-asset variants.", bold=True)
    _bullets(doc, [
        "WO Update against a single existing WO. Note and task added; the existing WO is modified in place.",
        "Change Order against an existing WO. The WO is updated with the delta data; the AIOA validation pattern runs on the inbound PO.",
        "Multi-asset update with consistent information. One CCC updates each WO referenced.",
        "Multi-asset update with inconsistent information. The case falls out to an SOM CSR for manual update.",
        "Inbound PO validation following the AIOA pattern. Flagged items enter the AI OA Fallout queue.",
        "Manual SOM email-from-WO-back-to-email loop preserved for routing corrections.",
        "Customer reply path where the update warrants one (schedule change, scope change). One-click CSR approval in the default tier.",
    ])

    _p(doc, "Use case 5. SOM WO Status and Inquiry.", bold=True)
    _bullets(doc, [
        "Specific WO status request. Salesforce lookup of the named WO; reply drafted in customer-friendly language with the KSP pointer.",
        "General WO inquiry. WO details retrieved; reply drafted if the question is cleanly answerable; otherwise routed to a CSR.",
        "Cannot classify cleanly (minimum information to classify, no specific WO named). The CCC shell is created and assigned to FCNV for follow-up.",
        "Inquiry cannot be auto-replied (for example a cal-cert retrieval). The case is created and assigned to a CSR for full handling.",
        "Multiple WO references in one email. The reply contains a per-WO status block.",
    ])

    _p(doc, "Use case 6. Service Contracts CCC Request Creation.", bold=True)
    _bullets(doc, [
        "Support Agreement Quote (multiple sub-types). The CCC opens with the appropriate status and stage and routes to the S+R Quote queue.",
        "Order Request (various Agreement sub-types). Same flow with Type set to Order Request; AIOA-pattern PO validation invoked if a PO is attached.",
        "FCNV Fallout when fields are missing. The case moves to Awaiting status with the missing fields surfaced in the review queue.",
        "CTA Scope fallout. The CTA reviews the technical specification before the case proceeds.",
        "AIOA Pass. The case moves to Assigned with stage Automation Complete; downstream S+R fulfilment begins.",
        "AIOA Fail. The case moves to Assigned with stage Review Required; the S+R CSR reviews via the AI OA Fallout queue.",
        "Begin process. The CCC is picked up by the S+R team for fulfilment.",
        "Renewal flow. Where the inbound is an agreement renewal request, the existing contract is looked up; if found, status transition; if not, treated as a new Quote.",
    ])

    _p(doc, "Use case 7. SSD Change Request.", bold=True)
    _bullets(doc, [
        "Happy path. CCC auto-created with Request Type Trade Order Modification, sub-type SSD Change, owner mapped to the Sales Order Owner or Direct Inquiries in Oracle.",
        "FCNV Fallout when fields are missing. The case stays with FCNV until the extracted fields are complete.",
        "Add SSD request to the CSR dashboard. Notification dispatched to the CSR and the relevant factory.",
        "Factory prepares SSD and triggers the CSR from the dashboard. A human-in-loop bracket is preserved around the factory and Oracle interactions per the diagram.",
        "CCC interaction finalises SSD. The dashboard hook writes the change into Oracle through Jitterbit.",
        "Auto-close on completion. The CCC moves to Closed.",
        "Customer notification. The outbound message is auto-drafted on close; CSR one-click publish in the default tier.",
    ])

    _h(doc, "3.3 Out of scope", 2)
    _p(doc,
       "The list below pins what is explicitly out of scope for this engagement, with the "
       "Q&A clarification reference where one applies.")
    _bullets(doc, [
        "Renewal-quote PO workflow automation (per Q&A item 7). Email classification is in scope for all requests, but workflow automation for renewal-quote POs is not in this phase. Renewal-quote POs route to a manual CSR via a CCC request after classification.",
        "Q2O (Quote-to-Order) workflow automation in Oracle (per Q&A item 10). The Salesforce-to-Oracle Q2O path already runs through Jitterbit at Keysight today; the solution does not rebuild it. The solution does classify and prepare the Q2O case, then hands off to the existing Q2O middleware flow.",
        "ITAR, GovCloud, and FedRAMP-scoped deployments (per Q&A item 8). Out of scope for this phase. The solution must still apply citizenship-based routing on US-government-customer email at the mailbox door per Q&A item 25, so that traffic is shielded from non-US-citizen access.",
        "Channels other than email (per Q&A item 9). EDI, fax, and multi-channel order entry are handled separately at Keysight and are not in this RFP.",
        "Direct database access or Informatica integrations to Oracle EBS (per Q&A item 56). Oracle EBS is reached exclusively through Jitterbit.",
        "Building a new customer-facing portal. Keysight Support Portal stays as the customer's read-only self-service surface.",
        "Building a new ServiceNow workflow (per Q&A item 4). ServiceNow is Keysight-owned; the solution integrates as a consumer only.",
        "Replacing Salesforce or Oracle EBS. Both stay as systems of record; the solution is a well-behaved client.",
        "Handling the classified or government-controlled environment itself. The solution routes eligible emails to the existing KSO queue; what happens inside that environment is out of our scope.",
        "Low-volume sub-types Keysight has already chosen to defer (for example Consumption Billing).",
        "Replacing Microsoft Outlook as the mailbox layer. The solution reads from your existing mailboxes and back-stamps folders; the mailbox infrastructure itself is unchanged.",
    ])

    _h(doc, "3.4 Assumptions", 2)
    _bullets(doc, [
        "Salesforce is a single global instance with API access per Q&A item 6. Keysight provides OAuth credentials and the CCC Request custom-object schema.",
        "Oracle EBS 12.2 is reached exclusively through Jitterbit per Q&A item 56. Keysight provides Jitterbit endpoint definitions, authentication, and the field mappings already used between Salesforce and Oracle.",
        "Microsoft Outlook is reachable through standard IMAP and Microsoft Graph. The solution operates transparently across the current 50 mailboxes (per Q&A item 12), reading all of them and applying the appropriate rule set as consolidation proceeds.",
        "Keysight's existing translation knowledge base is available to be merged into the per-language glossary at deployment per Q&A item 26. Enterprise systems remain English; translation happens at the email-processing layer per Q&A item 27.",
        "The deployment topology mirrors the ZBrain reference architecture (Section 9): AWS-hosted services (the Keysight standard per Q&A item 55), Azure AD for SSO and RBAC per Q&A item 55, Cloudflare for any internet-exposed agents per Q&A item 55, Azure Open AI as the LLM service (Keysight has no in-house custom LLM per Q&A item 51), Jenkins plus Bitbucket as the build toolchain. The enterprise governance toolchain is Keysight's choice; today Keysight does not operate a Purview-equivalent enterprise governance fabric per Q&A item 55, so application-level governance ships with the ZBrain Governance module and integrates with enterprise tooling as Keysight adopts it.",
        "Volume baseline per Keysight's Q&A (item 11): approximately 880 thousand emails per year, around 2,000 per day, with quarter-end and year-end bursts. Single email payloads typically run one to twenty MB including attachments. Solution sizing is anchored to the upper end of that range.",
        "User concurrency baseline per Q&A item 17: 600 to 700 total users across global sales operations, with 80 to 90 concurrent on this workflow at any moment, ramping during quarter-end and year-end.",
        "Four regions in scope per Q&A item 23: Americas, Europe, Asia-Pac, Japan. Single global rule book is the target; regional overrides (including Japan-specific fields) carried as configuration overlays.",
        "Language coverage follows the RFP list with the duplicate Spanish entry removed per Q&A item 52 (English, Canadian French and standard French, Spanish, Brazilian Portuguese, Korean, Simplified Chinese, Traditional Chinese, Vietnamese, Dutch, Swedish, Danish, Finnish, German, Italian, Czech, Japanese). Regional-language breakdown to follow as a Keysight addendum.",
        "ITAR, GovCloud, and FedRAMP workloads are out of scope for this phase per Q&A item 8. Capability statement is included for future planning; the solution must still apply citizenship-based routing on US-government-customer email at the mailbox door (per Q&A item 25), so that traffic is shielded from non-US-citizen access at all reasoning, logging, and replay surfaces.",
    ])

    _h(doc, "3.5 Risks and how we manage them", 2)
    _three_col_table(doc,
                     ("Risk", "Likelihood and impact", "Mitigation"),
                     [
                         ("Classification drift as Keysight's email mix evolves.",
                          "Medium probability, high impact",
                          "A drift detector watches the rolling baseline and alerts the team. CSR corrections are captured and replayed as proposed rule updates, so the rule book stays current."),
                         ("Hard-to-parse attachments (Outlook items embedded as files, animated images, very long filenames).",
                          "High probability, medium impact",
                          "Outlook item attachments are unrolled to surface the inner email. Vision-capable AI handles image-only PDFs. Long filenames are normalised. Low-confidence extractions land in the human-in-the-loop queue with the source attachment shown."),
                         ("Multi-asset emails. One CCC per asset rather than one per email.",
                          "High probability, medium impact",
                          "Asset detection at extraction time; fan-out when the per-asset information is consistent; CSR hand-off when it is not. Mirrors the rule documented in your published process."),
                         ("Distributor routing mistakes (standard customer ordering distributor product).",
                          "Medium probability, high impact",
                          "Distributor list and magic-SKU resolver applied at routing time. CSR sees a clear flag if the account and SKU disagree."),
                         ("Mailbox consolidation timing during the transition from approximately 50 mailboxes to 1 or 2.",
                          "Low probability, medium impact",
                          "The solution is mailbox-agnostic. Classification and routing do not depend on the mailbox of origin. Region metadata is preserved per mailbox."),
                         ("Government and restricted-customer compliance.",
                          "Low probability, high impact",
                          "The KSO redirect fires before any AI reads the body. Compliance audit trail is retained immutably."),
                         ("LLM provider outage or rate limit.",
                          "Low probability, high impact",
                          "Model-agnostic configuration with primary-and-secondary failover. The deterministic Outlook rules and the human-in-the-loop queue remain available in degraded mode."),
                         ("Volume spikes at quarter-end and year-end.",
                          "Annual, high impact",
                          "Queue-based intake with horizontal worker scaling. The baseline is roughly 2,000 emails per day (Q&A item 11); stress-test target is 50x baseline (100,000 emails per day sustained for 24 hours) with the worker pool elastic-scaling on queue depth, so Q-end and year-end bursts absorb in the queue rather than back-pressure into the mailbox layer."),
                     ])

    return _save(doc)


# ──────────────────────────────────────────────────────────────────
# 4. Implementation Approach
# ──────────────────────────────────────────────────────────────────
def build_implementation_approach() -> bytes:
    doc = _new_doc()
    _title_block(doc, "Section 4 of 13", "Implementation Approach")

    _h(doc, "4.0 Our approach in one paragraph", 2)
    _p(doc,
       "This engagement begins from a configured baseline. The Hackett Group has built Version 1 "
       "of the solution on the ZBrain platform. Version 1 covers all seven RFP use cases and the "
       "documented sub-paths, configured against our reading of the AS-IS process. Version 1 is "
       "not the production solution; it is a working baseline that serves as the input to "
       "Requirement Finalization, not the output of MVP Build. Across the ten-week engagement, "
       "Version 1 is adapted to Keysight's specific inbox patterns, rule book, translation "
       "knowledge base, and operational data. From MVP Build onward the work is tuning a system "
       "that already exists, not building one from scratch. This is what allows delivery of a "
       "working production pilot in ten weeks, against a green-field timeline that typically "
       "requires sixteen to twenty.")

    _h(doc, "4.1 Methodology", 2)
    _p(doc,
       "Delivery follows Hackett's AI Development Methodology, applied on every ZBrain enterprise "
       "engagement. The methodology has five named phases: Functional Design, Visual and "
       "Technical Design, MVP Development, Testing and Deployment, and Continuous Improvement. "
       "These compress into a ten-week engagement structured as a four-week Requirement "
       "Finalization phase, a four-week MVP Build phase, and a two-week Testing and Tuning "
       "phase. Each phase concludes with a working increment and an acceptance gate signed off "
       "by Keysight's named rule owners.")

    _h(doc, "4.2 The ten-week journey", 2)
    _three_col_table(doc,
                     ("Phase", "Weeks", "What ships at the end"),
                     [
                         ("Requirement Finalization",
                          "Weeks 1 to 4",
                          "Four weeks of joint working sessions with Keysight's named rule owners. Every use case and sub-path is reviewed against Version 1, deltas against current operations are captured, and the build backlog is locked. Outputs are signed-off intent definitions, Outlook pre-screen rules, routing rules and operational rule book, per-intent extraction schemas and per-region overlays, and the integration boundary with Salesforce, Jitterbit (the channel to Oracle EBS and DocuNet), Microsoft Outlook through Microsoft Graph, and Azure AD. Dev environment provisioned in Keysight's AWS account. Version 1 executes against an initial slice of Keysight data with accuracy reported by intent, language, and region. The Functional Design document is signed by Keysight's named rule owners."),
                         ("MVP Build",
                          "Weeks 5 to 8",
                          "Four-week build phase that adapts Version 1 to Keysight's data and rule book across all seven RFP use cases end to end. Operator HITL portal, Knowledge Base authoring screens, and the four-gate confidence model are live. Existing-CCC matrix verified across all ten statuses. Routing resolver complete: distributor list, magic-SKU table, FE and CSR override detection, region-aware overlays, citizenship-based KSO routing. Customer reply drafting executes in the customer's detected language. Oracle EBS write-back through Jitterbit operates against Keysight's existing field mappings. ZBrain Governance and Monitor modules are wired into the Keysight deployment."),
                         ("Testing and Tuning",
                          "Weeks 9 to 10",
                          "Two weeks of tuning, stress, and acceptance. The classifier, extraction schemas, confidence thresholds, and routing rules are tuned against the labelled corpus using a 60 / 20 / 20 train, validation, test split. Inbound throughput is stress-tested at fifty times Keysight's stated 2,000-emails-per-day baseline to absorb quarter-end and year-end bursts. UAT runs against synthetic mail and shadow mode on a subset of real mail; Keysight rule owners sign off classification, extraction, routing, reply quality, and audit trail. Drift detector and continuous-learning feedback loop run live. Governance and Monitor hooks are verified against the scope agreed with Keysight's governance team. First mailbox cutover into production runs with the rollback plan tested. Engagement hands off to the tiered support and hypercare model."),
                     ])

    _h(doc, "4.3 Data partitioning: the 60 / 20 / 20 split", 2)
    _p(doc,
       "Tuning the classifier and the extraction pipeline against Keysight data converts "
       "Version 1 into a Keysight-ready MVP. The dataset is partitioned three ways, each slice "
       "with a defined purpose, and the three slices sum to one hundred percent. The data is "
       "drawn from Keysight's existing labelled corpus from the prior POC, together with any "
       "production-representative mail Keysight can release under its data-handling rules.")
    _three_col_table(doc,
                     ("Partition", "Share", "Purpose"),
                     [
                         ("Training set",
                          "60 percent",
                          "Calibrates the classifier, the per-intent extraction schemas, the four-gate confidence thresholds, the routing rules, and the translation glossary. This is the data the system learns from during Testing and Tuning."),
                         ("Validation set",
                          "20 percent",
                          "Held out from training. Used to tune hyperparameters and per-gate thresholds, measure per-intent and per-language accuracy after the first tuning pass, and verify the 90 percent classification threshold the RFP names. Errors on this set drive the first feedback cycle to Keysight's named rule owners."),
                         ("Held-out test set",
                          "20 percent",
                          "Reserved for the Testing and Tuning acceptance gate and re-run before each production cutover. Confirms the system generalises to unseen data and protects Keysight against a model that overfits the training and validation sets. Not used during training or validation."),
                     ])
    _p(doc,
       "The three partitions remain strictly separated. No data in the validation or test sets "
       "is used to train. This is standard machine-learning practice and protects Keysight "
       "against a system that performs well on training data but fails on real inbound mail.",
       italic=True)

    _h(doc, "4.4 Engagement cadence", 2)
    _kv_table(doc,
              [
                  ("Weekly update call",
                   "Joint Keysight and Hackett project call at a fixed weekly time. Status against plan, working-increment demonstration, risk and dependency review, decisions captured, and outstanding items refreshed in the shared backlog. Written minutes posted to the shared channel after each call."),
                  ("Phase end",
                   "Phase review with acceptance-gate sign-off by Keysight's named rule owners at the close of Requirement Finalization, MVP Build, and Testing and Tuning. Phase retrospective covering delivery against plan, risks, and the forward backlog."),
                  ("Monthly",
                   "Steering committee with Keysight executive sponsors and Hackett senior management. Risks, dependencies, and scope changes formally reviewed and minuted."),
                  ("Continuous",
                   "Shared backlog visible to both teams; shared monitoring dashboards; shared support channel for asynchronous queries and decisions between weekly calls."),
              ])

    _h(doc, "4.5 Phase detail", 2)

    _p(doc, "Requirement Finalization (Weeks 1 to 4).", bold=True)
    _bullets(doc, [
        "Review every one of the seven RFP use cases with Keysight's named rule owners. Lock the in-scope sub-cases and the out-of-scope items.",
        "Populate the ZBrain Knowledge Base: intent definitions and language coverage; Outlook pre-screen rules; routing rules including the distributor list and the magic-SKU table; the operational rule book the FCNV operators apply today; per-language glossaries imported from Keysight's existing translation knowledge base.",
        "Document the per-intent extraction schemas, per-region overlays, and the four-gate confidence model thresholds against the AS-IS pattern.",
        "Provision the dev environment in Keysight's AWS account: VPC, Azure AD application registration, LLM service binding, Salesforce sandbox connection, Jitterbit channel to Oracle EBS and DocuNet, Microsoft Graph permissions for Outlook, Cloudflare edge configuration, Bitbucket repository, Jenkins pipeline.",
        "Execute the initial classifier against the labelled corpus and report accuracy by intent, language, and region.",
        "Sign off the Functional Design document with Keysight's named rule owners. Outputs are the build backlog and the acceptance criteria for the MVP Build phase.",
    ])

    _p(doc, "MVP Build (Weeks 5 to 8).", bold=True)
    _bullets(doc, [
        "Build all seven RFP use cases end to end against the Salesforce sandbox, configured from Version 1 against Keysight's signed-off rule book.",
        "Existing-CCC matrix verified across all ten statuses. Delta-amount calculation for Change Order clones in place.",
        "Multi-asset fan-out logic in place for the consistent-information case. The inconsistent-information case routes to an SOM CSR by design.",
        "AIOA validation pattern preserved on Trade Order Entry, SOM WO Update, and Service Contracts flows for attached POs.",
        "Routing resolver complete: distributor list, magic-SKU table, FE and CSR override detection, region-aware overlays, citizenship-based KSO routing.",
        "Operator HITL portal live: case list, case detail, per-gate confidence display, extracted fields, drafted reply, and approve / edit / reject actions.",
        "Knowledge Base authoring screens live so Keysight's rule owners can edit intents, routing rules, glossaries, and the operational rule book without a developer release.",
        "Communication agent live: reply drafting in the customer's detected language, SOA and acknowledgement document generation, CommunicationLog with full audit trail.",
        "Oracle EBS write-back through Jitterbit operates against Keysight's existing field mappings.",
        "ZBrain Governance and Monitor modules wired into the Keysight deployment, with policy and telemetry scope agreed at Functional Design.",
    ])

    _p(doc, "Testing and Tuning (Weeks 9 to 10).", bold=True)
    _bullets(doc, [
        "Tune the classifier, extraction schemas, confidence thresholds, and routing rules against the labelled corpus using the 60 / 20 / 20 train, validation, test split.",
        "Stress-test inbound throughput at fifty times Keysight's stated 2,000-emails-per-day baseline (sustained 100,000 emails per day for 24 hours, with elastic worker scaling on queue depth) to absorb quarter-end and year-end bursts without back-pressure into the mailbox layer.",
        "Run UAT against synthetic mail and shadow mode on a subset of real mail, covering classification, extraction, routing, reply quality, and audit trail. Keysight's named rule owners sign off each category.",
        "Activate the drift detector and the continuous-learning feedback loop. Every CSR correction generates a candidate rule update for review in the Governance and Analytics screen.",
        "Verify Governance and Monitor hooks against the scope agreed with Keysight's governance team at Functional Design.",
        "Execute the first mailbox cutover into production. Validate the rollback plan by selectively pausing the cutover mailbox.",
        "Hand off to the L1, L2, and L3 support tiers. Publish the run-book. Deliver CSR training.",
    ])

    _h(doc, "4.6 Continuous learning", 2)
    _p(doc,
       "Continuous learning is a first-class operational workflow that executes inside the "
       "ZBrain Monitor and Knowledge Base modules from the day of first mailbox cutover. Its "
       "purpose is to convert every CSR interaction and every data signal into an auditable "
       "improvement to the live system on a cycle measured in days. The stated Keysight "
       "principle is that one-hundred-percent automation is not the goal; the goal is that "
       "confidence in each gate is validated and autonomy is expanded use case by use case as "
       "the platform learns from corrections.")

    _p(doc, "Loop A. Signal capture (always-on).", bold=True)
    _bullets(doc, [
        "Every CSR action is captured with a structured diff: re-classifications, edited extraction fields, rejected drafts, manual routing overrides, force-publishes, and force-holds. The diff carries the original solution output, the CSR output, the source email and attachments, and the four-gate scores at decision time.",
        "Every external write (Salesforce, the Jitterbit calls to Oracle EBS and DocuNet, the customer reply send) is captured with its outcome (success, retry, hard fail) and any downstream consequence (CCC reopened, customer follow-up, AIOA fallout).",
        "Every drift signal from the Monitor module is captured: classification-mix shifts, confidence-distribution shifts, HITL-rate shifts, per-language accuracy shifts, per-intent SLA breaches.",
        "All three signal streams land in a single Learning Store inside ZBrain, indexed by email reference, intent, language, region, mailbox, and use case. Nothing is thrown away.",
    ])

    _p(doc, "Loop B. Opportunity identification (weekly).", bold=True)
    _bullets(doc, [
        "A weekly batch job clusters signals into Opportunities, each one carrying a fingerprint (\"PO emails from distributor X in Brazilian Portuguese are misclassifying as KSO 8 percent of the time\") and a proposed remedy (\"add this sender pattern as a distributor route exception\" or \"add these three glossary entries\" or \"raise the Extraction gate threshold for this intent\").",
        "Opportunities are scored on expected lift (how many emails per week the fix would affect), effort (configuration change vs new agent action vs full rule-book entry), and risk (will this change push high-confidence acts into the wrong cell). The list is ranked by lift-over-effort with risk as a brake.",
        "Opportunities surface in the ZBrain Governance and Analytics screen. Keysight's named rule owners review them on the weekly update call. Each opportunity is accepted, rejected, or deferred, with the reasoning logged.",
    ])

    _p(doc, "Loop C. A/B promotion (per accepted opportunity).", bold=True)
    _bullets(doc, [
        "An accepted opportunity is implemented as a shadow change: the candidate rule, threshold, glossary entry, or KB exemplar set is deployed alongside the current production version, not in place of it. Inbound emails run through both paths in parallel.",
        "Comparison mode logs the two outputs side by side. The candidate's output is never acted on, never sent to a customer, never written to Salesforce or Oracle. It is observed only.",
        "The A/B horizon is the smaller of (i) a fixed time window (typically one to two weeks) and (ii) a fixed sample size (typically 1,000 to 5,000 in-scope emails). The horizon is set when the candidate is promoted to shadow and is locked once it starts.",
        "Pre-defined success criteria gate the promotion: accuracy lift over the baseline at a defined confidence interval, no regression in any tracked sub-segment (per-region, per-language, per-intent), and no spike in the HITL rate. If all three pass, the candidate is promoted to production by a Keysight rule owner in the governance screen. If any fail, the candidate is retired and the reason is recorded.",
        "Promotion is reversible. Every promoted change carries a one-click rollback to the prior version. The rollback diffs the live state against the prior state and is logged.",
    ])

    _p(doc, "Loop D. Drift detection and circuit-breaker (continuous).", bold=True)
    _bullets(doc, [
        "Independent of the weekly cycle, a real-time drift detector monitors classification mix, confidence distribution, HITL rate, per-intent SLA adherence, and per-language extraction accuracy against a rolling baseline (default 30-day window).",
        "On a material deviation, the dashboard alerts the on-call team. If the deviation crosses a defined SLO floor, the agent's auto-action authority is paused for the affected segment and traffic routes to CSR review until a rule owner clears the issue. The same circuit-breaker contract is declared at deployment, enforced by the Governance module, and logged in the audit trail.",
        "Drift events also feed back into Loop B as candidate Opportunities, so the recovery action is captured for future use.",
    ])

    _p(doc, "Loop E. Operator-tunable knowledge bases (continuous).", bold=True)
    _bullets(doc, [
        "Routine rule changes are operator edits in the ZBrain Knowledge Base module, not developer releases. Intent definitions, routing rules, distributor lists, magic-SKU tables, per-region overlays, and translation glossary entries are all editable.",
        "Edits carry version history, named-owner attribution, and one-click rollback. A change-control record is created automatically; in the governance call the rule owners review the week's edits as a batch.",
    ])

    _h(doc, "4.7 Governance and Analytics screen", 2)
    _p(doc,
       "The Governance and Analytics screen is the operational home for continuous learning. It "
       "is a single surface inside the ZBrain platform UI that the Keysight governance team and "
       "rule owners use to run the weekly cycle and to monitor live health between cycles.")
    _bullets(doc, [
        "Health tiles: classification accuracy (live and 30-day rolling), automation rate (L4 / L3 / L2 split), HITL queue depth, per-intent SLA adherence, four-gate pass rates. Each tile is clickable for the underlying breakdown.",
        "Opportunity board: the ranked list from Loop B above. Each opportunity card shows the fingerprint, the proposed remedy, the expected lift, the risk note, the affected segment, and the accept / defer / reject actions. The card carries through to the A/B configuration screen when accepted.",
        "A/B experiments: every active shadow comparison with its horizon, sample count to date, accuracy delta with confidence interval, sub-segment regression flags, and the promotion gate status. Promotion is a single click for an authorised rule owner.",
        "Drift timeline: a rolling chart of every drift signal raised, with the alerted SLO and the resolution action.",
        "Change log: every promoted change, every Knowledge Base edit, every rule-owner sign-off, every rollback. Filterable by use case, by author, by date range. Exportable to Keysight's audit retention.",
        "Per-CSR view (for governance, not performance management): correction patterns by CSR, so the team can see whether a specific instruction in the rule book is creating recurring corrections. The intent is to fix the rule book, not to grade individual CSRs.",
    ])

    _h(doc, "4.8 Knowledge transfer", 2)
    _bullets(doc, [
        "Operator training for CSRs on the HITL portal: reading confidence scores, when to edit, when to reject, and how the solution incorporates corrections.",
        "Administrator training for the team that operates the solution day to day: editing knowledge bases, adding a new intent, updating routing rules, and publishing a rule-book change.",
        "Builder training for the team that extends the solution into new use cases post-MVP. The skills profile required to maintain the solution after go-live is a business analyst plus a low-code builder for routine changes (Knowledge Base edits, routing rules, glossary entries, intent definitions). A full-stack developer is needed only for new integration channels or sub-systems outside the seven RFP use cases.",
        "Run-book covering monitoring, alerts, escalation paths, the support model, and the rollback procedure for each mailbox.",
        "Documentation maintained inside Keysight's preferred knowledge repository, kept current across the engagement and into steady-state operation.",
    ])

    _h(doc, "4.9 Risk management", 2)
    _p(doc,
       "Risk is managed actively, not retrospectively. The risk register is reviewed on every "
       "weekly update call, expanded at every phase retrospective, and presented to the monthly "
       "steering committee. Each risk carries an owner, a likelihood, an impact, and a mitigation "
       "action. Where a risk crosses a defined threshold, the corresponding circuit-breaker "
       "policy activates and the team is notified.")

    _h(doc, "4.10 Post-deployment lifecycle", 2)
    _p(doc,
       "The engagement does not end at production cutover. The lifecycle below covers the period "
       "from first mailbox cutover through long-running steady-state operation. Implementation "
       "work is fixed-fee against the per-phase acceptance gates. Post-deployment work is time "
       "and materials against a named rate card, so Keysight pays only for the engineering "
       "actually consumed.")
    _three_col_table(doc,
                     ("Phase", "Window", "Coverage and commercial basis"),
                     [
                         ("Hypercare",
                          "Weeks 11 to 14 (first three months post-cutover)",
                          "Included in the implementation price; no separate fee. A focused subset of the delivery team remains engaged to absorb first-wave drift, close first-round CSR feedback, and onboard the remaining mailboxes against the rollback plan."),
                         ("Steady-state managed service",
                          "From Week 15 onward",
                          "Tier 1, Tier 2, and Tier 3 support, priced per deployed agent per month. Continuous Learning loops execute inside this phase. Includes the weekly update call, the monthly steering review, and a quarterly business review."),
                         ("Enhancements (T&M)",
                          "Continuous, on demand",
                          "New use cases beyond the RFP seven, new integration channels, new languages, model migrations, additional mailbox cutovers, and optimisation work surfaced by the Opportunity board. Engineering and AI / ML effort delivered on a time-and-materials basis at the named rate card, against agreed scope and acceptance criteria per request."),
                         ("Major version uplift",
                          "On platform releases",
                          "ZBrain platform major releases are tested in Keysight's UAT environment, regression-run against the held-out test set, and promoted to production with a rollback gate. Included in the platform subscription; T&M applies only where Keysight requests bespoke work alongside the uplift."),
                         ("Continuous improvement programme",
                          "Quarterly",
                          "A quarterly improvement plan based on the Opportunity board, signed off by the Keysight sponsor and delivered against a Quarterly Statement of Work. Scope is the highest-lift opportunities surfaced by Loops A and B. Priced T&M against the rate card."),
                     ])
    _p(doc,
       "T&M engagements draw against a published rate card by role (Solution Architect, AI / ML "
       "Lead, AI / ML Engineer, Integration Engineer, Backend Engineer, Frontend Engineer, QA, "
       "Project Manager). Each engagement begins with a written scope and acceptance criteria, "
       "tracks against a target hour estimate, and reports weekly burn. Keysight retains the "
       "right to pause, adjust scope, or terminate at the next weekly update call.",
       muted=True)

    return _save(doc)


# ──────────────────────────────────────────────────────────────────
# 5. AI/ML Capabilities
# ──────────────────────────────────────────────────────────────────
def build_aiml_capabilities() -> bytes:
    doc = _new_doc()
    _title_block(doc, "Section 5 of 13", "AI / ML Capabilities")

    _h(doc, "Our capability", 2)
    _p(doc,
       "We are The Hackett Group. We bring the methodology, the proprietary platform, and the "
       "engineering team to this engagement under a single firm. The capability you see here is "
       "not assembled from partners; it is one organisation's combined offering. The three "
       "pillars below set out what we put on the table.")

    _p(doc, "The Hackett AI Development Methodology.", bold=True)
    _p(doc,
       "Our methodology for AI engagements: Functional Design, Visual and Technical Design, MVP "
       "Development, Testing and Deployment, Continuous Improvement. The methodology is what we "
       "run this engagement against, and is the reason we can compress to eight weeks.")

    _p(doc, "ZBrain, our proprietary enterprise Agentic AI platform.", bold=True)
    _p(doc,
       "ZBrain is the Hackett platform we will build the solution on. It carries the AI "
       "primitives as configurable building blocks: knowledge-base authoring, multilingual "
       "classification, document understanding, confidence-gated decisioning, monitoring, "
       "guardrails, observability. These are platform capabilities we configure for Keysight, "
       "not code we write from scratch.")

    _p(doc, "Engineering continuity from Keysight's prior POC.", bold=True)
    _p(doc,
       "The Hackett engineers who delivered Keysight's prior POC are part of the proposed MVP "
       "delivery team. The institutional context on Keysight's CSR rule book, the AIOA "
       "validation pattern, and the multi-format attachment handling comes with them into the "
       "MVP build.")

    _p(doc,
       "The subsections below describe four specific AI capabilities the solution will deliver "
       "for Keysight's seven RFP use cases: classification, extraction, confidence scoring, and "
       "translation. They map onto the RFP's Core Functional Requirements (Email classification, "
       "Document Intelligence & Data Extraction, Confidence-Gated Decision Engine, Customer "
       "Communication). The broader catalogue of ZBrain platform capabilities sits in our "
       "response to the RFP Product Features sheet (Section 2).")
    _p(doc,
       "Each subsection follows the same shape: what needs to be true for Keysight, how we "
       "approach building it, how we measure quality, and the evidence that gives us confidence "
       "we can deliver. Implementation choices that are best deferred to Functional Design "
       "(specific model providers, specific OCR services, specific framework selections) are "
       "called out as such; locking those at proposal stage would be premature.")

    # ──────────────────────────────────────
    _h(doc, "5.1 Classification", 2)

    _p(doc, "What needs to be true.", bold=True)
    _p(doc,
       "The classifier has to distinguish, with high accuracy, between the seven RFP-named use "
       "cases and the operational triage classes Keysight already uses (KSO redirect, Brazil Tax, "
       "Portal Admin, Collections, Auto-Reply, Undeliverable, internal Keysight forwards, "
       "Others). It has to apply the same rules your CSRs apply in their heads today: a PO "
       "mentioned alongside a Work Order is a service request not a sales PO; an "
       "acknowledgement-only thread is not a new PO; a credit-card mention without a valid PO "
       "number is not promoted to Sales PO; an 'FYI' that hides the real request earlier in the "
       "thread is walked back through the thread. It has to do this in sixteen languages, across "
       "approximately 50 mailboxes, at quarter-end burst volume.")

    _p(doc, "How we approach building it.", bold=True)
    _bullets(doc, [
        "Knowledge-base authoring of intent definitions. Each intent carries its own keywords, sender patterns, examples, exceptions, and per-region applicability. The knowledge base is editable by Keysight rule owners; the agent re-reads it on every classification, so updates take effect immediately without a developer release cycle.",
        "A two-pass design. The first pass reads the email in full context and proposes an intent. The second pass cross-checks the proposal against Keysight's operational rules and either confirms it or overrides it with the reason recorded. Both passes are logged for audit. This two-pass pattern is what your existing POC uses.",
        "Per-source language detection, so an attached PO in Spanish does not pull the body classification off track. Each source carries its own language tag through the pipeline.",
        "Thread-fragment handling, so empty forwards, CAUTION banners, and quoted-only replies do not collapse a classification. The agent walks back through the thread until it finds genuine business content.",
        "CSR-instruction override detection, so a Keysight staff member's routing direction inside the email body (for example, 'please route to the SOW team') supersedes the default routing.",
        "Drift detection on a rolling baseline of classification mix and confidence distribution, with alerts on material deviation.",
    ])

    _p(doc, "How we measure quality.", bold=True)
    _bullets(doc, [
        "A regression corpus drawn from Keysight's existing labelled mail, run at every sprint cut. Acceptance is the 90 percent threshold the RFP names, with accuracy reported per intent, per language, and per region.",
        "Confidence-binned accuracy, so we see whether errors cluster at the high-confidence end (a serious issue) or the low-confidence end (the human-in-the-loop queue is doing its job).",
        "Per-CSR-correction tracking, so we can see whether a specific rule keeps surfacing edits and prioritise the rule-book fix.",
    ])

    _p(doc, "What gives us confidence we can deliver this.", bold=True)
    _p(doc,
       "The ZBrain platform's classification primitives are in production across enterprise "
       "email-automation use cases comparable to Keysight's. The two-pass pattern was validated "
       "in your prior POC; the engineers who carried that work are part of the proposed MVP "
       "delivery team, so the classification logic does not have to be re-derived from scratch. "
       "The knowledge-base authoring surface is a first-class platform capability, not a custom "
       "build for this engagement.")

    # ──────────────────────────────────────
    _h(doc, "5.2 Extraction and Document Intelligence", 2)

    _p(doc, "What needs to be true.", bold=True)
    _p(doc,
       "Every RFP use case needs a different set of fields. A PO Received case needs PO number, "
       "model, serial, ship-to, dollar amount, currency, Final Destination country, and quote "
       "reference. A Work Order case needs WO number plus the Model and Serial pairs (often "
       "multiple per email). A Service Contract case needs agreement type, sub-type, and term. "
       "The agent has to extract these reliably from the document formats your customers actually "
       "send: PDFs, Word, Excel (including the Trade Credit RMU and Summary tab patterns used "
       "for partner Rebates), scanned images, HTML, plain text, and embedded Outlook items.")

    _p(doc, "How we approach building it.", bold=True)
    _bullets(doc, [
        "Per-intent extraction schemas held in the knowledge base. Each use case carries the fields it needs, the type and format rules for each field, and any per-region overrides. Rule owners edit the schema; no code change is required to add a new field.",
        "Layered document intelligence. Industry-leading OCR and form-recognition services handle structured documents; vision-capable AI handles image-only content; native parsers handle Word, Excel, HTML, and plain text. Embedded Outlook items are unrolled so the inner email and its own attachments are surfaced.",
        "Multi-asset detection. Where the email contains multiple Model and Serial pairs, the agent extracts each pair and prepares the fan-out to one CCC per asset, mirroring Keysight's published rule.",
        "Magic-SKU detection that respects Keysight's existing conventions: CUSTOM PRODUCT for unresolved SKUs, SOWDUMMY for Statement-of-Work routing, EXPORTDUMMY for non-US destinations.",
        "Entity resolution into Salesforce. Sender email to Contact, name lookup fallback, PO and WO numbers to existing CCC. Where Salesforce does not yet hold the customer, the agent triggers the standard CMD activation request your team uses today.",
    ])

    _p(doc, "How we measure quality.", bold=True)
    _bullets(doc, [
        "Per-field accuracy on a labelled corpus, broken out by attachment type so we can see where the failure modes cluster.",
        "Per-language extraction accuracy, since OCR providers vary by language and we need to know it before production.",
        "Edge-case audit: a sample of complex emails (multi-asset, multi-language, deeply nested forwards, scanned-only attachments) is reviewed manually every sprint to confirm extraction quality at the long tail.",
    ])

    _p(doc, "What gives us confidence we can deliver this.", bold=True)
    _p(doc,
       "Multi-format document intelligence is a core ZBrain capability, deployed across "
       "enterprise customers handling attachment complexity comparable to Keysight's. The Rebate "
       "Trade Credit RMU pattern, the multi-asset fan-out rule, and the magic-SKU conventions "
       "were validated in your prior POC. Selecting the specific OCR provider is a Functional "
       "Design decision in Sprint 1; we keep that choice open so it can be made against Keysight's "
       "enterprise standards and the latest model landscape at delivery time.")

    # ──────────────────────────────────────
    _h(doc, "5.3 Confidence scoring", 2)

    _p(doc, "What needs to be true.", bold=True)
    _p(doc,
       "The Q&A call clarified that confidence is not a single weighted number. It is four "
       "independent per-transaction gates: Classification (did the agent identify the intent "
       "correctly?), Extraction (did the agent extract every required field with the right "
       "format and range?), Entity Resolution (did the agent find the matching Salesforce "
       "record?), and Action Feasibility (can the downstream action actually execute?). The "
       "autonomy tier follows directly from the lowest gate.")

    _three_col_table(doc,
                     ("Gate", "What it answers", "Pass criterion"),
                     [
                         ("Classification confidence",
                          "Did the agent classify the request correctly against Keysight's rule book?",
                          "Both classification passes agree; the proposed intent is consistent with the rules."),
                         ("Extraction confidence",
                          "Did the agent extract every field the downstream action needs?",
                          "All schema-required fields are populated and pass format and range validation."),
                         ("Entity Resolution confidence",
                          "Did the agent find the matching Salesforce records (Contact, Account, existing CCC, Quote, Order)?",
                          "Binary: the records are found or they are not."),
                         ("Action Feasibility confidence",
                          "Can the downstream write (CCC create, Order update, Docunet filing, customer reply) actually execute with what was resolved?",
                          "All Salesforce and Oracle write-path field requirements are satisfiable."),
                     ])

    _p(doc, "How we approach building it.", bold=True)
    _bullets(doc, [
        "Composite confidence is min(gates), not a weighted average. A weak link anywhere blocks autonomous action. This mirrors how a careful CSR thinks: if any one piece does not fit, the whole transaction is held.",
        "Per-gate scores are surfaced to the CSR in the human-in-the-loop view, so the team knows precisely which gate is the limiting factor.",
        "CSR override is always available. A CSR can force any case to full review regardless of the score; the override is captured as feedback into the rule book.",
        "Tier thresholds (the score boundaries for high, mid, and low confidence) are configurable per use case, so Keysight can tune autonomy by business risk.",
        "A drift detector watches the confidence distribution. If confidence-binning shifts materially over the rolling baseline, the operations team is alerted.",
    ])

    _p(doc, "How we measure quality.", bold=True)
    _bullets(doc, [
        "Per-tier post-hoc accuracy: actions taken at the high-confidence tier are sampled and reviewed; the acceptance bar on that sample is agreed with Keysight rule owners at Functional Design.",
        "Mid-tier approval rate: the share of mid-confidence cases the CSR approves without edit. A low approval rate signals that the mid-tier threshold is too generous.",
        "Low-tier turnaround: time from low-confidence surface to operator resolve, against the service level agreed with Keysight at Functional Design.",
    ])

    _p(doc, "What gives us confidence we can deliver this.", bold=True)
    _p(doc,
       "The four-gate model is consistent with how Keysight's POC scores classifier output; the "
       "extension to the other three gates is straightforward against the platform's monitor "
       "primitives. Tier-based autonomy is configuration, not code. The per-gate display in the "
       "operator portal is a known ZBrain pattern, not a bespoke build.")

    # ──────────────────────────────────────
    _h(doc, "5.4 Translation", 2)

    _p(doc, "What needs to be true.", bold=True)
    _p(doc,
       "Inbound customer email arrives in many languages; Keysight's internal systems are not "
       "standardised to English. The sixteen languages named in the RFP must be supported at "
       "deployment: English, French including Canadian French, Spanish, Brazilian Portuguese, "
       "Korean, Simplified Chinese, Traditional Chinese, Vietnamese, Dutch, Swedish, Danish, "
       "Finnish, German, Italian, Czech, and Japanese.")

    _p(doc, "How we approach building it.", bold=True)
    _bullets(doc, [
        "Per-source language detection. The body and each attachment are detected independently so that an English cover email with a Spanish PO is handled correctly.",
        "Multilingual reasoning. The agent reasons against a normalised internal representation; the operational rules apply consistently regardless of source language.",
        "Per-language glossary held in the knowledge base. Keysight's existing translation knowledge base is merged in at deployment; SalesOps-specific terminology is added; rule owners can edit per-term mappings without code.",
        "Reply drafting in the customer's detected language. On high-confidence cases the reply publishes automatically; on mid and low confidence a CSR reviews before publish.",
    ])

    _p(doc, "How we measure quality.", bold=True)
    _bullets(doc, [
        "Bilingual reply review at the CSR step in early sprints, especially for the higher-volume non-English languages.",
        "Sampled native-speaker review for languages where Keysight CSRs include fluent speakers.",
        "Glossary coverage measured against actual term occurrences in the inbound corpus, so the team sees which SalesOps-specific terms are still being translated generically.",
    ])

    _p(doc, "What gives us confidence we can deliver this.", bold=True)
    _p(doc,
       "Multilingual document understanding and reply drafting are first-class ZBrain platform "
       "capabilities. The major frontier-model families all carry production-grade coverage of "
       "the RFP language list. The glossary-merge pattern is a standard configuration step. "
       "Keysight already operates with translation reuse internally, so the operational pattern "
       "is familiar to the team.")

    # ──────────────────────────────────────
    _h(doc, "5.5 Continuous improvement", 2)

    _p(doc, "What needs to be true.", bold=True)
    _p(doc,
       "The agent has to get better at its job over time as Keysight's team corrects it, your "
       "inbound mix evolves, and new edge cases emerge. The correction loop has to be measured "
       "in sprints, not quarters.")

    _p(doc, "How we approach building it.", bold=True)
    _bullets(doc, [
        "Every CSR correction (re-classification, edited field, rejected reply, manual override) is captured with a before-and-after diff against the original AI output.",
        "A weekly batch surfaces the most common corrections as proposed updates to the knowledge base. Keysight rule owners review the proposals and decide whether to accept each one.",
        "An optional comparison mode runs a candidate rule update in parallel for A/B testing before promotion. The candidate output is logged but never acted on until Keysight promotes it.",
        "Drift detection ties the loop together. Material shifts in classification mix, confidence distribution, or human-in-the-loop rate are flagged within the rolling baseline window.",
        "Operator-tunable knowledge bases mean routine rule updates take minutes, not a release cycle.",
    ])

    _p(doc, "How we measure quality.", bold=True)
    _bullets(doc, [
        "Correction recurrence rate: the percentage of CSR corrections that repeat in the next two weeks. A high recurrence indicates the rule book needs an update that has not yet been made.",
        "Time-to-rule-fix: from CSR correction to rule-book update accepted by a Keysight rule owner. Target is one sprint cycle.",
        "Drift-to-detect time: from a material distribution shift to alert raised. Target is within the rolling baseline window of seven days.",
    ])

    _p(doc, "What gives us confidence we can deliver this.", bold=True)
    _p(doc,
       "The feedback-to-rule-update loop is wired into the ZBrain platform's Knowledge Base and "
       "Monitor modules. The prior POC used the same pattern at smaller scope; the MVP extends "
       "it across the full seven-use-case set with the drift detector and the optional "
       "comparison-mode pass layered on. The full lifecycle is described in Implementation "
       "Approach 4.6 and 4.7.")

    # ──────────────────────────────────────
    _h(doc, "5.6 Retrieval-Augmented Generation (RAG)", 2)

    _p(doc, "What needs to be true.", bold=True)
    _p(doc,
       "Q&A item 53 confirms that all three RAG types are required: structured database "
       "grounding, knowledge-graph traversal, and unstructured document search. The Knowledge "
       "Base module inside ZBrain is the single surface that supports all three retrieval "
       "patterns for this solution; each agent declares which retrieval types it needs and the "
       "platform binds them automatically. Keysight's rule owners author and curate the KB "
       "directly; no custom retrieval code is written per agent.")

    _three_col_table(doc,
                     ("RAG type", "How this solution uses it", "Where it is served from"),
                     [
                         ("Structured database grounding",
                          "Entity resolution into Salesforce: customer, contact, account, quote, existing CCC, order. PO line items reconciled against the matched quote. Asset and Serial validation against CMD records. Order status retrieval for WO Status and Inquiry use cases.",
                          "Served by the ZBrain Knowledge Base module as a structured-grounding retriever bound to live Salesforce APIs (and to the Jitterbit path for Oracle EBS data per Q&A item 56), with a read-through cache. Rule owners edit the bound field mappings inside the KB."),
                         ("Knowledge-graph traversal",
                          "Walking the relationships Keysight rule owners encode in the Knowledge Base: distributor list to magic-SKU mapping, Bill-to versus Ship-to account graph for routing, region overlays linked to mailbox metadata, intent ontology with parent and child relationships. Used in routing decisions and in multi-hop entity resolution.",
                          "Served by the ZBrain Knowledge Base module as a graph retriever. The graph itself is authored and curated by Keysight rule owners; promotion of changes flows through the Governance and Analytics screen described in 4.7."),
                         ("Unstructured document search",
                          "Per-language glossaries, intent-definition exemplars, historical reply templates, and the operational rule book itself. Retrieval grounds the classifier and the reply drafter against Keysight-authored content rather than the model's pre-training corpus.",
                          "Served by the ZBrain Knowledge Base module as a document retriever, with per-language indexing and per-region overlays. Keysight's existing translation KB is imported at deployment per Q&A item 26."),
                     ])

    _p(doc, "How we approach building it.", bold=True)
    _bullets(doc, [
        "All three retrieval surfaces are first-class features of the ZBrain Knowledge Base module. The solution does not stand up a parallel retrieval stack; it configures the KB module against Keysight's data and rule book.",
        "Each agent declares the RAG types it needs in its Solution Builder definition. The platform binds the agent to the configured KB retrievers automatically.",
        "Per Q&A item 26, Keysight's existing translation knowledge base is reusable for the language-translation use cases. It is imported into the ZBrain KB module at deployment and the unstructured-document RAG path is bound to it directly.",
        "All retrievals are cached at the per-email reference, so an audit pull replays the exact retrieval context that informed the decision. This is critical for the audit trail described in Section 10.5.",
        "Retrieval quality is monitored as a first-class metric in the Monitor module: per-query relevance, retrieval-grounded answer rate, and grounded-versus-hallucinated output rate.",
    ])

    # ──────────────────────────────────────
    _h(doc, "5.7 Agent-level RLHF (feedback to learning)", 2)

    _p(doc, "What needs to be true.", bold=True)
    _p(doc,
       "Per Q&A item 50, Keysight's primary interest is agent-level feedback: correcting a "
       "misclassification or a bad extraction so the agent improves on similar emails. This is "
       "the path that produces the largest day-to-day quality lift, and it is the path this "
       "solution invests in. Improvement happens at the agent and Knowledge Base layer, where "
       "Keysight's rule owners can see, review, and promote every change.")

    _p(doc, "How we approach building it.", bold=True)
    _bullets(doc, [
        "Every CSR correction is captured into the Learning Store described in 4.6 Loop A. The diff carries the original AI output, the CSR output, the source email and attachments, and the four-gate scores at decision time.",
        "The weekly Opportunity batch (4.6 Loop B) clusters corrections into proposed Knowledge Base updates: a new sender pattern, a new glossary entry, a routing exception, an intent-definition tweak, a confidence-threshold adjustment, a new exemplar attached to an intent.",
        "Each proposed update is evaluated through the A/B comparison-mode path described in 4.6 Loop C (shadow run, defined horizon, pre-set success criteria, sub-segment regression checks) before any rule owner promotes it to production.",
        "Promoted updates take effect on the next pipeline run. The agent re-reads the updated Knowledge Base on every classification, so a corrected rule starts shaping decisions immediately.",
        "All of this is reversible. Every promoted update carries one-click rollback through the Governance and Analytics screen.",
    ])

    _p(doc, "How we measure quality.", bold=True)
    _bullets(doc, [
        "Correction recurrence rate: the percentage of CSR corrections that repeat within two weeks. Target is monotonic decrease over the rolling baseline.",
        "Time-to-rule-fix: from CSR correction to rule-book or KB update accepted by a Keysight rule owner. Target is within the next weekly governance cycle.",
        "Sub-segment regression check: every promoted update is verified against the held-out test set with no regression in any tracked per-region, per-language, or per-intent slice.",
    ])

    # ──────────────────────────────────────
    _h(doc, "5.8 Human-in-the-loop policy", 2)

    _p(doc, "What needs to be true.", bold=True)
    _p(doc,
       "Per Q&A item 22, 100 percent automation is not the goal. The initial deployment posture "
       "is AI drafts, CSR reviews, CSR sends. As confidence in each gate is validated through "
       "production data, autonomy is increased per use case. Multi-intent emails that exceed one "
       "gate's confidence may be partially automated and partially routed to human review.")

    _p(doc, "How we approach building it.", bold=True)
    _bullets(doc, [
        "At cutover, every use case ships at the most conservative tier: AI drafts, CSR reviews, CSR sends. The HITL portal carries the four-gate breakdown so the CSR can see exactly why the AI proposed what it did and where to push back.",
        "Autonomy expansion is a deliberate Keysight decision, not an automatic ramp. Each use case has its own confidence-gate calibration history; once the rolling accuracy on a tier crosses the threshold Keysight sets and the rule owners sign off, the autonomy floor for that use case shifts up. The shift is logged and reversible.",
        "Multi-intent emails are decomposed at the Intake stage. Each detected intent carries its own four-gate score. The pipeline acts only where all four gates clear the autonomy threshold for that intent; the rest of the email is routed to CSR review with the auto-acted portion captured in the case context.",
        "CSR override is always available. A CSR can force any case to full review, force a hold, or force a publish; every override is captured as Loop A feedback and surfaces in the Opportunity board.",
        "Service level for the HITL queue is per-intent per Q&A item 18 (different SLAs apply to different email types; new-customer creation has its own SLA). Per-intent SLA targets are agreed at Functional Design once Keysight publishes the addendum.",
    ])

    return _save(doc)


# ──────────────────────────────────────────────────────────────────
# 6. Delivery Plan
# ──────────────────────────────────────────────────────────────────
def build_delivery_plan() -> bytes:
    doc = _new_doc()
    _title_block(doc, "Section 6 of 13", "Delivery Plan")

    _p(doc,
       "Eight weeks to a working production pilot, structured as four two-week sprints. Each "
       "sprint ends with a working increment and a written acceptance gate signed off by the "
       "named Keysight rule owners. The plan compresses the standard delivery cycle because the "
       "expensive upstream work (requirement gathering, architecture choice) has already been "
       "done by Keysight in the RFP and the prior POC engagement.")

    _h(doc, "6.1 Milestones", 2)
    _three_col_table(doc,
                     ("Milestone", "Week", "Acceptance criterion"),
                     [
                         ("Sprint 1 sign-off: Functional design and KB seeding",
                          "End of Week 2",
                          "Intent definitions, Outlook pre-screen rules, routing rules, and the operational rule book seeded and signed off by Keysight rule owners. Initial classifier accuracy reported against the regression corpus."),
                         ("Sprint 2 sign-off: Use Cases 1 and 3 end to end",
                          "End of Week 4",
                          "Trade Order Entry and SOM WO Automation working end to end against the Salesforce sandbox. Confidence model implements all four gates. HITL portal usable by a CSR for case review."),
                         ("Sprint 3 sign-off: Remaining five use cases",
                          "End of Week 6",
                          "Trade Sales Change Order, SOM WO Update, WO Status and Inquiry, Service Contracts, and SSD Change all working end to end. Existing-CCC matrix verified across all ten statuses. Customer reply drafting in the sixteen supported languages."),
                         ("Sprint 4 sign-off: UAT pass and first cutover",
                          "End of Week 8",
                          "Stress test at 50x the 2,000-per-day Keysight baseline (100,000 emails per day sustained for 24 hours, with elastic worker scaling). UAT cycle complete. Drift detector and feedback loop live. First mailbox cutover into production with rollback plan tested. Tiered support hand-off complete."),
                     ])

    _h(doc, "6.2 Acceptance criteria (cumulative)", 2)
    _bullets(doc, [
        "Classification accuracy at or above the 90 percent threshold the RFP names, measured against Keysight's regression corpus at every sprint cut.",
        "Automation rate (the share of cases handled at high confidence plus the share handled with one-click CSR approval) measured on the agreed UAT inbox, against the target agreed with Keysight at Functional Design and aligned to the RFP's 60 to 70 percent manual-effort reduction goal.",
        "Classification latency below 2 seconds at the 95th percentile; full-pipeline latency below 30 seconds at the 95th percentile.",
        "Sustained throughput of 50x Keysight's stated 2,000-per-day baseline (100,000 emails per day sustained for 24 hours under stress test), elastic-scaling on queue depth. Sizing target absorbs Q-end and year-end bursts (Q&A item 59) without back-pressure into the mailbox layer.",
        "Human-in-the-loop queue service level: turnaround from low-confidence surface to operator resolve, against the SLA agreed with Keysight at Functional Design.",
        "Every email produces a complete audit trail (reference identifier, every reading, every extracted field, every decision gate, every external write, every customer-facing draft).",
        "No customer-facing message leaves the system outside the tier ladder. Where a case is L4 (high confidence), the reply is published automatically; where it is L3 or L2, a CSR confirms before publish.",
        "All knowledge bases (intents, routing rules, operational rule book, translation glossaries) editable by Keysight rule owners without a developer release cycle, with version history and rollback.",
        "Enterprise governance hooks verified at the scope agreed with Keysight's governance team at Functional Design.",
    ])

    _h(doc, "6.3 Governance and cadence", 2)
    _kv_table(doc,
              [
                  ("Daily stand-up",
                   "Joint Keysight and Hackett team, 15 minutes, written summary into the shared channel."),
                  ("Weekly demo",
                   "Live demo of the week's working increment to the broader Keysight stakeholder group. Open Q&A."),
                  ("Sprint review (bi-weekly)",
                   "Acceptance gate sign-off by the named Keysight rule owners. Retrospective on the sprint."),
                  ("Monthly steering committee",
                   "Keysight executive sponsors plus Hackett senior management. Risks, dependencies, scope changes formally reviewed."),
                  ("Backlog",
                   "Shared between both teams; visible at all times; every item has a named owner and a target sprint."),
                  ("Change control",
                   "Any change to scope or acceptance criteria is logged, reviewed at the next sprint review, and signed off at the steering committee."),
              ])

    _h(doc, "6.4 Phase exit gates", 2)
    _p(doc,
       "Each sprint ends with an exit gate. The gate is a checklist. Items either pass or fail. "
       "No sprint closes with open Critical or High items unless the steering committee explicitly "
       "approves a carryover with a defined remediation date.")
    _bullets(doc, [
        "Gate items have severity (Critical, High, Medium, Low) and an owner.",
        "Critical items must be Pass before the sprint closes.",
        "High items must be Pass or Mitigated (with a documented mitigation owned by a named person).",
        "Medium and Low items can carry over with steering-committee approval.",
        "The gate checklist is published in advance of every sprint so both teams know what is being measured.",
    ])

    return _save(doc)


# ──────────────────────────────────────────────────────────────────
# 7. Team Composition
# ──────────────────────────────────────────────────────────────────
def build_team_composition() -> bytes:
    doc = _new_doc()
    _title_block(doc, "Section 7 of 13", "Team Composition")

    _p(doc,
       "A single accountable team for product and services delivery. Continuity is preserved "
       "with the engineers who carried out Keysight's earlier POC engagement, so context does not "
       "have to be rebuilt from scratch. Final contracting and named individuals are confirmed at "
       "MSA signature.")

    _h(doc, "7.1 Delivery team during the ten-week engagement", 2)
    _three_col_table(doc,
                     ("Role", "Count", "Responsibility"),
                     [
                         ("Solution Architect", "1",
                          "Overall design and architecture. Stakeholder alignment with Keysight. Single point of accountability for the delivery."),
                         ("AI / ML Lead", "1",
                          "Classification, extraction, confidence calibration, translation. Owns the Knowledge Bases and the operational rule book mapping."),
                         ("AI / ML Engineer", "1",
                          "Pair-builds the AI / ML capabilities. Builds the regression-test harness against the labelled corpus and runs the 60 / 20 / 20 train, validation, test split."),
                         ("Integration Engineer", "2",
                          "Salesforce, Outlook through Microsoft Graph plus IMAP and SMTP, and the Jitterbit channel Keysight provides to Oracle EBS and DocuNet. ServiceNow-side integration is owned by Keysight's dev team; these roles coordinate with that team."),
                         ("Backend Engineer", "1",
                          "Orchestrator services, HITL portal API, audit trail, Knowledge Base authoring API."),
                         ("Frontend Engineer", "1",
                          "Operator HITL portal, Knowledge Base authoring screens, analytics dashboards."),
                         ("QA and UAT Engineer", "1",
                          "Regression test corpus, performance and stress tests, UAT scripts, disaster-recovery verification."),
                         ("Project Manager", "1",
                          "Phase delivery, releases, change control, governance, weekly update call. Owns the shared backlog."),
                     ])
    _p(doc, "Nine people across the ten-week engagement.", muted=True)

    _h(doc, "7.2 Week-by-week resource utilisation", 2)
    _p(doc,
       "Resource allocation by week. Each cell is full-time-equivalent (FTE) days per week for "
       "that role: 5 indicates full-time on the engagement, 0 indicates not on the engagement "
       "that week. The Solution Architect, AI / ML Lead, and Project Manager are continuous "
       "across the engagement. Build roles ramp into MVP Build (Weeks 5 to 8) and wind down "
       "across Testing and Tuning (Weeks 9 to 10). QA and UAT capacity ramps from MVP Build "
       "into peak coverage across Testing and Tuning.")

    t = doc.add_table(rows=10, cols=12)
    t.style = "Light Grid Accent 1"
    hdr = ["Role", "Wk 1", "Wk 2", "Wk 3", "Wk 4", "Wk 5", "Wk 6", "Wk 7", "Wk 8", "Wk 9", "Wk 10", "Total"]
    for j, h in enumerate(hdr):
        cell = t.rows[0].cells[j]
        cell.text = h
        _shade(cell, "1A55F9")
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9.5)
    rows_data = [
        ("Solution Architect",        5, 5, 5, 5, 5, 5, 5, 5, 5, 5),
        ("AI / ML Lead",              5, 5, 5, 5, 5, 5, 5, 5, 5, 5),
        ("AI / ML Engineer",          3, 5, 5, 5, 5, 5, 5, 5, 5, 5),
        ("Integration Engineer x 2",  4, 6, 8, 8, 10, 10, 10, 10, 8, 6),
        ("Backend Engineer",          3, 5, 5, 5, 5, 5, 5, 5, 3, 2),
        ("Frontend Engineer",         2, 3, 5, 5, 5, 5, 5, 5, 3, 2),
        ("QA and UAT Engineer",       1, 2, 2, 3, 3, 4, 5, 5, 5, 5),
        ("Project Manager",           5, 5, 5, 5, 5, 5, 5, 5, 5, 5),
    ]
    role_rows: list[tuple] = []
    for name, *weeks in rows_data:
        total = sum(weeks)
        role_rows.append((name, *weeks, total))
    week_totals = [sum(row[1:11]) for row in role_rows]
    grand_total = sum(week_totals)
    role_rows.append(("Total FTE-days per week", *week_totals, grand_total))
    for i, row in enumerate(role_rows, start=1):
        for j, val in enumerate(row):
            cell = t.rows[i].cells[j]
            cell.text = str(val)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)
                    if j == 0:
                        run.bold = True
                    if i == len(role_rows):
                        run.bold = True

    _p(doc,
       "Engineering continuity from the prior POC sits inside the AI / ML Lead and AI / ML "
       "Engineer rows. The same individuals who carried that work are part of the proposed team.",
       muted=True)

    _h(doc, "7.3 Hypercare team (Weeks 11 to 14, post go-live)", 2)
    _p(doc,
       "Hypercare runs for the first three months following production cutover. The hypercare "
       "team is a focused subset of the delivery team, retained specifically to absorb the first "
       "wave of operational findings, the first weeks of real-mail drift, and the first round of "
       "CSR feedback at scale. Hypercare is included in the implementation price; no separate "
       "fee applies.")

    _three_col_table(doc,
                     ("Role", "Coverage during hypercare", "Responsibility"),
                     [
                         ("Hypercare Lead (Solution Architect)",
                          "0.5 FTE, Keysight business hours",
                          "Single point of contact for Keysight during hypercare. Triages every escalation. Owns the hypercare exit report at the end of Week 14."),
                         ("AI / ML Engineer (on-call)",
                          "0.5 FTE, Keysight business hours, 24x7 on-call for critical incidents",
                          "Tunes classifier and extraction schemas against the first real CSR corrections. Runs the weekly drift report. Promotes accepted rule-book updates."),
                         ("Integration Engineer (on-call)",
                          "0.5 FTE, Keysight business hours, 24x7 on-call for critical incidents",
                          "First-line response on any Salesforce, Outlook, or Jitterbit integration issue (Jitterbit being the channel Keysight provides to Oracle EBS and DocuNet). Coordinates with Keysight's enterprise-system teams, including the Keysight dev team that owns the ServiceNow side."),
                         ("Project Manager",
                          "0.25 FTE, Keysight business hours",
                          "Weekly hypercare update call. Weekly status report to Keysight executive sponsors. Coordinates mailbox-by-mailbox cutover for the remaining mailboxes."),
                         ("Account Manager",
                          "0.1 FTE, named contact",
                          "Single point of contact for Keysight business sponsors. Monthly business review during hypercare. Hand-off to the steady-state support model at Week 14."),
                     ])

    _h(doc, "7.4 Steady-state support model (from Week 15 onward)", 2)
    _p(doc,
       "After hypercare exit the team transitions to the steady-state support model. Tiered "
       "support with defined SLAs, scaled to deployed agent count and Keysight's ongoing "
       "operational pattern.")
    _three_col_table(doc,
                     ("Tier", "Coverage", "Responsibility"),
                     [
                         ("Tier 1 (L1) support",
                          "Keysight business hours",
                          "Monitoring, alert triage, first-line response, run-book execution."),
                         ("Tier 2 (L2) support",
                          "Keysight business hours",
                          "Knowledge-base updates, integration fixes, minor enhancements, regression-test runs."),
                         ("Tier 3 (L3) support",
                          "On-call 24 x 7 for critical incidents",
                          "Architecture issues, model changes, major incidents, security patches."),
                         ("Account Manager",
                          "Named contact",
                          "Single point of contact for Keysight. Monthly business review. Capacity planning. Roadmap input."),
                     ])

    _h(doc, "7.5 Engineering continuity from the prior POC", 2)
    _p(doc,
       "The same engineers who carried out Keysight's earlier POC engagement are part of the "
       "proposed delivery team. The POC work covered the upstream classification pattern, the "
       "operational rule book interpretation, the multi-attachment handling, and the Salesforce "
       "integration profile. Carrying those engineers into the MVP build means we do not have to "
       "rebuild the institutional context that the POC produced.")

    return _save(doc)


# ──────────────────────────────────────────────────────────────────
# 8. Pricing Model
# ──────────────────────────────────────────────────────────────────
def build_pricing_model() -> bytes:
    doc = _new_doc()
    _title_block(doc, "Section 8 of 13", "Pricing Model")

    _p(doc,
       "Cost structure follows the RFP Pricing sheet template and the Q&A pricing clarifications "
       "(items 36 through 43). Keysight asked for both a modular price per use case and a "
       "turnkey bundled price; both are provided in the accompanying Pricing schedule, with the "
       "bundling discount shown explicitly. Implementation cost is separated from ongoing or "
       "recurring cost per Keysight's stated preference (Q&A item 37). The single table in 8.1 "
       "sets out every line; the subsections that follow explain what is and is not in each.")

    _h(doc, "8.0 Pricing structure (modular and turnkey)", 2)
    _bullets(doc, [
        "Modular pricing per use case (Q&A item 36): each of the seven RFP use cases is priced as a standalone implementation. Keysight can elect to engage on any subset (for example, Email Classification first, then Document Parsing, then Q2O preparation, then Post-Booking, then Service Order, then Delivery / Schedule), and the implementation can be staged accordingly.",
        "Turnkey bundle: the full seven-use-case engagement is priced as a single fixed-fee programme with a stated bundling discount over the sum of the modular prices. The discount is on the Pricing schedule's first page.",
        "Implementation separated from ongoing (Q&A item 37): all front-loaded implementation, build, training, and hypercare costs are itemised under the Implementation section of the matrix. Ongoing platform subscription, support tiers, and pass-through consumption are itemised separately under Ongoing Support and Pass-Through.",
        "Final acceptance basis (Q&A item 38): the turnkey programme is priced against a single final-acceptance milestone covering design, build, UAT, and production deploy. Per-sprint progress payments are available as an alternative payment schedule if Keysight prefers.",
        "Outcome-based pricing option (Q&A item 41): Keysight indicated interest in outcome-based pricing. We have included an outcome-based variant in the schedule (priced per email correctly classified, per task completed end-to-end, and per multi-intent email decomposed) so Keysight can evaluate it alongside the consumption-based model.",
        "Audit and consumption transparency (Q&A item 40): platform-bundled LLM tokens are stated explicitly (with the included models and the bundled-cost basis); separately-billed token consumption is invoiced monthly with a per-model rate sheet and a consumption audit trail Keysight can pull on demand.",
        "Worked examples at Keysight volumes (Q&A item 39): the Pricing schedule includes worked Year 1, Year 2, and Year 3 examples sized to the Q&A baseline (880K emails per year, 2,000 per day, email size one to twenty MB).",
        "No fixed pricing floor (Q&A item 42): pricing is described in the format that best fits the model. Free-text explanation accompanies the structured rows.",
    ])

    _h(doc, "8.1 Pricing matrix", 2)
    _p(doc,
       "All prices live in the accompanying Pricing schedule. The matrix below sets out "
       "each line of the price model, the unit of pricing, the frequency, and the dependency or "
       "scope note.")

    # Build a 5-column table: Category, Line item, Unit, Frequency, Notes
    pricing_rows = [
        # Platform Licensing
        ("Platform Licensing", "ZBrain Builder . core platform", "Per tenant or per user", "Annual",
         "Choice of pricing basis at MSA. Modular: Knowledge Bases, Flows, Apps, Agents, Monitor, Guardrails priced independently if Keysight wants a subset."),
        ("Platform Licensing", "Additional environment(s) . Dev / UAT", "Per environment", "Annual",
         "Production environment is included in the core licence. Dev and UAT priced per environment."),
        ("Platform Licensing", "Additional region or data residency", "Per region", "Annual",
         "Per region beyond the primary deployment."),

        # Implementation - per phase
        ("Implementation . Per Phase", "Sprint 1 . Discovery (Weeks 1 to 2)", "Fixed fee", "One-time",
         "Joint discovery with Keysight rule owners. Functional design sign-off."),
        ("Implementation . Per Phase", "Sprint 2 . MVP Build, Use Cases 1 and 3 (Weeks 3 to 4)", "Fixed fee", "One-time",
         "Trade Order Entry and SOM Work Order Automation end to end."),
        ("Implementation . Per Phase", "Sprint 3 . MVP Build, Use Cases 2, 4, 5, 6, 7 (Weeks 5 to 6)", "Fixed fee", "One-time",
         "Remaining five use cases end to end."),
        ("Implementation . Per Phase", "Sprint 4 . Testing, Tuning, Cutover (Weeks 7 to 8)", "Fixed fee", "One-time",
         "80 / 20 / 20 tuning, UAT, first mailbox cutover."),

        # Implementation - per use case
        ("Implementation . Per Use Case", "Use Case 1 . Trade Order Entry (PO Received)", "Fixed fee", "One-time", "Detailed effort in the Pricing schedule."),
        ("Implementation . Per Use Case", "Use Case 2 . Trade Sales Change Order", "Fixed fee", "One-time", "Detailed effort in the Pricing schedule."),
        ("Implementation . Per Use Case", "Use Case 3 . SOM Work Order Automation", "Fixed fee", "One-time", "Detailed effort in the Pricing schedule."),
        ("Implementation . Per Use Case", "Use Case 4 . SOM WO Update", "Fixed fee", "One-time", "Detailed effort in the Pricing schedule."),
        ("Implementation . Per Use Case", "Use Case 5 . WO Status and Inquiry", "Fixed fee", "One-time", "Detailed effort in the Pricing schedule."),
        ("Implementation . Per Use Case", "Use Case 6 . Service Contracts", "Fixed fee", "One-time", "Detailed effort in the Pricing schedule."),
        ("Implementation . Per Use Case", "Use Case 7 . SSD Change Request", "Fixed fee", "One-time", "Detailed effort in the Pricing schedule."),

        # Hypercare
        ("Hypercare", "Weeks 9 to 12 (first three months post go-live)", "Fixed fee", "One-time",
         "Included in implementation. Team and coverage defined in the Team Composition document."),

        # Ongoing support
        ("Ongoing Support", "Tier 1 (L1) . monitoring and alert triage", "Per agent per month", "Monthly",
         "Keysight business hours. Scales with deployed agent count."),
        ("Ongoing Support", "Tier 2 (L2) . KB updates, integration fixes", "Per agent per month", "Monthly",
         "Keysight business hours."),
        ("Ongoing Support", "Tier 3 (L3) . architecture, model changes, critical incidents", "Per agent per month", "Monthly",
         "On-call 24 x 7 for critical incidents."),
        ("Ongoing Support", "Account Manager", "Flat fee", "Monthly",
         "Named contact. Monthly business review."),

        # Pass-through
        ("Pass-Through . Infrastructure", "AWS account consumption (compute, storage, queueing, networking)", "Pass-through, at cost", "Monthly",
         "Billed direct to Keysight by AWS. Includes ECS Fargate, EC2, RDS, DocumentDB, ElastiCache, S3, CloudFront, MQ, ALB, and supporting services."),
        ("Pass-Through . LLM", "Frontier model token consumption", "Pass-through, at cost or cost-plus", "Monthly",
         "Keysight may hold the provider contract directly (pass-through at cost) or The Hackett Group holds it (cost-plus a small operational margin)."),

        # Optional
        ("Optional Services", "Additional mailbox cutover (beyond first)", "Per mailbox", "Per cutover",
         "Mailbox onboarded post the initial production cutover."),
        ("Optional Services", "Additional language pack", "Per language", "One-time",
         "Languages beyond the RFP-named sixteen."),
        ("Optional Services", "New use case (beyond the RFP seven)", "Fixed fee", "Per use case",
         "Scope, effort, and acceptance criteria defined at change request."),
        ("Optional Services", "Extended SLA . 24 x 7 follow-the-sun", "Uplift on ongoing support", "Monthly",
         "Available for the deployed agent count."),
    ]

    t = doc.add_table(rows=1 + len(pricing_rows), cols=5)
    t.style = "Light Grid Accent 1"
    headers = ["Category", "Line item", "Unit of pricing", "Frequency", "Notes / dependency"]
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = h
        _shade(cell, "1A55F9")
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9.5)
    prev_cat = ""
    for i, row in enumerate(pricing_rows, start=1):
        cat = row[0] if row[0] != prev_cat else ""
        prev_cat = row[0]
        for j, val in enumerate((cat,) + row[1:]):
            cell = t.rows[i].cells[j]
            cell.text = str(val)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8.5)
                    if j == 0:
                        run.bold = True

    _h(doc, "8.2 What sits inside implementation cost", 2)
    _bullets(doc, [
        "Sprint 1 Discovery: joint discovery with Keysight rule owners, environment provisioning, knowledge-base seeding, Version 1 alignment against Keysight inbox patterns.",
        "Sprints 2 and 3 Build: adapting Version 1 to your data and rule book across all seven use cases with the full sub-case catalogue.",
        "Sprint 4 Testing and Cutover: 60 / 20 / 20 train / validation / test tuning, stress test, UAT, first mailbox cutover.",
        "Integrations: Salesforce, Outlook through Microsoft Graph plus IMAP and SMTP, and use of the Jitterbit channel Keysight provides to Oracle EBS and Docunet. ServiceNow-side integration work is on the Keysight dev team and is not in our scope.",
        "Knowledge transfer: operator training, administrator training, builder training, run-book.",
        "Hypercare (Weeks 9 to 12): first three months post go-live, with the team and coverage defined in the Team Composition document.",
    ])

    _h(doc, "8.3 What sits outside implementation cost", 2)
    _bullets(doc, [
        "Additional mailbox cutovers beyond the first one. Priced per mailbox as a small fixed fee or rolled into ongoing support.",
        "New use cases beyond the seven named in the RFP. Treated as change requests with defined scope, effort, and acceptance criteria.",
        "Azure consumption. Billed direct to Keysight by Microsoft.",
        "LLM consumption where Keysight holds the provider contract directly.",
        "Material scope changes signed off by the steering committee under change control.",
    ])

    _h(doc, "8.4 Payment terms", 2)
    _bullets(doc, [
        "Milestone-based against per-sprint acceptance signed by Keysight rule owners.",
        "Discovery deposit at MSA signature.",
        "Net 30 against acceptance certificate.",
        "Hypercare invoiced monthly during Weeks 9 to 12, included in the implementation price.",
        "Ongoing support invoiced monthly from Week 13 onward.",
    ])

    return _save(doc)


# ──────────────────────────────────────────────────────────────────
# 9. Infrastructure
# ──────────────────────────────────────────────────────────────────
def build_infrastructure() -> bytes:
    from docx.shared import Inches
    from pathlib import Path as _Path

    doc = _new_doc()
    _title_block(doc, "Section 9 of 13", "Infrastructure")

    _p(doc,
       "This section sets out the deployment architecture for the proposed solution under the "
       "hybrid build-and-deploy model described in the Executive Summary and Solution Overview. "
       "Design and build happen inside ZBrain Solution Builder on the Hackett-operated platform "
       "tenancy. The production solution deploys separately onto Keysight's own AWS account "
       "(the Keysight cloud standard per Q&A item 55), with ZBrain services running in a VPC "
       "behind private subnets, Azure AD for SSO and RBAC (the Keysight enterprise identity "
       "standard per Q&A item 55), Azure Open AI bound through a private network path as the "
       "currently configured LLM service (model-agnostic; alternatives can be configured), "
       "Cloudflare at the edge for any internet-exposed agents (the Keysight standard per Q&A "
       "item 55), and Jenkins plus Bitbucket as the build and release toolchain. Production-grade "
       "controls (IAM, Secrets Manager, CloudWatch, Inspector, Security Hub) are present from "
       "day one.")

    _p(doc,
       "Material point: prompts, customer data, audit trails, and Knowledge Base content all "
       "live inside Keysight's AWS boundary at runtime. The Solution Builder tenancy at Hackett "
       "is used for design, build, and release-management only — once a release is cut, the "
       "artefacts ship to Keysight's account and the agents execute there. The boundary line is "
       "agreed at Functional Design: which signals (anonymised telemetry, error counts) can "
       "flow back to the Hackett tenancy for support, and which (customer content, decisions, "
       "audit trail) stay inside Keysight's AWS.",
       italic=True)

    _h(doc, "9.1 Reference architecture", 2)
    _p(doc,
       "The diagram below is the ZBrain reference architecture as deployed today. The Keysight "
       "deployment follows the same shape, configured to Keysight's account, VPC, and IAM "
       "structure. Final region and account boundaries are confirmed at Functional Design.")

    # Embed the architecture diagram.
    _arch = _Path(__file__).resolve().parents[3] / "frontend" / "public" / "asis-diagrams" / "zbrain-architecture.jpg"
    if _arch.exists():
        try:
            doc.add_picture(str(_arch), width=Inches(6.5))
            # Centre the picture
            last_para = doc.paragraphs[-1]
            last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _p(doc, "ZBrain reference architecture (AWS-hosted, Azure-side LLM and identity, Jenkins + Bitbucket toolchain).", italic=True, muted=True)
        except Exception:
            _p(doc, "[ZBrain reference architecture diagram — see accompanying image file zbrain-architecture.jpg]", italic=True, muted=True)
    else:
        _p(doc, "[ZBrain reference architecture diagram — see accompanying image file zbrain-architecture.jpg]", italic=True, muted=True)

    _h(doc, "9.2 Component summary", 2)
    _three_col_table(doc,
                     ("Layer", "Component", "Role"),
                     [
                         ("Edge", "Cloudflare", "Per Q&A item 55, the Keysight standard for internet-exposed agents. Edge WAF, DDoS, bot management, and where required mTLS for the HITL portal, the Knowledge Base authoring screens, the Governance and Analytics screen, and any inbound webhook surface."),
                         ("Edge", "CloudFront + S3", "Static-asset delivery and CDN for the operator portal."),
                         ("Edge", "Route 53 (DNS) + ACM (certificates)", "Custom domain and TLS termination."),
                         ("Edge", "Application Load Balancer", "Distributes traffic across the orchestrator and ZBrain service pools."),
                         ("ZBrain platform module", "Solution Builder", "Visual orchestration surface for the end-to-end agents (Intake, Extract, Reconcile, Decide, Execute, Communicate). Configuration runs on the Hackett-operated ZBrain tenancy during design and build; releases ship to Keysight's AWS account."),
                         ("ZBrain platform module", "Knowledge Base", "Stores and serves all three RAG retrievers (structured DB grounding, knowledge-graph traversal, unstructured document search) per Q&A item 53. Editable by Keysight rule owners."),
                         ("ZBrain platform module", "Governance", "Application-level governance per agent: RBAC, allow-listed tool access, prompt-injection defences, citizenship-based access enforcement, audit forwarding."),
                         ("ZBrain platform module", "Monitor", "Real-time observability of every agent: classification mix, confidence distribution, four-gate pass rates, HITL queue depth, drift signals."),
                         ("Compute", "Amazon ECS on Fargate", "Serverless containers for the Orchestrator Service and the supporting ZBrain Services. Auto Scaling adjusts capacity to load."),
                         ("Compute", "Amazon EC2 (where Fargate is not used)", "Long-running components that need direct EC2 hosting."),
                         ("State", "Amazon RDS (PostgreSQL)", "Case state, audit trail, configuration, knowledge-base entries."),
                         ("State", "Amazon DocumentDB", "MongoDB-compatible document store for unstructured content and intermediate AI outputs."),
                         ("State", "Amazon ElastiCache for Redis", "Low-latency cache and session store."),
                         ("State", "Amazon S3", "Attachment storage, archive, static asset hosting, build artefacts."),
                         ("Messaging", "Amazon MQ (RabbitMQ)", "Message broker between the inbound queue and the worker pool. Absorbs burst spikes without back-pressure into the mailbox."),
                         ("Identity", "Microsoft Azure AD", "Per Q&A item 55, the Keysight SSO and RBAC standard. Single sign-on for the HITL portal, the Knowledge Base authoring screens, and the Governance and Analytics screen, with conditional-access policies on risky operations."),
                         ("Identity", "AWS IAM", "Service-level access controls for the AWS resources behind the ZBrain services."),
                         ("Secrets", "AWS Secrets Manager", "Credentials, API keys, integration tokens. No secrets in source code or in container images."),
                         ("Network", "Amazon VPC + private subnets", "Network isolation. ZBrain services run inside the VPC; only the Cloudflare-fronted edge and CloudFront are reachable externally."),
                         ("Monitoring (cloud)", "Amazon CloudWatch", "Logs, metrics, alarms across the AWS stack."),
                         ("Monitoring (cloud)", "Amazon Inspector", "Automated security vulnerability assessment for the AWS workloads."),
                         ("Monitoring (cloud)", "AWS Security Hub", "Centralised security and compliance signal for the AWS account."),
                         ("Container registry", "Amazon ECR", "Private container image registry feeding the ECS / Fargate services."),
                         ("LLM provider", "Azure Open AI (currently configured)", "Model-agnostic platform; alternative providers can be configured. Keysight has no in-house custom LLM per Q&A item 51."),
                         ("CI/CD", "Jenkins (Jenkins Master) + Bitbucket", "Source control and the deployment pipeline. Build artefacts published to ECR; deployments rolled out to ECS."),
                     ])

    _h(doc, "9.3 How the architecture meets Keysight's stated needs", 2)
    _kv_table(doc,
              [
                  ("Scalability to Keysight's stated volumes",
                   "ECS Fargate with AWS Auto Scaling handles the inbound queue and worker pools. Amazon MQ decouples the IMAP poller from the workers so burst spikes (quarter-end and year-end per Q&A item 59) absorb in the queue rather than pushing back into the mailbox layer. Baseline sized at Keysight's stated 2,000 emails per day, ~880K per year (Q&A item 11); stress-test target is 50x baseline sustained for 24 hours, with elastic worker scaling on queue depth. Concurrent operator load (80 to 90 baseline per Q&A item 17) is handled by the HITL portal's separate horizontal scaling."),
                  ("Resilience and continuity",
                   "Multi-AZ deployment within the region. Cross-region disaster recovery position confirmed against Keysight's continuity standards at Functional Design. RDS automated backups and snapshots; S3 versioning for attachments and artefacts."),
                  ("Network isolation",
                   "All ZBrain services run inside a VPC with private subnets. Only the Application Load Balancer (fronted by ACM for TLS) and CloudFront are reachable externally. Outbound calls to Azure Open AI, Salesforce, Jitterbit (Keysight's middleware channel to Oracle EBS and Docunet), and Microsoft Graph for Outlook go through controlled egress paths."),
                  ("Identity and access",
                   "Operator and end-user identity through Azure AD. AWS IAM enforces least-privilege access for the AWS services themselves. AWS Secrets Manager holds every external credential."),
                  ("Observability",
                   "CloudWatch handles logs, metrics, and alarms; Inspector handles vulnerability assessment; Security Hub aggregates security posture. ZBrain Monitor layers application-level telemetry on top. Exports to Keysight's existing SIEM are agreed at Functional Design."),
                  ("Deployment and release",
                   "Multi-stage pipeline orchestrated through Jenkins, with builds promoted from dev to UAT to production using the same container artefact stored in ECR. Quality gates at every promotion; canary or blue-green rollout into production with health-check-driven rollback. Release alignment with Keysight's existing change windows; freezes respected."),
                  ("Compliance posture",
                   "Encryption at rest and in transit by default; AWS KMS managed keys (with customer-managed keys available for sensitive workloads); data residency by region selection; audit trail exportable to Keysight's SIEM; ITAR-ready in a customer-controlled deployment with restricted administrator access and enhanced audit logging."),
              ])

    _h(doc, "9.4 Integration channels", 2)
    _three_col_table(doc,
                     ("Target system", "Integration path", "Notes"),
                     [
                         ("Salesforce",
                          "REST and SOQL through Salesforce APIs.",
                          "Operator credentials and service-principal credentials managed in AWS Secrets Manager."),
                         ("Oracle EBS",
                          "Through Jitterbit, your existing integration platform.",
                          "We do not introduce a parallel middleware. We use Jitterbit the same way your downstream CSR work uses it today."),
                         ("Docunet (document store)",
                          "Reached through Jitterbit, the same middleware Keysight uses for Oracle EBS. The Docunet integration is provided by Keysight on the middleware side.",
                          "Files are filed in Docunet with Doc type FCNV, per Keysight's existing convention."),
                         ("ServiceNow",
                          "ServiceNow-side integration is built and owned by Keysight's dev team.",
                          "Keysight's existing change-order approval and reminder workflow stays as-is. We coordinate with the Keysight dev team rather than building this integration ourselves."),
                         ("Microsoft Outlook",
                          "IMAP for inbound mail; Microsoft Graph for mailbox metadata and folder operations.",
                          "Mailbox-agnostic. Works transparently across the current set of mailboxes during the consolidation transition."),
                         ("Azure Open AI",
                          "Azure Open AI service endpoint via secure outbound from the VPC.",
                          "Currently the configured LLM service. ZBrain remains model-agnostic; alternative providers can be configured if Keysight chooses."),
                         ("Azure AD",
                          "OIDC or SAML federation for operator login.",
                          "Provides single sign-on; conditional-access policies enforced by Azure AD."),
                     ])

    _h(doc, "9.5 What stays out of scope at the infrastructure layer", 2)
    _bullets(doc, [
        "Changes to the Jitterbit topology Keysight already operates.",
        "Changes to the Azure AD or Microsoft Graph permission model beyond what is needed to add the application registration.",
        "Changes to Keysight's existing ServiceNow workflows.",
        "Changes to the Outlook tenant or mailbox infrastructure.",
        "Standing up parallel observability or SIEM tooling. We integrate with what Keysight already operates.",
    ])

    return _save(doc)


# ──────────────────────────────────────────────────────────────────
# 10. Security and Governance
# ──────────────────────────────────────────────────────────────────
def build_security_governance() -> bytes:
    doc = _new_doc()
    _title_block(doc, "Section 10 of 13", "Security and Governance")

    _p(doc,
       "Security and governance are designed in from the beginning, not added at the end. The "
       "model has two tiers: enterprise-level governance that aligns with Keysight's existing "
       "enterprise tooling (today, Microsoft for SSO and RBAC, AWS as the cloud standard, "
       "Cloudflare for any internet-exposed agents, per Q&A item 55), and application-level "
       "governance per agent, delivered by the ZBrain Governance and Monitor modules running "
       "inside Keysight's own AWS deployment under the hybrid build-and-deploy model.")

    _p(doc,
       "Per Q&A item 55, Keysight does not operate a Purview-equivalent enterprise AI "
       "governance fabric today. The model below describes what the ZBrain Governance module "
       "delivers at deployment, and how the surface integrates with enterprise governance tooling "
       "as Keysight adopts it.",
       italic=True)

    _h(doc, "10.1 Enterprise-level governance", 2)
    _kv_table(doc,
              [
                  ("Identity (Microsoft Azure AD)",
                   "Per Q&A item 55, Keysight uses Microsoft tooling for SSO and RBAC. Azure AD provides single sign-on for the HITL portal, the Knowledge Base authoring screens, and the ZBrain Governance and Analytics screen, with conditional-access policies on risky operations (KB rule promotion, A/B candidate promotion, manual override of a low-confidence decision, force-publish of a customer reply)."),
                  ("Cloud (AWS standard)",
                   "Per Q&A item 55, AWS is the Keysight cloud standard. Under the hybrid build-and-deploy model, the production solution is deployed onto Keysight's own AWS account. ZBrain Solution Builder runs separately on the Hackett platform tenancy for design and build; the deployed runtime runs entirely inside Keysight's AWS. AWS IAM enforces service-level access for the AWS resources behind the ZBrain services in Keysight's account; AWS Secrets Manager holds every external credential; AWS CloudWatch, Inspector, and Security Hub provide cloud-level observability and compliance signal."),
                  ("Internet edge (Cloudflare)",
                   "Per Q&A item 55, Cloudflare is used for any internet-exposed agents at Keysight. The solution's external surfaces (HITL portal, Knowledge Base authoring screens, Governance and Analytics screen, and any internet-facing webhook used for inbound integrations) inherit Cloudflare's edge controls (WAF, DDoS, bot management, mTLS where required)."),
                  ("Audit forwarding",
                   "Every classification, extraction, decision gate, external write, and customer-facing draft is logged with a per-email reference and forwarded to the SIEM Keysight already operates. Retention is approximately ten years per Q&A item 45, which is the audit requirement Keysight has named."),
                  ("Data classification",
                   "Inbound mail and extracted fields tagged by sensitivity. KSO and restricted-customer content is tagged before the AI reads the body, per 10.3 below."),
                  ("Enterprise governance roadmap",
                   "As Keysight adopts an enterprise AI governance fabric (Purview-equivalent), the ZBrain Governance module exposes the integration points needed (agent registration, data-class export, policy-decision-point hooks). The integration scope is agreed jointly with Keysight's governance team at Functional Design and revisited at the quarterly business review."),
              ])

    _h(doc, "10.2 Application-level governance per agent (ZBrain Governance and Monitor modules)", 2)
    _kv_table(doc,
              [
                  ("Role-based access control",
                   "Privilege rings per agent in the ZBrain Governance module. Operators see what they need; rule owners can edit knowledge bases; builders can deploy new agents; auditors get read-only with export. No one role carries end-to-end privilege by default."),
                  ("Service-level objectives",
                   "Each agent ships with a written SLO in the Monitor module covering accuracy, latency, automation rate, and human-in-the-loop service level. SLOs are monitored continuously; deviations trigger alerts in the Governance and Analytics screen."),
                  ("Circuit-breaker policies",
                   "Each agent has automated pause thresholds. If classification accuracy falls below the SLO floor, or the HITL rate spikes above an upper bound, the agent pauses auto-action for the affected segment and routes everything to CSR review until a rule owner clears the issue. The breaker is segment-aware (per region, per language, per intent) so a regional anomaly does not shut down the global pipeline."),
                  ("Tool access controls",
                   "Agents are allow-listed for the external tools they may call (Salesforce, Jitterbit which carries Oracle EBS and DocuNet, Microsoft Graph for Outlook). ServiceNow-side integration is owned by Keysight's dev team. New tool additions require a Governance-module review."),
                  ("Prompt-injection defences",
                   "Inbound content is filtered for known prompt-injection patterns before it reaches the AI. The reasoning path is bounded to the operational rule book; the AI cannot be talked into executing actions outside its scope."),
                  ("Change control and rollback",
                   "Every promoted change in the Governance and Analytics screen (KB edit, A/B promotion, rule-book update, threshold adjustment, model swap) carries version history, named-owner attribution, and one-click rollback. The change log is exported to Keysight's audit retention nightly."),
              ])

    _h(doc, "10.3 Citizenship-based access and KSO routing", 2)
    _p(doc,
       "Q&A item 25 sets the rule: emails from US-government customers must be readable only by "
       "US-citizen reviewers, and non-US-citizen access must be blocked across model inference, "
       "logging, replay, and any vendor support staff with system access. The solution implements "
       "this end to end.")
    _bullets(doc, [
        "At the mailbox door, the KSO routing rule fires before any AI reads the body. The decision is deterministic and based on sender domain and body-keyword signals already documented in Keysight's existing Outlook rules. The redirect is logged immutably for compliance audit.",
        "Inside the platform, restricted-customer email is tagged with a citizenship-restricted classification at ingress. The tag travels with the email reference through every stage.",
        "Reasoning paths: the model inference call for a citizenship-restricted email is constrained to a US-resident inference endpoint, with the request audit-logged including the inference region and the model identifier.",
        "Logging: structured logs and trace data for citizenship-restricted emails are written to a separate index with RBAC restricted to US-citizen operators and auditors only. The standard log index does not receive a copy.",
        "Replay: the comparison-mode replay path (used for A/B candidate evaluation) refuses to replay citizenship-restricted emails outside the US-citizen access boundary. The candidate is evaluated only on the non-restricted slice or, where Keysight authorises, only by US-citizen operators inside the restricted index.",
        "Vendor support: any Hackett support engineer accessing citizenship-restricted data must hold US citizenship; access is provisioned through the Governance module's restricted role, with every access event logged. Access requests are reviewed and approved by Keysight before being granted.",
    ])

    _h(doc, "10.4 Data security", 2)
    _bullets(doc, [
        "Encryption at rest and in transit by default. Modern cipher suites only. AWS KMS managed keys, with customer-managed keys available for workloads that require Keysight-controlled cryptographic material.",
        "Data masking in non-production environments. UAT and dev environments never see production customer content; representative synthetic data is used.",
        "Data residency by region. Keysight picks the deployment region per regulatory requirement; cross-region transfer is explicit and audited.",
        "Right-to-erasure and right-to-access endpoints available for regulatory data-subject requests.",
        "Data retention: approximately ten years per Q&A item 45. The retention policy is enforced by the Governance module against the audit store and the SIEM export; expiry is automatic.",
    ])

    _h(doc, "10.5 Audit trail", 2)
    _p(doc,
       "Every inbound email produces a complete audit trail with a single reference identifier. "
       "The trail records every reading, every extracted field with the source attachment "
       "reference, every decision gate score, every external write (Salesforce, plus the "
       "Jitterbit calls to Oracle EBS and DocuNet), every customer-facing draft, and every "
       "CSR-driven change. The trail is immutable, append-only, exported to the SIEM Keysight "
       "operates, and retained for the ten-year audit window named in Q&A item 45.")

    return _save(doc)


# ──────────────────────────────────────────────────────────────────
# 11. References
# ──────────────────────────────────────────────────────────────────
def build_references() -> bytes:
    doc = _new_doc()
    _title_block(doc, "Section 11 of 13", "References")

    _h(doc, "11.1 Prior Keysight engagement", 2)
    _p(doc,
       "The Hackett Group has previously delivered a POC engagement for Keysight on the ZBrain platform "
       "covering the upstream classification pattern, the multi-format attachment handling, and "
       "the operational-rule-book interpretation. The engineers from that engagement are part of "
       "the proposed MVP delivery team, so institutional context is preserved into the production "
       "build phase.")

    _h(doc, "11.2 Comparable implementations", 2)
    _p(doc,
       "Per Q&A item 34, Keysight's reference criteria are explicit: references must be in a "
       "similar domain (semiconductor, industrial test, or comparable high-tech vertical) AND "
       "for sales-operations or order-management use cases. Reference calls focus on this use "
       "case only and are not used to cover the vendor's broader platform footprint at the "
       "reference customer. Other Keysight teams (for example supply chain or services) are not "
       "on the same reference call.")
    _p(doc,
       "We hold references that match both halves of the criteria, available under NDA at the "
       "vendor capability review stage:")
    _bullets(doc, [
        "Sales-operations and order-management email automation at semiconductor or industrial-test vendors with comparable annual volume and integration footprint.",
        "Document-intelligence implementations in the same domain involving structured extraction from multi-format attachments (PDF, Word, Excel, image, embedded inbox items).",
        "Multi-language customer-communication automation across the language list named in the RFP, deployed inside a sales-ops or order-management workflow.",
        "Enterprise integrations spanning Salesforce, Oracle EBS through Jitterbit, and Microsoft 365 in a single sales-operations workflow.",
    ])

    _h(doc, "11.3 Reference selection process", 2)
    _bullets(doc, [
        "At the vendor capability review stage, Keysight confirms the specific reference profile of greatest interest (sub-vertical, deployment scale, integration scope).",
        "The Hackett Group nominates two to three customer references matching Q&A item 34's combined criteria (similar domain + sales-ops or order-management use case).",
        "Reference calls are scoped to this use case only, per Q&A item 34. The broader platform footprint with the reference customer is not part of the conversation.",
        "Other Keysight teams (supply chain, services) are not on the same call, per Q&A item 34. If a multi-track conversation is needed, separate calls are scheduled.",
        "Final reference selection is confirmed before contract signature.",
    ])

    return _save(doc)


# ──────────────────────────────────────────────────────────────────
# 12. Demo
# ──────────────────────────────────────────────────────────────────
def build_demo() -> bytes:
    doc = _new_doc()
    _title_block(doc, "Section 12 of 13", "Demo")

    _p(doc,
       "Per Q&A item 32, Keysight has asked for a demo of production-deployed work for similar "
       "customers, not a custom build for Keysight. The demo is a walk-through of an actual "
       "customer implementation already running on ZBrain in the SalesOps and email-automation "
       "space: the timeline to go-live, the data analysed, the outcomes achieved, and how the "
       "implementation maps to Keysight's seven RFP use cases. Live test emails and bespoke POC "
       "builds are not part of the demo, in line with Keysight's stated preference.")

    _h(doc, "12.1 Demo format and slot", 2)
    _p(doc,
       "Per Q&A item 32, Keysight has allocated a 90-minute slot for the vendor demo. Our "
       "proposed split is 15 to 20 minutes of demo followed by Q&A using the remaining time. "
       "If Keysight prefers a longer demo with shorter Q&A, we can rebalance.")
    _three_col_table(doc,
                     ("Segment", "Duration", "Content"),
                     [
                         ("Demo walk-through",
                          "15 to 20 minutes",
                          "Live screen-share through a production-deployed ZBrain implementation in a similar domain. Walk the email-classification, document-extraction, four-gate confidence model, HITL portal, multilingual reply drafting, and Knowledge Base authoring flows. Each section is anchored to one of the seven RFP use cases so the mapping is explicit."),
                         ("Outcomes review",
                          "10 minutes",
                          "The headline outcomes achieved on the production deployment: automation rate, accuracy, time-to-go-live, post-cutover learning curve. Data shown is from the live production instance, redacted where customer confidentiality requires."),
                         ("Q&A",
                          "Remaining time",
                          "Open Q&A with the Keysight evaluation team, including the project team, Head of Global Apps, the Chief Data and Analytics Officer, and the Services representative (per the audience confirmed in Q&A item 33). Architecture, governance, integration, and pricing questions welcome."),
                     ])

    _h(doc, "12.2 What the demo will cover", 2)
    _bullets(doc, [
        "The production-deployed AI assistant working through inbound mail in a domain analogous to Keysight's, covering classification, extraction, decision, and reply drafting end to end.",
        "Classification edge cases similar to those in Keysight's rule book (multi-intent emails, PO mentioned alongside a Work Order, acknowledgement-only threads, generic forwards that hide a real request earlier in the thread).",
        "Extraction across attachment formats relevant to Keysight (PDF including scanned forms, Word, Excel, images, embedded inbox items).",
        "The four-gate confidence display in the operator portal, so the evaluation team can see how the AI explains itself.",
        "Customer reply drafting in multiple languages, using a production glossary.",
        "The Knowledge Base authoring flow: a live edit during the demo, and the next pipeline run picking up the change without a restart.",
        "The Governance and Analytics screen (described in Implementation Approach 4.7): health tiles, the Opportunity board, active A/B experiments, the drift timeline, and the change log.",
        "A walk-through of how the seven RFP use cases would map onto the production deployment shown.",
    ])

    _h(doc, "12.3 Reference implementation shown", 2)
    _p(doc,
       "The specific customer implementation used for the demo is confirmed under NDA at the "
       "vendor capability review stage. The selection criteria are aligned to Q&A item 34: same "
       "domain (semiconductor, industrial test, or comparable high-tech vertical) and same use "
       "case profile (sales operations and order management). The reference focus is this use "
       "case only; broader platform footprint at that customer is not part of the demo per "
       "Keysight's stated preference.")

    _h(doc, "12.4 Supplementary materials", 2)
    _p(doc,
       "Per Q&A item 35, Keysight has confirmed that supplementary documents (roadmap, "
       "architecture, governance posture) may be submitted alongside the RFP response. These are "
       "available on request: the ZBrain product roadmap, the deployment reference architecture "
       "in detail, and the governance integration approach with Microsoft enterprise tooling.")

    _h(doc, "12.5 Demo outcomes and next steps", 2)
    _bullets(doc, [
        "The Keysight evaluation team sees the solution working against material that is representative of the production workload at a comparable customer.",
        "Open questions and feedback from the demo are captured into a shared backlog for the engagement.",
        "Any change to the proposed scope is taken into Sprint 1 functional design.",
        "Reference customer introductions are scheduled for the vendor capability review stage. Reference calls focus on the SalesOps and order-management use case only, per Q&A item 34.",
    ])

    return _save(doc)


# ──────────────────────────────────────────────────────────────────
# 13. RACI Matrix (role by provider)
# ──────────────────────────────────────────────────────────────────
def build_raci_matrix() -> bytes:
    from docx.enum.section import WD_ORIENT

    doc = _new_doc()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    sec.left_margin = Cm(1.5)
    sec.right_margin = Cm(1.5)

    _title_block(doc, "Section 13 of 13", "RACI Matrix (role by provider)")

    _p(doc,
       "Responsibility assignment across the delivery and operate roles for the Keysight "
       "SalesOps engagement. The matrix is read role by provider, not activity by role. Each "
       "cell carries the responsibility the named provider holds for the role on the row.")

    _h(doc, "Legend", 2)
    _kv_table(doc,
              [
                  ("R", "Responsible. Assigned to do the work."),
                  ("A", "Accountable. Ultimately owns the work getting done."),
                  ("C", "Consultant. Must review or approve before completion."),
                  ("I", "Informed. Must be communicated to before completion."),
                  ("A / R", "Accountable first, responsible for the work in smaller capacity."),
                  ("R / A", "Responsible first, accountable to partner. Teach-to-fish role."),
              ],
              key_header="Code", val_header="Meaning")

    _h(doc, "Provider columns", 2)
    _kv_table(doc,
              [
                  ("KS ITS", "Keysight IT, Infrastructure, Security, and Governance team."),
                  ("KS Business", "Keysight rule owners, FCNV operator leads, business sponsors, and steering committee."),
                  ("THG Tech", "The Hackett Group technical delivery team (Solution Architect, AI / ML Lead, AI / ML Engineer, Integration Engineer, Backend Engineer, Frontend Engineer, QA, Evaluation Engineer)."),
                  ("THG Business", "The Hackett Group commercial and delivery management team (Project Manager, Program Manager, Account Manager)."),
                  ("CSP", "Cloud Service Provider (AWS). Tenancy operation and platform incident response."),
              ],
              key_header="Provider", val_header="Definition")

    _h(doc, "Role by provider", 2)

    raci_rows: list[tuple[str, str, str, str, str, str]] = [
        ("Product Owner (Functional Expert, Business Solution Lead)",     "C",   "A / R", "I",   "R / A", ""),
        ("Solution Architect",                                            "A",   "C",    "R",   "I",     "I"),
        ("Data / Knowledge Engineer (KB owner)",                          "A / R", "C",  "R / A", "I",   ""),
        ("Infra and Platform Engineering",                                "A / R", "I",  "R",   "I",     "C"),
        ("AI Engineer",                                                   "A / R", "I",  "R / A", "I",   ""),
        ("Prompt Engineer",                                               "A / R", "C",  "R / A", "C",   ""),
        ("Full Stack Engineer (HITL portal, KB authoring)",               "A / R", "I",  "R / A", "I",   ""),
        ("QA",                                                             "A / R", "C",  "R / A", "C",   ""),
        ("UAT",                                                            "I",   "A / R", "C",   "C",    ""),
        ("Evaluation Engineer (accuracy, precision, recall, drift)",       "A / R", "C",  "R / A", "C",   ""),
        ("UI / UX Engineer",                                              "A / R", "C",  "R / A", "C",   ""),
        ("Analytics Engineer (Monitor module)",                           "A / R", "C",  "R / A", "C",   ""),
        ("Business Analytics Engineer",                                   "C",   "A / R", "C",   "R / A", ""),
        ("Deployment",                                                     "A / R", "C",  "R / A", "C",   "I"),
        ("Hypercare",                                                      "A / R", "C",  "R / A", "C",   "I"),
        ("Data SME (defining and validating data)",                       "R",   "A",    "I",   "C",     ""),
        ("AWS Cloud SME",                                                  "R / A", "I",  "C",   "I",    "A / R"),
        ("Salesforce SME",                                                 "R / A", "C",  "C",   "I",    ""),
        ("Security",                                                       "A / R", "C",  "I",   "C",    "I"),
        ("Privacy",                                                        "A / R", "C",  "I",   "C",    ""),
        ("Compliance",                                                     "A / R", "C",  "C",   "C",    ""),
        ("Legal Review",                                                   "A / R", "C",  "I",   "C",    ""),
        ("ServiceNow Programmer (Keysight-owned)",                        "A / R", "C",  "I",   "I",    ""),
        ("Program Manager",                                                "C",   "I",    "C",   "A / R", "I"),
        ("Use Case Project Manager",                                       "I",   "C",    "C",   "A / R", "I"),
        ("Scrum Master (weekly update call facilitator)",                  "I",   "C",    "A / R", "R / A", ""),
        ("MDM",                                                            "A / R", "C",  "I",   "I",    ""),
        ("Data Governance",                                                "A / R", "C",  "I",   "I",    ""),
    ]

    providers = ["Role", "KS ITS", "KS Business", "THG Tech", "THG Business", "CSP"]
    t = doc.add_table(rows=1 + len(raci_rows), cols=len(providers))
    t.style = "Light Grid Accent 1"
    for j, h in enumerate(providers):
        cell = t.rows[0].cells[j]
        cell.text = h
        _shade(cell, "1A55F9")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9.5)
    for i, row in enumerate(raci_rows, start=1):
        for j, val in enumerate(row):
            cell = t.rows[i].cells[j]
            cell.text = val
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            for p in cell.paragraphs:
                if j > 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
                    if j == 0:
                        run.bold = True

    t.columns[0].width = Cm(8.0)
    for c in range(1, len(providers)):
        t.columns[c].width = Cm(3.6)

    _p(doc,
       "Final named individuals per role are confirmed at MSA signature. Empty cells indicate "
       "no defined responsibility for that provider on that role; the role does not require "
       "input or notification from that provider as part of the standard engagement model.",
       muted=True)

    return _save(doc)


# ──────────────────────────────────────────────────────────────────
# Section registry and combined-doc builder
# ──────────────────────────────────────────────────────────────────

SECTIONS: list[dict] = [
    {"slug": "cover-letter", "label": "Cover Letter",
     "filename": "Keysight_RFP_00_Cover_Letter.docx", "builder": build_cover_letter},
    {"slug": "executive-summary", "label": "1. Executive Summary",
     "filename": "Keysight_RFP_01_Executive_Summary.docx", "builder": build_executive_summary},
    {"slug": "solution-overview", "label": "2. Solution Overview",
     "filename": "Keysight_RFP_02_Solution_Overview.docx", "builder": build_solution_overview},
    {"slug": "scope", "label": "3. Scope",
     "filename": "Keysight_RFP_03_Scope.docx", "builder": build_scope},
    {"slug": "implementation-approach", "label": "4. Implementation Approach",
     "filename": "Keysight_RFP_04_Implementation_Approach.docx", "builder": build_implementation_approach},
    {"slug": "ai-ml-capabilities", "label": "5. AI / ML Capabilities",
     "filename": "Keysight_RFP_05_AI_ML_Capabilities.docx", "builder": build_aiml_capabilities},
    {"slug": "delivery-plan", "label": "6. Delivery Plan",
     "filename": "Keysight_RFP_06_Delivery_Plan.docx", "builder": build_delivery_plan},
    {"slug": "team-composition", "label": "7. Team Composition",
     "filename": "Keysight_RFP_07_Team_Composition.docx", "builder": build_team_composition},
    {"slug": "pricing-model", "label": "8. Pricing Model",
     "filename": "Keysight_RFP_08_Pricing_Model.docx", "builder": build_pricing_model},
    {"slug": "infrastructure", "label": "9. Infrastructure",
     "filename": "Keysight_RFP_09_Infrastructure.docx", "builder": build_infrastructure},
    {"slug": "security-governance", "label": "10. Security and Governance",
     "filename": "Keysight_RFP_10_Security_and_Governance.docx", "builder": build_security_governance},
    {"slug": "references", "label": "11. References",
     "filename": "Keysight_RFP_11_References.docx", "builder": build_references},
    {"slug": "demo", "label": "12. Demo",
     "filename": "Keysight_RFP_12_Demo.docx", "builder": build_demo},
    {"slug": "raci-matrix", "label": "13. RACI Matrix",
     "filename": "Keysight_RFP_13_RACI_Matrix.docx", "builder": build_raci_matrix},
]


def get_section(slug: str) -> dict | None:
    for s in SECTIONS:
        if s["slug"] == slug:
            return s
    return None


# Backwards-compat alias kept for the previous combined-doc URL.
def build_rfp_reply_docx() -> bytes:
    """Combined doc kept for backwards compatibility with /api/docs/rfp-reply.docx.
    Returns the Executive Summary so the old URL still resolves to something useful."""
    return build_executive_summary()
# === v1.1 RFP-REPLY END ===
