"""In-process XLSX → PDF and DOCX → PDF conversion.

The Lambda extractor (Azure Document Intelligence wrapper) only accepts a
publicly reachable `pdf_url`. To run XLSX/DOCX through the SAME Azure
pipeline, we render their content into a quick PDF first and ship that
PDF to the Lambda.

The output PDF is intentionally simple — clean text + tables, sized for
3-page Stage 1 light extraction. No styling, since Azure DocIntel will
rewrite layout anyway.

We use libraries already in the project's requirements:
  * openpyxl    — read XLSX cell grids
  * python-docx — read DOCX paragraphs and tables
  * reportlab   — render the resulting PDF
"""
from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from openpyxl import load_workbook
from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


def _styles() -> dict[str, ParagraphStyle]:
    base = ParagraphStyle("base", fontName="Helvetica", fontSize=10, leading=13)
    return {
        "title": ParagraphStyle("title", parent=base, fontName="Helvetica-Bold", fontSize=14, leading=18, spaceAfter=8),
        "h2": ParagraphStyle("h2", parent=base, fontName="Helvetica-Bold", fontSize=11, leading=14, spaceBefore=8, spaceAfter=4),
        "para": base,
        "kv": ParagraphStyle("kv", parent=base, fontSize=10, leading=12),
    }


def xlsx_to_pdf(src: Path, dest: Path) -> None:
    """Render every populated cell of every sheet into a basic table PDF."""
    wb = load_workbook(filename=str(src), data_only=True, read_only=True)
    s = _styles()
    flow: list = []
    flow.append(Paragraph(f"Workbook: {src.name}", s["title"]))

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        flow.append(Paragraph(f"Sheet: {sheet_name}", s["h2"]))
        rows: list[list[str]] = []
        for row in ws.iter_rows(values_only=True):
            cleaned = ["" if v is None else str(v) for v in row]
            if any(cell.strip() for cell in cleaned):
                rows.append(cleaned)
        if not rows:
            flow.append(Paragraph("(empty sheet)", s["para"]))
            flow.append(Spacer(1, 6))
            continue
        max_cols = max(len(r) for r in rows)
        rows = [r + [""] * (max_cols - len(r)) for r in rows]
        table = Table(rows, hAlign="LEFT")
        table.setStyle(
            TableStyle([
                ("FONT", (0, 0), (-1, -1), "Helvetica", 9),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#CCCCCC")),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F2F2")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ])
        )
        flow.append(table)
        flow.append(Spacer(1, 8))

    doc = SimpleDocTemplate(
        str(dest), pagesize=LETTER,
        leftMargin=0.6 * inch, rightMargin=0.6 * inch,
        topMargin=0.6 * inch, bottomMargin=0.6 * inch,
        title=src.name,
    )
    doc.build(flow)


def _docx_iter_blocks(doc: Document) -> Iterable[tuple[str, object]]:
    """Yield ('para', Paragraph) and ('table', Table) in document order."""
    body = doc.element.body
    for child in body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            for para in doc.paragraphs:
                if para._element is child:
                    yield ("para", para)
                    break
        elif tag == "tbl":
            for table in doc.tables:
                if table._element is child:
                    yield ("table", table)
                    break


def docx_to_pdf(src: Path, dest: Path) -> None:
    """Render DOCX paragraphs + tables into a basic PDF preserving order."""
    docx_doc = Document(str(src))
    s = _styles()
    flow: list = []
    flow.append(Paragraph(f"Document: {src.name}", s["title"]))

    for kind, obj in _docx_iter_blocks(docx_doc):
        if kind == "para":
            text = (obj.text or "").strip()
            if not text:
                flow.append(Spacer(1, 4))
                continue
            style = obj.style.name if obj.style else ""
            sty = s["h2"] if "Heading" in style else s["para"]
            try:
                flow.append(Paragraph(text.replace("&", "&amp;").replace("<", "&lt;"), sty))
            except Exception:
                flow.append(Paragraph(text, sty))
        elif kind == "table":
            rows: list[list[str]] = []
            for row in obj.rows:
                rows.append([(cell.text or "").strip() for cell in row.cells])
            if not rows:
                continue
            max_cols = max(len(r) for r in rows) if rows else 0
            rows = [r + [""] * (max_cols - len(r)) for r in rows]
            table = Table(rows, hAlign="LEFT")
            table.setStyle(
                TableStyle([
                    ("FONT", (0, 0), (-1, -1), "Helvetica", 9),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#CCCCCC")),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F2F2")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ])
            )
            flow.append(table)
            flow.append(Spacer(1, 6))

    doc = SimpleDocTemplate(
        str(dest), pagesize=LETTER,
        leftMargin=0.6 * inch, rightMargin=0.6 * inch,
        topMargin=0.6 * inch, bottomMargin=0.6 * inch,
        title=src.name,
    )
    doc.build(flow)


def to_pdf(src: Path, dest_dir: Path) -> Path | None:
    """Convert XLSX/DOCX to PDF, write to dest_dir, return the PDF path.
    Returns None if the source format isn't supported here."""
    ext = src.suffix.lower()
    dest_dir.mkdir(parents=True, exist_ok=True)
    out = dest_dir / f"{src.stem}__converted.pdf"
    if ext in (".xlsx", ".xls"):
        xlsx_to_pdf(src, out)
        return out
    if ext in (".docx", ".doc"):
        docx_to_pdf(src, out)
        return out
    return None
