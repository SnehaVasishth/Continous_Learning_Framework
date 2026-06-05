"""Stage 2 — Document Intelligence & Data Extraction.

For PO/Q2O intents: opens attached PDF/XLSX/image and extracts structured fields.
For other intents: extracts the relevant operational fields from the email body.
"""
from __future__ import annotations

import json
import os
from pathlib import Path

import docx
import openpyxl
from pypdf import PdfReader

from .llm import ask_llm

PO_SCHEMA = (
    "{\"po_number\": string, \"quote_number\": string|null, \"customer_name\": string, "
    "\"requested_ship_date\": ISO date string, "
    "\"payment_terms\": string, \"bill_to\": string, \"ship_to\": string, "
    "\"line_items\": [{\"sku\": string, \"description\": string, \"qty\": int, \"unit_price\": number}], "
    "\"total\": number, \"notes\": string}"
)

OPS_SCHEMA = (
    "{\"order_number\": string|null, \"quote_number\": string|null, \"work_order_number\": string|null, "
    "\"asset_serial\": string|null, \"requested_action\": string, \"new_ship_date\": ISO date string|null, "
    "\"service_type\": string|null}"
)

CHANGE_ORDER_SCHEMA = (
    "{\"order_number\": string, \"customer_po\": string|null, \"requested_action\": string, "
    "\"line_changes\": [{\"sku\": string|null, \"description\": string|null, \"change_kind\": one of [\"qty\", \"price\", \"add\", \"remove\", \"swap\"], "
    "\"new_qty\": int|null, \"new_unit_price\": number|null, \"new_sku\": string|null, \"reason\": string|null}], "
    "\"new_bill_to\": string|null, \"new_ship_to\": string|null, \"notes\": string|null}"
)

SSD_SCHEMA = (
    "{\"order_number\": string|null, \"customer_po\": string|null, \"line_skus\": [string]|null, "
    "\"current_ship_date\": ISO date string|null, \"new_ship_date\": ISO date string, "
    "\"direction\": one of [\"push_out\", \"pull_in\", \"partial\"], \"reason\": string, \"notes\": string|null}"
)

SOM_CREATE_SCHEMA = (
    "{\"service_type\": one of [\"calibration\", \"repair\", \"installation\", \"on_site_service\", \"pm\"], "
    "\"standards_referenced\": [string] (e.g. [\"ISO/IEC 17025\", \"ANSI/NCSL Z540.3\"]), "
    "\"assets\": [{\"asset_serial\": string|null, \"sku\": string|null, \"description\": string|null, \"location\": string|null, \"last_cal_date\": ISO date|null, \"oot_observed\": bool|null, \"notes\": string|null}], "
    "\"requested_completion_date\": ISO date|null, \"on_site_required\": bool, \"po_reference\": string|null, "
    "\"contract_reference\": string|null, \"notes\": string|null}"
)

SOM_UPDATE_SCHEMA = (
    "{\"work_order_number\": string|null, \"order_number\": string|null, \"requested_action\": string, "
    "\"add_assets\": [{\"asset_serial\": string|null, \"sku\": string|null, \"description\": string|null}], "
    "\"add_note\": string|null, \"add_task\": string|null, \"notes\": string|null}"
)

SOM_INQUIRY_SCHEMA = (
    "{\"work_order_numbers\": [string], \"asset_serials\": [string], \"customer_po\": string|null, "
    "\"requested_info\": one of [\"status\", \"eta\", \"as_found_data\", \"cert_expiry\", \"all\"], "
    "\"urgency\": one of [\"urgent\", \"normal\", \"low\"], \"notes\": string|null}"
)

SERVICE_CONTRACT_SCHEMA = (
    "{\"contract_type\": one of [\"calibration_plan\", \"onsite_service_plan\", \"pm_plan\", \"warranty_extension\", \"unknown\"], "
    "\"requested_action\": one of [\"quote\", \"renew\", \"order\", \"info\"], "
    "\"existing_contract_number\": string|null, \"asset_count_estimate\": int|null, "
    "\"included_skus\": [string], \"asset_serials\": [string], \"term_months\": int|null, "
    "\"sla_tier_requested\": string|null, \"start_date\": ISO date|null, \"notes\": string|null}"
)

SYSTEM_PO = (
    "You are a document-intelligence agent. Extract structured PO data from the email body and any provided "
    "attachments (text from PDFs, Excel BOMs, DOCX specs, or images). "
    "Return strict JSON matching this schema: " + PO_SCHEMA + ". "
    "Important: quote_number is the referenced quote ID from the email or attachments. "
    "it usually starts with 'Q-', 'QT-', or 'QUOTE-' (e.g. 'QT-AURA-AUTO-119-DEMO'). "
    "Pull it whether it's mentioned in the email body, the BOM, or the PO. "
    "If a field is genuinely missing, use null. Do not invent values."
)

SYSTEM_OPS = (
    "You are an operations-fields extraction agent. Read the customer email and pull out actionable fields. "
    "Return strict JSON matching this schema: " + OPS_SCHEMA + ". "
    "Use null for any field not present. requested_action should be a short imperative phrase."
)

SYSTEM_CHANGE_ORDER = (
    "You are a Trade Sales Change Order extraction agent. The customer is asking to MODIFY an EXISTING booked order. "
    "Pull every change they're requesting line-by-line. Return strict JSON matching this schema: "
    + CHANGE_ORDER_SCHEMA + ". "
    "change_kind values: qty (quantity change), price (negotiated price change), add (add new line), remove (cancel line), swap (replace SKU). "
    "If the customer doesn't specify the order_number, leave it null; downstream will fuzzy-match by customer + customer_po."
)

SYSTEM_SSD = (
    "You are a Ship Schedule Date (SSD) change extraction agent. "
    "Customer is asking to change the ship date on an existing order. Return strict JSON matching this schema: "
    + SSD_SCHEMA + ". "
    "direction: 'push_out' (later than current), 'pull_in' (earlier than current), 'partial' (split shipment). "
    "If multiple orders are referenced, choose the primary one for order_number and list affected SKUs in line_skus."
)

SYSTEM_SOM_CREATE = (
    "You are a Service Order Management (SOM) extraction agent for NEW work-order creation requests. "
    "Customer is asking for calibration, repair, installation, or on-site service. They may include MULTIPLE assets "
    "(in the email body or as a spreadsheet attachment). Return strict JSON matching this schema: "
    + SOM_CREATE_SCHEMA + ". "
    "If the customer attaches a spreadsheet listing multiple instruments, return ONE object per asset in the assets[] array. "
    "Pull standards from the email; common references include ISO/IEC 17025, ANSI/NCSL Z540.3, A2LA, MIL-STD-810. "
    "on_site_required: true if they mention on-site / field service / 'come to our lab'."
)

SYSTEM_SOM_UPDATE = (
    "You are a SOM update extraction agent. Customer is asking to update an EXISTING work order: "
    "add a note, add a task, add additional assets to an open WO. Return strict JSON matching this schema: "
    + SOM_UPDATE_SCHEMA + "."
)

SYSTEM_SOM_INQUIRY = (
    "You are a SOM status-inquiry extraction agent. Pull every WO number, asset serial, and PO ref the customer "
    "is asking about. Return strict JSON matching this schema: " + SOM_INQUIRY_SCHEMA + ". "
    "urgency: 'urgent' if the email contains words like URGENT, ASAP, audit Friday, escalation; 'normal' otherwise."
)

SYSTEM_SERVICE_CONTRACT = (
    "You are a service-contract extraction agent. Customer is asking about a service plan / cal contract / "
    "support agreement. Could be a quote request, renewal, order, or information question. "
    "Return strict JSON matching this schema: " + SERVICE_CONTRACT_SCHEMA + "."
)


_FALLBACK_SYSTEM_BY_INTENT = {
    "po_intake": (SYSTEM_PO, True),
    "quote_to_order": (SYSTEM_PO, True),
    "trade_change_order": (SYSTEM_CHANGE_ORDER, False),
    "ssd_change_request": (SYSTEM_SSD, False),
    "delivery_change": (SYSTEM_SSD, False),
    "hold_release": (SYSTEM_OPS, False),
    "service_order": (SYSTEM_SOM_CREATE, True),
    "wo_update_request": (SYSTEM_SOM_UPDATE, False),
    "wo_status_inquiry": (SYSTEM_SOM_INQUIRY, False),
    "service_contract_request": (SYSTEM_SERVICE_CONTRACT, False),
    "general_inquiry": (SYSTEM_OPS, False),
}

_INTENTS_ALLOWING_IMAGES = {"po_intake", "quote_to_order", "service_order"}


def _system_prompt_from_kb_schema(kb_body: dict) -> str:
    """Build a fresh SYSTEM prompt from a KB extract_schema rule body."""
    sp = (kb_body.get("system_prompt") or "").strip()
    fields = kb_body.get("fields") or []
    if not fields:
        return sp
    field_lines = []
    for f in fields:
        name = f.get("name", "?")
        ftype = f.get("type", "any")
        req = "required" if f.get("required") else "optional"
        desc = f.get("description") or ""
        field_lines.append(f"  - {name} ({ftype}, {req}): {desc}")
    return (
        sp
        + "\n\nReturn strict JSON with these fields:\n"
        + "\n".join(field_lines)
        + "\n\nIf a field is genuinely missing, use null. Do not invent values."
    )


def _read_pdf_text(path: Path) -> str:
    try:
        reader = PdfReader(str(path))
        return "\n\n".join((p.extract_text() or "") for p in reader.pages)
    except Exception as e:
        return f"[PDF read failed: {e}]"


def _read_docx_text(path: Path) -> str:
    try:
        doc = docx.Document(str(path))
        parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception as e:
        return f"[DOCX read failed: {e}]"


def _read_xlsx_text(path: Path) -> str:
    wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    out = []
    for ws in wb.worksheets:
        out.append(f"# Sheet: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(cells):
                out.append(" | ".join(cells))
    wb.close()
    return "\n".join(out)


def _make_preview(text: str, kind: str) -> str:
    """Trimmed snippet (first ~12 lines / 800 chars) used for the trace UI."""
    if not text:
        return ""
    lines = [ln for ln in text.splitlines() if ln.strip()]
    snippet = "\n".join(lines[:14])
    if len(snippet) > 900:
        snippet = snippet[:900] + " …"
    return snippet


def _attachment_block(
    attachments: list[dict], on_event: callable | None = None
) -> tuple[str, list[Path]]:
    text_parts: list[str] = []
    image_paths: list[Path] = []
    for a in attachments or []:
        path = Path(a.get("path", ""))
        kind = (a.get("type") or "").lower()
        name = a.get("name", "")
        meta: dict = {"name": name, "type": kind}
        if not path.exists():
            meta["status"] = "missing"
            if on_event:
                on_event(name, kind, meta, "missing")
            continue
        size_bytes = path.stat().st_size
        meta["size_bytes"] = size_bytes
        try:
            if kind == "pdf":
                text = _read_pdf_text(path)
                text_parts.append(f"--- PDF ATTACHMENT: {name} ---\n{text}")
                meta["chars_extracted"] = len(text)
                meta["preview"] = _make_preview(text, kind)
                meta["status"] = "extracted"
                if on_event:
                    on_event(name, kind, meta, "extracted")
            elif kind == "xlsx":
                text = _read_xlsx_text(path)
                text_parts.append(f"--- XLSX ATTACHMENT: {name} ---\n{text}")
                meta["chars_extracted"] = len(text)
                meta["preview"] = _make_preview(text, kind)
                meta["status"] = "extracted"
                if on_event:
                    on_event(name, kind, meta, "extracted")
            elif kind == "docx":
                text = _read_docx_text(path)
                text_parts.append(f"--- DOCX ATTACHMENT: {name} ---\n{text}")
                meta["chars_extracted"] = len(text)
                meta["preview"] = _make_preview(text, kind)
                meta["status"] = "extracted"
                if on_event:
                    on_event(name, kind, meta, "extracted")
            elif kind == "image":
                image_paths.append(path)
                text_parts.append(
                    f"--- IMAGE ATTACHMENT (will OCR via vision): {name} at {path.resolve().as_posix()} ---"
                )
                meta["status"] = "queued_for_vision_ocr"
                if on_event:
                    on_event(name, kind, meta, "queued_for_vision_ocr")
            else:
                meta["status"] = "unknown_type"
                if on_event:
                    on_event(name, kind, meta, "unknown_type")
        except Exception as e:
            meta["status"] = "error"
            meta["error"] = str(e)[:200]
            if on_event:
                on_event(name, kind, meta, "error")
    return ("\n\n".join(text_parts), image_paths)


def _resolve_system(intent: str) -> tuple[str, bool]:
    """Pull the active extract schema for this intent from the KB.
    Falls back to the original hard-coded schema if no KB rule covers it."""
    from .. import kb as _kb

    kb_body = _kb.extract_schema_for(intent)
    if kb_body and (kb_body.get("system_prompt") or kb_body.get("fields")):
        return _system_prompt_from_kb_schema(kb_body), intent in _INTENTS_ALLOWING_IMAGES
    return _FALLBACK_SYSTEM_BY_INTENT.get(intent, (SYSTEM_OPS, False))


def run_extract(
    *,
    email: dict,
    intake: dict,
    on_attachment: callable | None = None,
    thread_summary: str | None = None,
) -> dict:
    """Schema-driven structured extraction. Returns a dict matching the intent's KB extract_schema.

    Provider chain:
      1. OpenAI gpt-5.2 with response_format=json_object (when OPENAI_API_KEY set) — preferred
      2. Claude Agent SDK (legacy fallback)

    Embeds an extra `_provider`, `_provider_meta`, `_prompt_system`, `_prompt_user`,
    and `_provider_response_raw` in the result so Stage 2 sub-step 2.2 can surface
    the full LLM trace in the UI.

    When `thread_summary` is provided (multi-message conversation), the LLM is shown
    the chronological chain so it can extract fields that surface in *replies* — e.g.
    revised qty in message 5 of an 8-message PO thread, signed end-use statements
    sent in a later reply, etc.
    """
    from ..services import openai_client

    intent = intake.get("intent") or "general_inquiry"
    attachments_text, image_paths = _attachment_block(email.get("attachments") or [], on_event=on_attachment)
    body = email["body"]
    common = (
        f"FROM: {email['from']}\n"
        f"SUBJECT: {email['subject']}\n"
        f"INTENT: {intent}\n"
        f"BODY:\n{body}\n"
        f"\nATTACHMENTS:\n{attachments_text or '(none)'}\n"
    )
    if thread_summary:
        # Place the thread BEFORE the latest-envelope block so the LLM reads
        # the full conversation context first. Latest message remains explicit
        # so it knows what triggered this pipeline run.
        common = (
            "EMAIL THREAD (chronological; extract from the entire conversation, "
            "not just the latest message; later replies may amend / supersede earlier values):\n"
            f"{thread_summary}\n\n"
            "LATEST MESSAGE THAT TRIGGERED THIS RUN:\n" + common
        )
    user_prompt = common + "\nReturn JSON only. No prose, no code fences, no extra keys."
    system, allow_images = _resolve_system(intent)

    out: dict = {}
    provider_used: str | None = None
    provider_meta: dict = {}
    raw_response: str = ""

    if openai_client.is_configured():
        try:
            from openai import OpenAI
            client = OpenAI()
            model = os.environ.get("OPENAI_MODEL", "gpt-5.2")
            for attempt_model in [model, "gpt-5", "gpt-4.1", "gpt-4o"]:
                try:
                    resp = client.chat.completions.create(
                        model=attempt_model,
                        messages=[
                            {"role": "system", "content": system},
                            {"role": "user", "content": user_prompt},
                        ],
                        response_format={"type": "json_object"},
                        temperature=0.0,
                    )
                    raw_response = (resp.choices[0].message.content or "").strip()
                    try:
                        out = json.loads(raw_response)
                        provider_used = f"OpenAI ({attempt_model}) · response_format=json_object"
                        provider_meta = {
                            "provider": provider_used,
                            "model": attempt_model,
                            "finish_reason": resp.choices[0].finish_reason,
                            "usage": getattr(resp, "usage", None) and {
                                "prompt_tokens": resp.usage.prompt_tokens,
                                "completion_tokens": resp.usage.completion_tokens,
                                "total_tokens": resp.usage.total_tokens,
                            },
                        }
                        try:
                            from .llm import record_llm_cost
                            _u = provider_meta.get("usage") or {}
                            record_llm_cost(
                                model_hint=attempt_model,
                                tokens_in=_u.get("prompt_tokens"),
                                tokens_out=_u.get("completion_tokens"),
                                system=system,
                                user=user_prompt,
                                raw_response=raw_response,
                                tool="extract_openai",
                            )
                        except Exception:
                            import logging as _logging
                            _logging.getLogger(__name__).exception("extract openai cost metering failed")
                        break
                    except Exception as je:
                        logger_msg = f"openai_json_parse_failed_on_{attempt_model}: {je}"
                        provider_meta["last_parse_error"] = logger_msg
                        continue
                except Exception as me:
                    msg = str(me)
                    if any(needle in msg.lower() for needle in (
                        "model_not_found", "does not exist", "model_does_not_exist",
                        "unsupported_value", "is not supported with",
                    )):
                        continue
                    provider_meta["last_call_error"] = f"{type(me).__name__}: {msg[:200]}"
                    break
        except Exception as e:
            provider_meta["fatal_error"] = f"{type(e).__name__}: {str(e)[:200]}"

    # Fallback: Claude Agent SDK (legacy)
    if not out:
        try:
            out = ask_llm(
                system=system,
                user=user_prompt,
                json_only=True,
                image_paths=[str(p) for p in image_paths] if (allow_images and image_paths) else None,
            )
            provider_used = provider_used or "ZBrain LLM (Claude Opus 4.7)"
            provider_meta.setdefault("provider", provider_used)
        except ValueError as e:
            return {
                "_extract_error": str(e)[:300],
                "_attachments_seen": [a.get("name") for a in (email.get("attachments") or [])],
                "_provider": provider_used,
                "_provider_meta": provider_meta,
                "_prompt_system": system,
                "_prompt_user": user_prompt,
                "_provider_response_raw": raw_response,
            }

    if isinstance(out, dict):
        out.setdefault("_provider", provider_used)
        out.setdefault("_provider_meta", provider_meta)
        out.setdefault("_prompt_system", system)
        out.setdefault("_prompt_user", user_prompt)
        out.setdefault("_provider_response_raw", raw_response)
    return out
