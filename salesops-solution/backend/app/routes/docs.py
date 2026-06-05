"""Serves the SOLUTION.md and SOLUTION_OVERVIEW.md docs for in-app viewer pages."""
import io
from datetime import datetime
from html import escape as _html_escape
from pathlib import Path
from typing import Optional

from fastapi import APIRouter, Depends, HTTPException, Request
from fastapi.responses import HTMLResponse, JSONResponse, PlainTextResponse, Response
from pydantic import BaseModel
from sqlalchemy.orm import Session

from ..config import ROOT
from ..db import get_db
# === v1.1 RFP-REPLY ===
from ..services.rfp_reply_docx import SECTIONS as RFP_SECTIONS, build_rfp_reply_docx, get_section
# === v1.1 RFP-VIEWER ===
try:
    import mammoth  # type: ignore[import-not-found]
except Exception:  # pragma: no cover - optional dependency in some local envs
    mammoth = None
# === v1.1 DOC-FEEDBACK ===
from ..models import DocFeedback

router = APIRouter()

SOLUTION_MD = ROOT.parent / "SOLUTION.md"
SOLUTION_OVERVIEW_MD = ROOT.parent / "SOLUTION_OVERVIEW.md"
# === v1.1 AS-IS-DOC START ===
AS_IS_MD = ROOT.parent / "AS_IS_PROCESS_SOP.md"
# === v1.1 AS-IS-DOC END ===


@router.get("/solution", response_class=PlainTextResponse)
def solution() -> str:
    if not SOLUTION_MD.exists():
        raise HTTPException(404, "SOLUTION.md not found")
    return SOLUTION_MD.read_text(encoding="utf-8")


@router.get("/solution-overview", response_class=PlainTextResponse)
def solution_overview() -> str:
    if not SOLUTION_OVERVIEW_MD.exists():
        raise HTTPException(404, "SOLUTION_OVERVIEW.md not found")
    return SOLUTION_OVERVIEW_MD.read_text(encoding="utf-8")


# === v1.1 AS-IS-DOC START ===
@router.get("/as-is.md", response_class=PlainTextResponse)
def as_is_markdown() -> PlainTextResponse:
    if not AS_IS_MD.exists():
        raise HTTPException(404, "AS_IS_PROCESS_SOP.md not found")
    return PlainTextResponse(
        AS_IS_MD.read_text(encoding="utf-8"),
        headers={"Content-Disposition": 'attachment; filename="AS_IS_PROCESS_SOP.md"'},
        media_type="text/markdown; charset=utf-8",
    )


@router.get("/as-is", response_class=HTMLResponse)
def as_is_html() -> str:
    if not AS_IS_MD.exists():
        raise HTTPException(404, "AS_IS_PROCESS_SOP.md not found")
    md = AS_IS_MD.read_text(encoding="utf-8")
    md_json = __import__("json").dumps(md)
    return f"""<!doctype html>
<html lang=en>
<head>
<meta charset=utf-8>
<title>AS-IS Process SOP — Keysight SalesOps</title>
<meta name=viewport content="width=device-width, initial-scale=1">
<script src="https://cdn.jsdelivr.net/npm/marked/marked.min.js"></script>
<style>
  :root {{ --ink: #131426; --rule: #e5e7eb; --accent: #1A55F9; --muted: #6B7280; --surface: #f8fafc; }}
  body {{ font-family: -apple-system, BlinkMacSystemFont, Inter, "Segoe UI", Roboto, sans-serif;
         max-width: 920px; margin: 40px auto; padding: 0 24px 80px; line-height: 1.55; color: var(--ink); }}
  h1, h2, h3 {{ line-height: 1.25; }}
  h1 {{ border-bottom: 2px solid var(--accent); padding-bottom: .35em; }}
  h2 {{ border-bottom: 1px solid var(--rule); padding-bottom: .25em; margin-top: 2em; }}
  table {{ border-collapse: collapse; margin: 1em 0; font-size: .92em; }}
  th, td {{ border: 1px solid var(--rule); padding: 6px 10px; text-align: left; vertical-align: top; }}
  th {{ background: var(--surface); }}
  code {{ background: #f5f5f7; padding: 1px 5px; border-radius: 3px; font-size: .92em; }}
  pre code {{ display: block; padding: 12px; overflow-x: auto; }}
  blockquote {{ border-left: 4px solid var(--accent); margin: 1em 0; padding: .25em 1em; background: var(--surface); color: #374151; }}
  a {{ color: var(--accent); }}
  img {{ max-width: 100%; height: auto; display: block; margin: 1em auto; border-radius: 4px; }}

  /* Sticky download bar (hidden in print) */
  .toolbar {{
    position: sticky; top: 0; z-index: 10;
    background: rgba(255,255,255,.94); backdrop-filter: blur(8px);
    border-bottom: 1px solid var(--rule);
    margin: -40px -24px 24px; padding: 12px 24px;
    display: flex; gap: 8px; align-items: center; flex-wrap: wrap;
  }}
  .toolbar .title {{ font-weight: 600; font-size: 14px; color: var(--muted); margin-right: auto; }}
  .toolbar button, .toolbar a.btn {{
    appearance: none; border: 1px solid var(--rule); background: white; color: var(--ink);
    font: inherit; font-size: 13px; font-weight: 500; padding: 7px 14px; border-radius: 8px;
    cursor: pointer; text-decoration: none; transition: all .15s ease;
  }}
  .toolbar button.primary {{ background: var(--accent); color: white; border-color: var(--accent); }}
  .toolbar button:hover, .toolbar a.btn:hover {{ box-shadow: 0 1px 3px rgba(0,0,0,.08); transform: translateY(-1px); }}

  /* PRINT — Save-as-PDF flow.
     Page size A4 with margins; toolbar hidden; images kept inside a single page,
     never split across pages; headings stay with their following paragraph. */
  @page {{ size: A4; margin: 18mm 14mm; }}
  @media print {{
    body {{ margin: 0; padding: 0; max-width: 100%; color: #000; }}
    .toolbar {{ display: none !important; }}
    h1 {{ font-size: 22pt; }}
    h2 {{ font-size: 16pt; page-break-before: auto; break-before: auto; page-break-after: avoid; break-after: avoid; }}
    h3, h4 {{ page-break-after: avoid; break-after: avoid; }}

    /* The critical rule: keep every diagram on a single page. */
    img {{
      max-width: 100%;
      max-height: 92vh;
      page-break-inside: avoid;
      break-inside: avoid;
      page-break-before: auto;
      page-break-after: auto;
      display: block;
      margin: 0 auto;
    }}
    /* Same protection on figures / picture wrappers if the renderer adds them */
    figure, picture {{ page-break-inside: avoid; break-inside: avoid; }}

    /* Keep tables together unless they're truly long */
    table {{ page-break-inside: avoid; break-inside: avoid; }}
    tr {{ page-break-inside: avoid; break-inside: avoid; }}

    pre, blockquote {{ page-break-inside: avoid; break-inside: avoid; }}
    p {{ orphans: 3; widows: 3; }}
    a {{ color: #000; text-decoration: none; }}
  }}
</style>
</head>
<body>
<div class="toolbar">
  <span class="title">AS-IS Process SOP · Keysight SalesOps</span>
  <button class="primary" onclick="window.print()" title="Use your browser's Save-as-PDF in the print dialog">⬇ Save as PDF</button>
  <a class="btn" href="/api/docs/as-is.md" download="AS_IS_PROCESS_SOP.md">Download .md</a>
</div>
<div id=doc></div>
<script>
  document.getElementById('doc').innerHTML = marked.parse({md_json});
</script>
</body>
</html>"""
# === v1.1 AS-IS-DOC END ===


# === v1.1 RFP-REPLY START ===
@router.get("/rfp-reply.docx")
def rfp_reply_combined_docx() -> Response:
    """Kept for compatibility; the response is now twelve separate documents.
    This URL returns the Executive Summary so the legacy link still resolves."""
    blob = build_rfp_reply_docx()
    return Response(
        content=blob,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": 'attachment; filename="Keysight_RFP_01_Executive_Summary.docx"',
            "Cache-Control": "no-store",
        },
    )


@router.get("/rfp-reply/{slug}.docx")
def rfp_reply_section_docx(slug: str) -> Response:
    section = get_section(slug)
    if section is None:
        raise HTTPException(404, f"Unknown section: {slug}")
    blob = section["builder"]()
    return Response(
        content=blob,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{section["filename"]}"',
            "Cache-Control": "no-store",
        },
    )


# === v1.1 RFP-VIEWER START ===
def _rfp_viewer_shell(slug: str, label: str, inner_html: str) -> str:
    """Wrap the rendered section HTML in a viewer template with sidebar nav,
    download button, and print-friendly CSS."""
    # Build sidebar with all 12 sections plus the AS-IS link.
    nav_items = []
    for s in RFP_SECTIONS:
        active = "active" if s["slug"] == slug else ""
        parts = s["label"].split(".", 1)
        if len(parts) == 2 and parts[0].strip().isdigit():
            navnum, navlabel = parts[0].strip(), parts[1].strip()
        else:
            navnum, navlabel = "", s["label"].strip()
        nav_items.append(
            f'<a class="navitem {active}" href="/api/docs/rfp-reply/{s["slug"]}">'
            f'<span class="navnum">{navnum}</span>'
            f'<span class="navlabel">{navlabel}</span>'
            f'</a>'
        )
    sidebar = "\n".join(nav_items)

    # Previous / next links
    idx = next((i for i, s in enumerate(RFP_SECTIONS) if s["slug"] == slug), 0)
    prev_link = (
        f'<a class="pn-btn" href="/api/docs/rfp-reply/{RFP_SECTIONS[idx-1]["slug"]}">'
        f'&larr; {RFP_SECTIONS[idx-1]["label"]}</a>'
        if idx > 0 else '<span class="pn-btn disabled">&larr; First</span>'
    )
    next_link = (
        f'<a class="pn-btn" href="/api/docs/rfp-reply/{RFP_SECTIONS[idx+1]["slug"]}">'
        f'{RFP_SECTIONS[idx+1]["label"]} &rarr;</a>'
        if idx < len(RFP_SECTIONS) - 1 else '<span class="pn-btn disabled">Last &rarr;</span>'
    )

    return f"""<!doctype html>
<html lang=en>
<head>
<meta charset=utf-8>
<title>{label} &middot; Keysight RFP Response</title>
<meta name=viewport content="width=device-width, initial-scale=1">
<style>
  :root {{ --ink: #131426; --rule: #e5e7eb; --accent: #1A55F9; --muted: #6B7280;
           --surface: #f8fafc; --sidebar: #f9fafb; }}
  * {{ box-sizing: border-box; }}
  html, body {{ margin: 0; padding: 0; }}
  body {{ font-family: -apple-system, BlinkMacSystemFont, Inter, "Segoe UI", Roboto, sans-serif;
         color: var(--ink); background: white; line-height: 1.55; }}

  /* Layout: sidebar + main */
  .app {{ display: grid; grid-template-columns: 280px 1fr; min-height: 100vh; }}
  .sidebar {{
    background: var(--sidebar); border-right: 1px solid var(--rule);
    padding: 16px 12px 24px; position: sticky; top: 0; height: 100vh;
    overflow-y: auto;
  }}
  .sidebar h1 {{ font-size: 13px; margin: 8px 12px 14px; color: var(--muted);
                 text-transform: uppercase; letter-spacing: .08em; font-weight: 600; }}
  .navitem {{
    display: flex; align-items: baseline; gap: 8px;
    padding: 9px 12px; margin: 2px 0; border-radius: 8px;
    color: var(--ink); text-decoration: none; font-size: 13.5px; line-height: 1.35;
    transition: background .12s;
  }}
  .navitem:hover {{ background: white; }}
  .navitem.active {{ background: var(--accent); color: white; font-weight: 600; }}
  .navitem.active .navnum {{ color: white; opacity: .85; }}
  .navnum {{ flex: 0 0 24px; color: var(--muted); font-weight: 600;
             font-variant-numeric: tabular-nums; }}
  .navlabel {{ flex: 1 1 auto; }}
  .sidebar-foot {{ margin-top: 18px; padding: 14px 12px 0;
                   border-top: 1px solid var(--rule); }}
  .sidebar-foot a {{ display: block; color: var(--accent); font-size: 13px;
                     text-decoration: none; padding: 6px 0; }}
  .sidebar-foot a:hover {{ text-decoration: underline; }}

  /* Main column */
  .main {{ min-width: 0; }}
  .toolbar {{
    position: sticky; top: 0; z-index: 10;
    background: rgba(255,255,255,.94); backdrop-filter: blur(8px);
    border-bottom: 1px solid var(--rule);
    padding: 14px 32px; display: flex; align-items: center; gap: 10px;
  }}
  .toolbar .crumb {{ font-size: 12px; color: var(--muted); margin-right: auto; }}
  .toolbar button, .toolbar a.btn {{
    appearance: none; border: 1px solid var(--rule); background: white; color: var(--ink);
    font: inherit; font-size: 13px; font-weight: 500; padding: 7px 14px;
    border-radius: 8px; cursor: pointer; text-decoration: none;
    transition: all .15s ease; display: inline-flex; align-items: center; gap: 6px;
  }}
  .toolbar button.primary {{ background: var(--accent); color: white;
                              border-color: var(--accent); }}
  .toolbar button:hover, .toolbar a.btn:hover {{
    box-shadow: 0 1px 3px rgba(0,0,0,.08); transform: translateY(-1px);
  }}

  .content {{ max-width: 880px; padding: 28px 48px 80px; margin: 0 auto; }}

  /* Content typography (mammoth-rendered) */
  .content h1 {{ font-size: 26px; color: var(--accent);
                 border-bottom: 2px solid var(--accent);
                 padding-bottom: .35em; margin-top: 0; }}
  .content h2 {{ font-size: 19px; border-bottom: 1px solid var(--rule);
                 padding-bottom: .2em; margin-top: 1.8em; }}
  .content h3 {{ font-size: 16px; margin-top: 1.4em; }}
  .content p {{ font-size: 14.5px; }}
  .content ul, .content ol {{ font-size: 14.5px; padding-left: 22px; }}
  .content li {{ margin: 4px 0; }}
  .content table {{ border-collapse: collapse; margin: 1.2em 0;
                    font-size: 13px; width: 100%; }}
  .content th, .content td {{ border: 1px solid var(--rule);
                              padding: 8px 11px; text-align: left;
                              vertical-align: top; }}
  .content th {{ background: var(--accent); color: white; font-weight: 600; }}
  .content code {{ background: #f5f5f7; padding: 1px 5px;
                   border-radius: 3px; font-size: .92em; }}
  .content em {{ color: #374151; }}
  .content strong {{ color: var(--ink); }}

  .pagenav {{ display: flex; gap: 12px; margin-top: 48px;
              padding-top: 24px; border-top: 1px solid var(--rule);
              justify-content: space-between; }}
  .pn-btn {{
    flex: 0 1 auto; padding: 10px 16px; border: 1px solid var(--rule);
    border-radius: 8px; text-decoration: none; color: var(--ink);
    font-size: 13.5px; transition: all .15s; max-width: 45%;
  }}
  .pn-btn:hover:not(.disabled) {{ border-color: var(--accent);
                                   color: var(--accent); }}
  .pn-btn.disabled {{ color: var(--muted); cursor: not-allowed; opacity: .5; }}

  /* === v1.1 DOC-FEEDBACK START === floating feedback chat */
  .fb-fab {{
    position: fixed; right: 22px; bottom: 22px; z-index: 30;
    background: var(--accent); color: white; border: none; cursor: pointer;
    width: 56px; height: 56px; border-radius: 50%; font-size: 22px;
    box-shadow: 0 6px 18px rgba(26,85,249,.32); display: grid; place-items: center;
    transition: all .15s ease;
  }}
  .fb-fab:hover {{ transform: translateY(-2px); box-shadow: 0 10px 26px rgba(26,85,249,.4); }}
  .fb-fab .fb-count {{
    position: absolute; top: -4px; right: -4px; background: #C97A0B; color: white;
    border-radius: 999px; min-width: 20px; height: 20px; font-size: 11.5px;
    font-weight: 600; line-height: 20px; padding: 0 6px; border: 2px solid white;
  }}
  .fb-panel {{
    position: fixed; right: 22px; bottom: 92px; z-index: 30;
    width: 380px; max-width: calc(100vw - 44px);
    max-height: calc(100vh - 130px);
    background: white; border: 1px solid var(--rule); border-radius: 14px;
    box-shadow: 0 12px 40px rgba(15,30,80,.18);
    display: none; flex-direction: column; overflow: hidden;
  }}
  .fb-panel.open {{ display: flex; }}
  .fb-head {{
    padding: 14px 16px; border-bottom: 1px solid var(--rule);
    display: flex; align-items: center; gap: 8px;
    background: linear-gradient(180deg, #fafbfd, white);
  }}
  .fb-head .fb-title {{ font-weight: 600; font-size: 14px; color: var(--ink); }}
  .fb-head .fb-sub {{ font-size: 11.5px; color: var(--muted); margin-left: auto; }}
  .fb-head .fb-close {{
    background: transparent; border: none; cursor: pointer; color: var(--muted);
    font-size: 20px; padding: 0 4px; margin-left: 6px;
  }}
  .fb-stream {{
    flex: 1 1 auto; overflow-y: auto; padding: 14px 16px;
    display: flex; flex-direction: column; gap: 10px;
    background: #fafbfd;
  }}
  .fb-empty {{ text-align: center; color: var(--muted); font-size: 12.5px; padding: 28px 8px; }}
  .fb-msg {{
    background: white; border: 1px solid var(--rule); border-radius: 12px;
    padding: 9px 12px; font-size: 13px; color: var(--ink); line-height: 1.45;
    word-wrap: break-word;
  }}
  .fb-msg.addressed {{ border-color: var(--auto, #1F8A4C); background: #E1F4E8; }}
  .fb-msg.closed {{ opacity: .55; }}
  .fb-meta {{
    display: flex; justify-content: space-between; gap: 6px;
    font-size: 10.5px; color: var(--muted); margin-top: 6px;
  }}
  .fb-section {{
    display: inline-block; background: var(--accent); color: white;
    padding: 1px 7px; border-radius: 999px; font-size: 10.5px; margin-bottom: 4px;
  }}
  .fb-actions {{ display: inline-flex; gap: 4px; }}
  .fb-tiny {{
    background: white; border: 1px solid var(--rule); cursor: pointer;
    font-size: 10.5px; padding: 2px 7px; border-radius: 6px; color: var(--muted);
  }}
  .fb-tiny:hover {{ color: var(--ink); border-color: var(--ink); }}
  .fb-form {{
    padding: 12px 14px; border-top: 1px solid var(--rule); background: white;
    display: flex; flex-direction: column; gap: 8px;
  }}
  .fb-form input, .fb-form textarea {{
    width: 100%; border: 1px solid var(--rule); border-radius: 8px;
    padding: 8px 10px; font: inherit; font-size: 13px; color: var(--ink);
    background: white; resize: none;
  }}
  .fb-form input:focus, .fb-form textarea:focus {{
    outline: none; border-color: var(--accent);
    box-shadow: 0 0 0 3px rgba(26,85,249,.12);
  }}
  .fb-form textarea {{ min-height: 64px; max-height: 140px; }}
  .fb-form-row {{ display: flex; gap: 8px; align-items: center; }}
  .fb-send {{
    background: var(--accent); color: white; border: none; border-radius: 8px;
    padding: 8px 14px; font: inherit; font-size: 13px; font-weight: 600;
    cursor: pointer; margin-left: auto;
  }}
  .fb-send:hover {{ filter: brightness(1.08); }}
  .fb-send:disabled {{ opacity: .5; cursor: not-allowed; }}
  /* === v1.1 DOC-FEEDBACK END === */

  /* PRINT: hide sidebar + toolbar, keep diagrams whole */
  @page {{ size: A4; margin: 18mm 14mm; }}
  @media print {{
    .sidebar, .toolbar, .pagenav, .fb-fab, .fb-panel {{ display: none !important; }}
    .app {{ display: block; }}
    .main {{ width: 100%; }}
    .content {{ max-width: 100%; padding: 0; margin: 0; }}
    .content h2 {{ page-break-before: auto; page-break-after: avoid;
                   break-after: avoid; }}
    .content h3, .content h4 {{ page-break-after: avoid; break-after: avoid; }}
    .content img {{ max-width: 100%; max-height: 92vh;
                    page-break-inside: avoid; break-inside: avoid;
                    display: block; margin: 0 auto; }}
    .content table {{ page-break-inside: avoid; break-inside: avoid;
                      font-size: 10px; }}
    .content tr {{ page-break-inside: avoid; break-inside: avoid; }}
    .content p {{ orphans: 3; widows: 3; }}
  }}

  /* Responsive: collapse sidebar on narrow screens */
  @media (max-width: 880px) {{
    .app {{ grid-template-columns: 1fr; }}
    .sidebar {{ position: relative; height: auto;
                border-right: none; border-bottom: 1px solid var(--rule); }}
    .content {{ padding: 20px; }}
  }}
</style>
</head>
<body>
<div class="app">
  <aside class="sidebar">
    <h1>RFP Response</h1>
    {sidebar}
    <div class="sidebar-foot">
      <a href="/api/docs/as-is">AS-IS Process SOP</a>
      <a href="/api/docs/as-is.md">AS-IS as Markdown</a>
      <a href="/api/docs/rfp-reply">Download centre</a>
    </div>
  </aside>
  <main class="main">
    <div class="toolbar">
      <span class="crumb">{label} &middot; Keysight RFP Response</span>
      <button class="primary" onclick="window.print()" title="Use your browser's Save as PDF in the print dialog">&#9662; Save as PDF</button>
      <a class="btn" href="/api/docs/rfp-reply/{slug}.docx">&#9662; Download .docx</a>
    </div>
    <div class="content">
      {inner_html}
      <nav class="pagenav">
        {prev_link}
        {next_link}
      </nav>
    </div>
  </main>
</div>

<!-- === v1.1 DOC-FEEDBACK START === floating reviewer feedback / chat -->
<button id="fbFab" class="fb-fab" title="Leave feedback on this section">
  <span aria-hidden="true">&#128172;</span>
  <span id="fbCount" class="fb-count" style="display:none;">0</span>
</button>

<div id="fbPanel" class="fb-panel" role="dialog" aria-label="Feedback panel">
  <div class="fb-head">
    <span class="fb-title">Feedback &middot; {label}</span>
    <span class="fb-sub" id="fbSubCount">0 comments</span>
    <button class="fb-close" id="fbClose" aria-label="Close">&times;</button>
  </div>
  <div id="fbStream" class="fb-stream">
    <div class="fb-empty" id="fbEmpty">No feedback yet. Drop a note below &mdash; it'll be picked up and used to update the content.</div>
  </div>
  <form id="fbForm" class="fb-form" onsubmit="return fbSubmit(event)">
    <input type="text" id="fbAnchor" placeholder="Section ref (optional, e.g. &sect;5.1 Classification)" maxlength="160" />
    <textarea id="fbText" placeholder="What should change or what's missing?" required></textarea>
    <div class="fb-form-row">
      <span style="font-size:11px;color:var(--muted);">Comments are stored locally; pull them with <code style="font-size:11px;">GET /api/docs/feedback?doc={slug}</code></span>
      <button class="fb-send" type="submit" id="fbSend">Send</button>
    </div>
  </form>
</div>

<script>
  const FB_DOC = "{slug}";
  const fbFab = document.getElementById('fbFab');
  const fbPanel = document.getElementById('fbPanel');
  const fbClose = document.getElementById('fbClose');
  const fbForm = document.getElementById('fbForm');
  const fbText = document.getElementById('fbText');
  const fbAnchor = document.getElementById('fbAnchor');
  const fbStream = document.getElementById('fbStream');
  const fbEmpty = document.getElementById('fbEmpty');
  const fbCount = document.getElementById('fbCount');
  const fbSubCount = document.getElementById('fbSubCount');
  const fbSend = document.getElementById('fbSend');

  function fbOpen() {{ fbPanel.classList.add('open'); fbText.focus(); }}
  function fbCloseFn() {{ fbPanel.classList.remove('open'); }}
  fbFab.addEventListener('click', () => {{
    if (fbPanel.classList.contains('open')) fbCloseFn(); else fbOpen();
  }});
  fbClose.addEventListener('click', fbCloseFn);

  function fbEsc(s) {{ return (s||'').replace(/[&<>"']/g, c => ({{'&':'&amp;','<':'&lt;','>':'&gt;','"':'&quot;',"'":'&#39;'}})[c]); }}
  function fbRelTime(iso) {{
    if (!iso) return '';
    const s = Math.max(0, Math.floor((Date.now() - new Date(iso).getTime())/1000));
    if (s < 60) return s + 's ago';
    if (s < 3600) return Math.floor(s/60) + 'm ago';
    if (s < 86400) return Math.floor(s/3600) + 'h ago';
    return Math.floor(s/86400) + 'd ago';
  }}

  async function fbLoad() {{
    try {{
      const r = await fetch('/api/docs/feedback?doc=' + encodeURIComponent(FB_DOC), {{credentials:'include'}});
      if (!r.ok) return;
      const list = await r.json();
      fbStream.innerHTML = '';
      if (!list.length) {{
        fbStream.appendChild(fbEmpty); fbEmpty.style.display='block';
        fbCount.style.display='none';
        fbSubCount.textContent = '0 comments';
        return;
      }}
      fbEmpty.style.display='none';
      fbCount.style.display = ''; fbCount.textContent = list.length;
      fbSubCount.textContent = list.length + (list.length === 1 ? ' comment' : ' comments');
      // Newest at bottom (chat-style)
      list.slice().reverse().forEach(fb => {{
        const m = document.createElement('div');
        m.className = 'fb-msg ' + (fb.status === 'addressed' ? 'addressed' : fb.status === 'closed' ? 'closed' : '');
        const anchor = fb.section_anchor ? '<div class="fb-section">' + fbEsc(fb.section_anchor) + '</div>' : '';
        m.innerHTML = anchor + fbEsc(fb.comment_text)
          + '<div class="fb-meta">'
          + '<span>' + fbEsc(fb.author || 'reviewer') + ' &middot; ' + fbRelTime(fb.created_at) + ' &middot; <em>' + fbEsc(fb.status || 'open') + '</em></span>'
          + '<span class="fb-actions">'
          + (fb.status !== 'addressed' ? '<button class="fb-tiny" data-id="' + fb.id + '" data-action="addressed">Mark addressed</button>' : '')
          + (fb.status !== 'closed'    ? '<button class="fb-tiny" data-id="' + fb.id + '" data-action="closed">Close</button>' : '')
          + '</span></div>';
        fbStream.appendChild(m);
      }});
      fbStream.querySelectorAll('button.fb-tiny').forEach(b => {{
        b.addEventListener('click', async (e) => {{
          const id = e.currentTarget.getAttribute('data-id');
          const action = e.currentTarget.getAttribute('data-action');
          await fetch('/api/docs/feedback/' + id, {{
            method:'PATCH', credentials:'include',
            headers:{{'Content-Type':'application/json'}},
            body: JSON.stringify({{status: action}})
          }});
          fbLoad();
        }});
      }});
      fbStream.scrollTop = fbStream.scrollHeight;
    }} catch (e) {{ console.error('feedback load', e); }}
  }}

  async function fbSubmit(ev) {{
    ev.preventDefault();
    const text = fbText.value.trim();
    if (!text) return false;
    fbSend.disabled = true;
    try {{
      const r = await fetch('/api/docs/feedback', {{
        method:'POST', credentials:'include',
        headers:{{'Content-Type':'application/json'}},
        body: JSON.stringify({{
          doc_slug: FB_DOC,
          section_anchor: fbAnchor.value.trim() || null,
          comment_text: text
        }})
      }});
      if (r.ok) {{
        fbText.value = ''; fbAnchor.value = '';
        await fbLoad();
      }} else {{
        alert('Could not save feedback (' + r.status + ')');
      }}
    }} finally {{ fbSend.disabled = false; }}
    return false;
  }}

  fbLoad();
  setInterval(fbLoad, 15000);
</script>
<!-- === v1.1 DOC-FEEDBACK END === -->

</body>
</html>"""


@router.get("/rfp-reply/cost-calculator", response_class=HTMLResponse)
def rfp_reply_cost_calculator() -> str:
    """Interactive cost calculator for the Keysight RFP response. Every input
    is editable; cost figures recompute on change. Anchored to the 880k email
    annual baseline with attachments, OCR pages, LLM token mix, and unit
    prices all under the operator's control."""
    from ._cost_calculator_html import COST_CALCULATOR_HTML
    return COST_CALCULATOR_HTML


@router.get("/rfp-reply/benefit-case", response_class=HTMLResponse)
def rfp_reply_benefit_case() -> str:
    """Interactive benefit-case calculator. Pulls the platform run cost from
    the cost calculator default and combines it with the operational baseline
    (650 FTE x $15k, 65% efficiency target) to produce a three-year NPV,
    ROI %, payback months, and FTE redeployment count."""
    from ._benefit_calculator_html import BENEFIT_CALCULATOR_HTML
    return BENEFIT_CALCULATOR_HTML


@router.get("/rfp-reply/benefit-case-simple", response_class=HTMLResponse)
def rfp_reply_benefit_case_simple() -> str:
    """Simplified, customer-facing benefit calculator. Six inputs, four
    headline tiles (payback, annual savings, 3-year net, ROI), and a
    side-by-side ZBrain vs systems-integrator comparison. Same math as the
    full calculator but with token-mix, OCR, translation, and ramp inputs
    hidden so the page reads in 30 seconds, not 5 minutes."""
    from ._benefit_calculator_simple_html import BENEFIT_CALCULATOR_SIMPLE_HTML
    return BENEFIT_CALCULATOR_SIMPLE_HTML


@router.get("/rfp-reply/value-capture", response_class=HTMLResponse)
def rfp_reply_value_capture() -> str:
    """Customer-perspective value-capture calculator. Four inputs, a
    22-week timeline visualisation, and a head-to-head with a typical
    systems integrator. Lead story: ZBrain delivers in 8 weeks so FTE
    savings start week 9; an SI delivers in 6 months so the customer
    keeps paying full FTE cost through year-end. The 16-week head start
    is the real economic argument."""
    from ._benefit_calculator_value_html import BENEFIT_CALCULATOR_VALUE_HTML
    return BENEFIT_CALCULATOR_VALUE_HTML


@router.get("/rfp-reply/{slug}", response_class=HTMLResponse)
def rfp_reply_section_html(slug: str) -> str:
    section = get_section(slug)
    if section is None:
        raise HTTPException(404, f"Unknown section: {slug}")
    blob = section["builder"]()
    # Convert DOCX bytes to HTML using mammoth. Same content as the .docx
    # download, just rendered for browser viewing.
    result = _mammoth_to_html(blob)
    return _rfp_viewer_shell(slug, section["label"], result)


def _mammoth_to_html(docx_bytes: bytes) -> str:
    """Convert DOCX bytes to HTML via mammoth, with style mappings that
    preserve our heading hierarchy and table semantics."""
    if mammoth is None:
        # Graceful fallback for environments where mammoth isn't available.
        # Keeps the backend bootable and still renders readable section text.
        from docx import Document as _DocxDocument

        doc = _DocxDocument(io.BytesIO(docx_bytes))
        out: list[str] = []
        for p in doc.paragraphs:
            txt = (p.text or "").strip()
            if not txt:
                continue
            style_name = ((p.style.name if p.style else "") or "").lower()
            tag = "p"
            if style_name.startswith("heading 1"):
                tag = "h1"
            elif style_name.startswith("heading 2"):
                tag = "h2"
            elif style_name.startswith("heading 3"):
                tag = "h3"
            out.append(f"<{tag}>{_html_escape(txt)}</{tag}>")
        if doc.tables:
            for t in doc.tables:
                rows_html: list[str] = []
                for r in t.rows:
                    cells = "".join(f"<td>{_html_escape((c.text or '').strip())}</td>" for c in r.cells)
                    rows_html.append(f"<tr>{cells}</tr>")
                out.append("<table>" + "".join(rows_html) + "</table>")
        if not out:
            return "<p>(No content)</p>"
        return "\n".join(out)

    style_map = """
    p[style-name='Heading 1'] => h1:fresh
    p[style-name='Heading 2'] => h2:fresh
    p[style-name='Heading 3'] => h3:fresh
    p[style-name='List Bullet'] => ul > li:fresh
    """
    result = mammoth.convert_to_html(io.BytesIO(docx_bytes), style_map=style_map)
    return result.value
# === v1.1 RFP-VIEWER END ===


@router.get("/rfp-reply", response_class=HTMLResponse)
def rfp_reply_landing() -> str:
    cards = []
    cards.append(
        '<div class="row" style="border-color:#1A55F9;background:#F4F7FD;">'
        '<div class="rowmeta">'
        '<a class="rowlabel" href="/api/docs/rfp-reply/benefit-case">Benefit case calculator (interactive)</a>'
        '<code>End-to-end · FTE redeployment · 3-year NPV · ROI · payback</code>'
        '</div>'
        '<div class="rowactions">'
        '<a class="btn primary" href="/api/docs/rfp-reply/benefit-case">Open benefit case</a>'
        '</div>'
        '</div>'
    )
    cards.append(
        '<div class="row" style="border-color:#1A55F9;background:#F4F7FD;">'
        '<div class="rowmeta">'
        '<a class="rowlabel" href="/api/docs/rfp-reply/cost-calculator">Cost calculator (interactive)</a>'
        '<code>Live recompute · editable inputs · anchored to 880k baseline</code>'
        '</div>'
        '<div class="rowactions">'
        '<a class="btn primary" href="/api/docs/rfp-reply/cost-calculator">Open calculator</a>'
        '</div>'
        '</div>'
    )
    for s in RFP_SECTIONS:
        cards.append(
            f'<div class="row">'
            f'<div class="rowmeta">'
            f'<a class="rowlabel" href="/api/docs/rfp-reply/{s["slug"]}">{s["label"]}</a>'
            f'<code>{s["filename"]}</code>'
            f'</div>'
            f'<div class="rowactions">'
            f'<a class="btn primary" href="/api/docs/rfp-reply/{s["slug"]}">View</a>'
            f'<a class="btn" href="/api/docs/rfp-reply/{s["slug"]}.docx">Download .docx</a>'
            f'</div>'
            f'</div>'
        )
    cards_html = "\n".join(cards)
    return f"""<!doctype html>
<html lang=en>
<head>
<meta charset=utf-8>
<title>Keysight RFP Response &middot; section index</title>
<meta name=viewport content="width=device-width, initial-scale=1">
<style>
  :root {{ --ink: #131426; --rule: #e5e7eb; --accent: #1A55F9; --muted: #6B7280; }}
  body {{ font-family: -apple-system, BlinkMacSystemFont, Inter, "Segoe UI", Roboto, sans-serif;
         max-width: 920px; margin: 56px auto; padding: 0 24px 80px;
         color: var(--ink); line-height: 1.55; }}
  h1 {{ color: var(--accent); border-bottom: 2px solid var(--accent);
       padding-bottom: .35em; margin: 0 0 .4em; }}
  .intro {{ color: #374151; }}
  .row {{
    display: flex; align-items: center; justify-content: space-between; gap: 16px;
    padding: 14px 18px; border: 1px solid var(--rule); border-radius: 10px;
    margin: 8px 0; background: white; transition: all .15s;
  }}
  .row:hover {{ box-shadow: 0 1px 3px rgba(0,0,0,.06); border-color: #d1d5db; }}
  .rowmeta {{ flex: 1 1 auto; min-width: 0; }}
  .rowlabel {{ font-weight: 600; font-size: 15px; color: var(--ink);
              text-decoration: none; display: block; }}
  .rowlabel:hover {{ color: var(--accent); }}
  .rowmeta code {{ display: block; color: var(--muted); font-size: 11.5px;
                  margin-top: 2px; word-break: break-all;
                  font-family: ui-monospace, SFMono-Regular, monospace; }}
  .rowactions {{ display: flex; gap: 8px; flex: 0 0 auto; }}
  .btn {{
    display: inline-flex; align-items: center; gap: 5px;
    border: 1px solid var(--rule); background: white; color: var(--ink);
    padding: 8px 14px; border-radius: 8px; font-weight: 500; font-size: 13px;
    text-decoration: none; transition: all .15s;
  }}
  .btn.primary {{ background: var(--accent); color: white; border-color: var(--accent); }}
  .btn:hover {{ transform: translateY(-1px); box-shadow: 0 2px 6px rgba(0,0,0,.08); }}
  .btn.primary:hover {{ box-shadow: 0 3px 10px rgba(26,85,249,.25); }}
  .muted {{ color: var(--muted); font-size: 13px; }}
  .footer {{ margin-top: 28px; padding-top: 16px; border-top: 1px solid var(--rule); }}
  a {{ color: var(--accent); }}
  .legend {{ background: #f8fafc; border: 1px solid var(--rule);
            border-radius: 10px; padding: 14px 18px; margin: 12px 0 22px;
            font-size: 13px; color: #374151; }}
</style>
</head>
<body>
<h1>Keysight RFP Response</h1>
<p class="intro">The response is broken into twelve separate documents so each section has the space it needs. Click <strong>View</strong> on any row to read in the browser; click <strong>Download .docx</strong> to grab the editable Word file. Each download is generated on demand from the latest content.</p>

<div class="legend">
<strong>View</strong> opens a reading view with a sidebar that lets you flip between sections, plus a <strong>Save as PDF</strong> button that uses your browser's print dialog. The DOCX is the same content; pick whichever is easier for your reviewer.
</div>

{cards_html}

<div class="footer">
<p class="muted">
Companion documents: <a href="/api/docs/as-is">AS-IS Process SOP (rendered)</a> &middot; <a href="/api/docs/as-is.md">AS-IS SOP (Markdown)</a>
</p>
</div>
</body>
</html>"""
# === v1.1 RFP-REPLY END ===


# === v1.1 DOC-FEEDBACK START === Reviewer feedback / chat for the RFP-response viewer
class FeedbackIn(BaseModel):
    doc_slug: str
    section_anchor: Optional[str] = None
    comment_text: str
    author: Optional[str] = None


class FeedbackPatch(BaseModel):
    status: Optional[str] = None


def _serialize(fb: DocFeedback) -> dict:
    return {
        "id": fb.id,
        "doc_slug": fb.doc_slug,
        "section_anchor": fb.section_anchor,
        "comment_text": fb.comment_text,
        "author": fb.author or "reviewer",
        "status": fb.status or "open",
        "created_at": fb.created_at.isoformat() if fb.created_at else None,
        "updated_at": fb.updated_at.isoformat() if fb.updated_at else None,
    }


def _basic_auth_user(req: Request) -> str:
    """Pull the HTTP Basic Auth user as the comment author. Falls back to 'reviewer'."""
    import base64
    h = req.headers.get("authorization", "")
    if h.lower().startswith("basic "):
        try:
            raw = base64.b64decode(h.split(" ", 1)[1]).decode("utf-8", errors="ignore")
            user = raw.split(":", 1)[0]
            return user or "reviewer"
        except Exception:
            return "reviewer"
    return "reviewer"


@router.post("/feedback")
def feedback_add(body: FeedbackIn, req: Request, db: Session = Depends(get_db)) -> dict:
    text = (body.comment_text or "").strip()
    if not text:
        raise HTTPException(400, "comment_text is required")
    fb = DocFeedback(
        doc_slug=(body.doc_slug or "all").strip(),
        section_anchor=(body.section_anchor or None),
        comment_text=text,
        author=body.author or _basic_auth_user(req),
        status="open",
    )
    db.add(fb)
    db.commit()
    db.refresh(fb)
    return _serialize(fb)


@router.get("/feedback")
def feedback_list(doc: Optional[str] = None, status: Optional[str] = None, db: Session = Depends(get_db)) -> list[dict]:
    q = db.query(DocFeedback)
    if doc:
        q = q.filter(DocFeedback.doc_slug == doc)
    if status:
        q = q.filter(DocFeedback.status == status)
    rows = q.order_by(DocFeedback.created_at.desc()).all()
    return [_serialize(r) for r in rows]


@router.patch("/feedback/{fb_id}")
def feedback_patch(fb_id: int, body: FeedbackPatch, db: Session = Depends(get_db)) -> dict:
    fb = db.query(DocFeedback).filter(DocFeedback.id == fb_id).first()
    if not fb:
        raise HTTPException(404, f"feedback {fb_id} not found")
    if body.status is not None:
        fb.status = body.status
    fb.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(fb)
    return _serialize(fb)


@router.delete("/feedback/{fb_id}")
def feedback_delete(fb_id: int, db: Session = Depends(get_db)) -> dict:
    fb = db.query(DocFeedback).filter(DocFeedback.id == fb_id).first()
    if not fb:
        raise HTTPException(404, f"feedback {fb_id} not found")
    db.delete(fb)
    db.commit()
    return {"ok": True}
# === v1.1 DOC-FEEDBACK END ===
