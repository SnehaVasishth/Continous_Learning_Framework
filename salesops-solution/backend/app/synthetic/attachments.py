"""Build synthetic attachment files: PO PDFs, BOM XLSXs, scanned-PO PNGs, spec DOCX.

Enterprise-grade layout — descriptions wrap inside table cells, two-column
header (vendor/issuer + Bill-To/Ship-To metadata), subtotal/tax/total rows,
restrained color palette, no fictional company branding.
"""
from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.shared import Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfgen.canvas import Canvas
from reportlab.platypus import (
    KeepTogether,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

INK = colors.HexColor("#131426")
MUTED = colors.HexColor("#5B6275")
DIVIDER = colors.HexColor("#D9DCE4")
TINT = colors.HexColor("#F1F4F9")


def _po_styles():
    base = ParagraphStyle("base", fontName="Helvetica", fontSize=9, leading=12, textColor=INK)
    return {
        "title": ParagraphStyle("title", parent=base, fontName="Helvetica-Bold", fontSize=18, leading=22, spaceAfter=2),
        "kicker": ParagraphStyle("kicker", parent=base, fontSize=8, textColor=MUTED, leading=10),
        "label": ParagraphStyle("label", parent=base, fontSize=7.5, textColor=MUTED, leading=9, spaceAfter=2),
        "value": ParagraphStyle("value", parent=base, fontName="Helvetica", fontSize=9.5, leading=12),
        "value_b": ParagraphStyle("value_b", parent=base, fontName="Helvetica-Bold", fontSize=9.5, leading=12),
        "small": ParagraphStyle("small", parent=base, fontSize=8.5, leading=11),
        "small_b": ParagraphStyle("small_b", parent=base, fontName="Helvetica-Bold", fontSize=8.5, leading=11),
        "right": ParagraphStyle("right", parent=base, fontSize=9, leading=12, alignment=TA_RIGHT),
        "right_b": ParagraphStyle("right_b", parent=base, fontName="Helvetica-Bold", fontSize=9, leading=12, alignment=TA_RIGHT),
        "footnote": ParagraphStyle("footnote", parent=base, fontSize=7.5, textColor=MUTED, leading=10),
    }


def _font(size: int = 18):
    candidates = [
        r"C:\Windows\Fonts\segoeui.ttf",
        r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\calibri.ttf",
    ]
    for p in candidates:
        if Path(p).exists():
            try:
                return ImageFont.truetype(p, size)
            except Exception:
                continue
    return ImageFont.load_default()


def make_po_pdf(
    out_path: Path,
    *,
    customer_name: str,
    customer_addr: str,
    po_number: str,
    issue_date: date,
    line_items: Iterable[dict],
    payment_terms: str = "Net 45",
    requested_ship: str | None = None,
    bill_to: str | None = None,
    ship_to: str | None = None,
    note: str | None = None,
    quote_reference: str | None = None,
    buyer_contact: str | None = None,
) -> Path:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=LETTER,
        leftMargin=0.55 * inch,
        rightMargin=0.55 * inch,
        topMargin=0.55 * inch,
        bottomMargin=0.55 * inch,
        title=f"Purchase Order {po_number}",
        author=customer_name,
    )
    s = _po_styles()
    story = []

    header_left = [
        Paragraph("PURCHASE ORDER", s["title"]),
        Paragraph(customer_name.upper(), s["kicker"]),
        Spacer(1, 4),
        Paragraph(customer_addr, s["small"]),
    ]
    if buyer_contact:
        header_left.append(Paragraph(buyer_contact, s["small"]))

    meta_rows = [
        [Paragraph("PO Number", s["label"]), Paragraph(po_number, s["value_b"])],
        [Paragraph("Issue Date", s["label"]), Paragraph(issue_date.isoformat(), s["value"])],
    ]
    if requested_ship:
        meta_rows.append([Paragraph("Requested Ship Date", s["label"]), Paragraph(requested_ship, s["value"])])
    meta_rows.append([Paragraph("Payment Terms", s["label"]), Paragraph(payment_terms, s["value"])])
    if quote_reference:
        meta_rows.append([Paragraph("Quote Reference", s["label"]), Paragraph(quote_reference, s["value_b"])])

    meta_table = Table(meta_rows, colWidths=[1.3 * inch, 1.85 * inch])
    meta_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), TINT),
                ("BOX", (0, 0), (-1, -1), 0.4, DIVIDER),
                ("INNERGRID", (0, 0), (-1, -1), 0.25, DIVIDER),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]
        )
    )

    header_table = Table([[header_left, meta_table]], colWidths=[4.05 * inch, 3.20 * inch])
    header_table.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ]
        )
    )
    story.append(header_table)
    story.append(Spacer(1, 0.18 * inch))

    bill_block = [Paragraph("BILL TO", s["label"]), Paragraph(bill_to or customer_name, s["small"])]
    ship_block = [Paragraph("SHIP TO", s["label"]), Paragraph(ship_to or customer_name, s["small"])]
    addr_table = Table([[bill_block, ship_block]], colWidths=[3.625 * inch, 3.625 * inch])
    addr_table.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ("BOX", (0, 0), (-1, -1), 0.4, DIVIDER),
                ("LINEAFTER", (0, 0), (0, -1), 0.4, DIVIDER),
            ]
        )
    )
    story.append(addr_table)
    story.append(Spacer(1, 0.20 * inch))

    rows = [
        [
            Paragraph("SKU", s["small_b"]),
            Paragraph("Description", s["small_b"]),
            Paragraph("Qty", s["right_b"]),
            Paragraph("Unit Price", s["right_b"]),
            Paragraph("Extended", s["right_b"]),
        ]
    ]
    subtotal = 0.0
    for li in line_items:
        ext = li["qty"] * li["unit_price"]
        subtotal += ext
        rows.append(
            [
                Paragraph(li["sku"], s["small"]),
                Paragraph(li["description"], s["small"]),
                Paragraph(str(li["qty"]), s["right"]),
                Paragraph(f"${li['unit_price']:,.2f}", s["right"]),
                Paragraph(f"${ext:,.2f}", s["right"]),
            ]
        )

    tax_rate = 0.0
    tax = round(subtotal * tax_rate, 2)
    total = subtotal + tax

    spacer_row = ["", "", "", "", ""]
    rows.append(spacer_row)
    rows.append(["", "", "", Paragraph("Subtotal", s["right"]), Paragraph(f"${subtotal:,.2f}", s["right"])])
    rows.append(
        ["", "", "", Paragraph(f"Tax ({tax_rate * 100:.0f}%)", s["right"]), Paragraph(f"${tax:,.2f}", s["right"])]
    )
    rows.append(["", "", "", Paragraph("TOTAL", s["right_b"]), Paragraph(f"${total:,.2f}", s["right_b"])])

    line_count = len(rows) - 4
    table = Table(
        rows,
        colWidths=[0.95 * inch, 3.05 * inch, 0.55 * inch, 1.30 * inch, 1.40 * inch],
        repeatRows=1,
    )
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), TINT),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("LINEBELOW", (0, 0), (-1, 0), 0.6, INK),
                ("LINEBELOW", (0, 1), (-1, line_count), 0.25, DIVIDER),
                ("LINEABOVE", (3, line_count + 2), (-1, line_count + 2), 0.4, DIVIDER),
                ("LINEABOVE", (3, line_count + 4), (-1, line_count + 4), 0.6, INK),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 0.18 * inch))

    if note:
        story.append(Paragraph("NOTES", s["label"]))
        story.append(Paragraph(note, s["small"]))
        story.append(Spacer(1, 0.12 * inch))

    story.append(Spacer(1, 0.10 * inch))
    story.append(
        Paragraph(
            "This purchase order is issued subject to the buyer's standard terms and conditions of purchase. "
            "Acknowledgment of receipt is requested within 2 business days. "
            "All shipments must reference the PO Number on packing slips and invoices.",
            s["footnote"],
        )
    )

    doc.build(story)
    return out_path


def make_soa_pdf(
    out_path: Path,
    *,
    po_number: str,
    soa_number: str | None = None,
    acknowledged_date: date,
    customer_name: str,
    customer_addr: str | None = None,
    bill_to: str | None = None,
    ship_to: str | None = None,
    line_items: Iterable[dict],
    payment_terms: str | None = None,
    requested_ship_date: str | None = None,
    promised_ship_date: str | None = None,
    incoterms: str | None = None,
    sales_rep: str | None = None,
    csr_owner: str | None = None,
    notes: str | None = None,
    currency: str = "USD",
) -> Path:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=LETTER,
        leftMargin=0.55 * inch,
        rightMargin=0.55 * inch,
        topMargin=0.55 * inch,
        bottomMargin=0.55 * inch,
        title=f"Sales Order Acknowledgment {soa_number or po_number}",
        author="Keysight Technologies",
    )
    s = _po_styles()
    story = []

    issuer_block = [
        Paragraph("SALES ORDER ACKNOWLEDGMENT", s["title"]),
        Paragraph("KEYSIGHT TECHNOLOGIES, INC.", s["kicker"]),
        Spacer(1, 4),
        Paragraph(
            "1400 Fountaingrove Parkway<br/>Santa Rosa, CA 95403, USA<br/>Tax ID: 26-2818193",
            s["small"],
        ),
    ]

    meta_rows = [
        [Paragraph("PO Number", s["label"]), Paragraph(po_number, s["value_b"])],
    ]
    if soa_number:
        meta_rows.append([Paragraph("SOA Number", s["label"]), Paragraph(soa_number, s["value_b"])])
    meta_rows.append(
        [Paragraph("Acknowledged", s["label"]), Paragraph(acknowledged_date.isoformat(), s["value"])]
    )
    if requested_ship_date:
        meta_rows.append([Paragraph("Requested Ship", s["label"]), Paragraph(requested_ship_date, s["value"])])
    if promised_ship_date:
        meta_rows.append([Paragraph("Promised Ship", s["label"]), Paragraph(promised_ship_date, s["value_b"])])
    if payment_terms:
        meta_rows.append([Paragraph("Payment Terms", s["label"]), Paragraph(payment_terms, s["value"])])
    if incoterms:
        meta_rows.append([Paragraph("Incoterms", s["label"]), Paragraph(incoterms, s["value"])])
    if sales_rep:
        meta_rows.append([Paragraph("Sales Rep", s["label"]), Paragraph(sales_rep, s["value"])])

    meta_table = Table(meta_rows, colWidths=[1.3 * inch, 1.85 * inch])
    meta_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), TINT),
                ("BOX", (0, 0), (-1, -1), 0.4, DIVIDER),
                ("INNERGRID", (0, 0), (-1, -1), 0.25, DIVIDER),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]
        )
    )

    header_table = Table([[issuer_block, meta_table]], colWidths=[4.05 * inch, 3.20 * inch])
    header_table.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ]
        )
    )
    story.append(header_table)
    story.append(Spacer(1, 0.18 * inch))

    cust_block = [
        Paragraph("CUSTOMER", s["label"]),
        Paragraph(customer_name, s["value_b"]),
    ]
    if customer_addr:
        cust_block.append(Paragraph(customer_addr, s["small"]))
    if csr_owner:
        cust_block.append(Spacer(1, 4))
        cust_block.append(Paragraph(f"<b>CSR:</b> {csr_owner}", s["small"]))

    if bill_to or ship_to:
        bill_inner = [Paragraph("BILL TO", s["label"]), Paragraph(bill_to or customer_name, s["small"])]
        ship_inner = [Paragraph("SHIP TO", s["label"]), Paragraph(ship_to or customer_name, s["small"])]
        addr_table = Table(
            [[cust_block, bill_inner, ship_inner]],
            colWidths=[2.45 * inch, 2.4 * inch, 2.4 * inch],
        )
        addr_table.setStyle(
            TableStyle(
                [
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 10),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                    ("TOPPADDING", (0, 0), (-1, -1), 8),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                    ("BOX", (0, 0), (-1, -1), 0.4, DIVIDER),
                    ("LINEAFTER", (0, 0), (1, -1), 0.4, DIVIDER),
                ]
            )
        )
        story.append(addr_table)
    else:
        cust_table = Table([[cust_block]], colWidths=[7.25 * inch])
        cust_table.setStyle(
            TableStyle(
                [
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 10),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                    ("TOPPADDING", (0, 0), (-1, -1), 8),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                    ("BOX", (0, 0), (-1, -1), 0.4, DIVIDER),
                ]
            )
        )
        story.append(cust_table)
    story.append(Spacer(1, 0.20 * inch))

    rows = [
        [
            Paragraph("SKU", s["small_b"]),
            Paragraph("Description", s["small_b"]),
            Paragraph("Qty", s["right_b"]),
            Paragraph("Unit Price", s["right_b"]),
            Paragraph("Extended", s["right_b"]),
        ]
    ]
    subtotal = 0.0
    for li in line_items:
        try:
            qty = int(li.get("qty") or 0)
            price = float(li.get("unit_price") or 0.0)
        except Exception:
            qty, price = 0, 0.0
        ext = qty * price
        subtotal += ext
        rows.append(
            [
                Paragraph(str(li.get("sku") or ""), s["small"]),
                Paragraph(str(li.get("description") or ""), s["small"]),
                Paragraph(str(qty), s["right"]),
                Paragraph(f"${price:,.2f}", s["right"]),
                Paragraph(f"${ext:,.2f}", s["right"]),
            ]
        )

    tax_rate = 0.0
    tax = round(subtotal * tax_rate, 2)
    total = subtotal + tax
    line_count = len(rows) - 1

    rows.append(["", "", "", "", ""])
    rows.append(["", "", "", Paragraph("Subtotal", s["right"]), Paragraph(f"${subtotal:,.2f}", s["right"])])
    rows.append(
        ["", "", "", Paragraph(f"Tax ({tax_rate * 100:.0f}%)", s["right"]), Paragraph(f"${tax:,.2f}", s["right"])]
    )
    rows.append(["", "", "", Paragraph(f"TOTAL ({currency})", s["right_b"]), Paragraph(f"${total:,.2f}", s["right_b"])])

    table = Table(
        rows,
        colWidths=[0.95 * inch, 3.05 * inch, 0.55 * inch, 1.30 * inch, 1.40 * inch],
        repeatRows=1,
    )
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), TINT),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("LINEBELOW", (0, 0), (-1, 0), 0.6, INK),
                ("LINEBELOW", (0, 1), (-1, line_count), 0.25, DIVIDER),
                ("LINEABOVE", (3, line_count + 2), (-1, line_count + 2), 0.4, DIVIDER),
                ("LINEABOVE", (3, line_count + 4), (-1, line_count + 4), 0.6, INK),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 0.18 * inch))

    if notes:
        story.append(Paragraph("NOTES", s["label"]))
        story.append(Paragraph(notes, s["small"]))
        story.append(Spacer(1, 0.12 * inch))

    story.append(Spacer(1, 0.10 * inch))
    story.append(
        Paragraph(
            "This Sales Order Acknowledgment confirms receipt of your purchase order and is governed by Keysight's "
            "standard terms and conditions of sale. The promised ship date is an estimate and may be revised by mutual agreement. "
            "Please reference the PO Number on all correspondence and packing materials.",
            s["footnote"],
        )
    )

    doc.build(story)
    return out_path


def make_bom_xlsx(
    out_path: Path,
    *,
    customer_name: str,
    quote_number: str,
    line_items: Iterable[dict],
) -> Path:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    ws = wb.active
    ws.title = "BOM"

    ink = Font(color="131426", bold=True, size=12)
    label = Font(color="5B6275", bold=False, size=9)
    header = Font(color="FFFFFF", bold=True, size=10)
    money = Font(size=10, name="Calibri")
    fill_blue = PatternFill(start_color="1A55F9", end_color="1A55F9", fill_type="solid")
    fill_tint = PatternFill(start_color="F1F4F9", end_color="F1F4F9", fill_type="solid")
    thin = Side(style="thin", color="D9DCE4")
    border = Border(top=thin, left=thin, right=thin, bottom=thin)
    right = Alignment(horizontal="right", vertical="center")
    left = Alignment(horizontal="left", vertical="center", wrap_text=True)

    ws["A1"] = "Bill of Materials"
    ws["A1"].font = ink
    ws.row_dimensions[1].height = 22

    ws["A3"] = "Customer"
    ws["A3"].font = label
    ws["B3"] = customer_name
    ws["A4"] = "Quote #"
    ws["A4"].font = label
    ws["B4"] = quote_number
    ws["B4"].font = Font(bold=True, size=10)

    headers = ["SKU", "Description", "Qty", "Unit Price (USD)", "Extended (USD)"]
    for col_idx, h in enumerate(headers, start=1):
        c = ws.cell(row=6, column=col_idx, value=h)
        c.font = header
        c.fill = fill_blue
        c.alignment = Alignment(horizontal="center", vertical="center")
        c.border = border

    subtotal = 0.0
    row_idx = 7
    for li in line_items:
        ext = li["qty"] * li["unit_price"]
        subtotal += ext
        ws.cell(row=row_idx, column=1, value=li["sku"]).alignment = left
        ws.cell(row=row_idx, column=2, value=li["description"]).alignment = left
        ws.cell(row=row_idx, column=3, value=li["qty"]).alignment = right
        c4 = ws.cell(row=row_idx, column=4, value=li["unit_price"])
        c4.alignment = right
        c4.number_format = "$#,##0.00"
        c5 = ws.cell(row=row_idx, column=5, value=ext)
        c5.alignment = right
        c5.number_format = "$#,##0.00"
        for col in range(1, 6):
            ws.cell(row=row_idx, column=col).border = border
            ws.cell(row=row_idx, column=col).font = money
        row_idx += 1

    total_row = row_idx + 1
    label_cell = ws.cell(row=total_row, column=4, value="Subtotal")
    label_cell.font = Font(bold=True, size=10)
    label_cell.alignment = right
    label_cell.fill = fill_tint
    sub_cell = ws.cell(row=total_row, column=5, value=subtotal)
    sub_cell.number_format = "$#,##0.00"
    sub_cell.alignment = right
    sub_cell.font = Font(bold=True, size=10)
    sub_cell.fill = fill_tint

    grand_row = total_row + 1
    g_label = ws.cell(row=grand_row, column=4, value="TOTAL")
    g_label.font = Font(bold=True, size=11, color="131426")
    g_label.alignment = right
    g = ws.cell(row=grand_row, column=5, value=subtotal)
    g.number_format = "$#,##0.00"
    g.font = Font(bold=True, size=11, color="131426")
    g.alignment = right

    widths = {"A": 18, "B": 52, "C": 8, "D": 18, "E": 18}
    for col, w in widths.items():
        ws.column_dimensions[col].width = w

    wb.save(out_path)
    return out_path


def make_spec_docx(
    out_path: Path,
    *,
    title: str,
    sections: list[tuple[str, str]],
) -> Path:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    h = doc.add_heading(title, level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x13, 0x14, 0x26)
    for heading, body in sections:
        sub = doc.add_heading(heading, level=2)
        for run in sub.runs:
            run.font.color.rgb = RGBColor(0x13, 0x14, 0x26)
        p = doc.add_paragraph(body)
        for run in p.runs:
            run.font.size = Pt(10)
    doc.save(str(out_path))
    return out_path


def make_scanned_po_png(
    out_path: Path,
    *,
    customer_name: str,
    customer_addr: str,
    po_number: str,
    issue_date: str,
    line_items: Iterable[dict],
    payment_terms: str = "Net 45",
    requested_ship: str | None = None,
) -> Path:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    W, H = 1240, 1600
    img = Image.new("RGB", (W, H), (252, 252, 250))
    draw = ImageDraw.Draw(img)

    f_title = _font(38)
    f_h = _font(22)
    f_b = _font(16)
    f_label = _font(13)
    f_small = _font(13)

    margin = 80
    y = margin

    draw.text((margin, y), "PURCHASE ORDER", font=f_title, fill=(19, 20, 38))
    y += 50
    draw.text((margin, y), customer_name.upper(), font=f_label, fill=(91, 98, 117))
    y += 30
    draw.text((margin, y), customer_addr, font=f_b, fill=(19, 20, 38))
    y += 50

    box_x = W - margin - 360
    box_y = margin + 8
    box_w = 360
    rows = [
        ("PO NUMBER", po_number, True),
        ("ISSUE DATE", issue_date, False),
    ]
    if requested_ship:
        rows.append(("REQUESTED SHIP DATE", requested_ship, False))
    rows.append(("PAYMENT TERMS", payment_terms, False))

    box_h = 40 + 38 * len(rows)
    draw.rectangle([box_x, box_y, box_x + box_w, box_y + box_h], outline=(217, 220, 228), width=2, fill=(241, 244, 249))
    by = box_y + 16
    for label, val, bold in rows:
        draw.text((box_x + 14, by), label, font=f_label, fill=(91, 98, 117))
        draw.text((box_x + 14, by + 14), val, font=f_b if not bold else _font(17), fill=(19, 20, 38))
        by += 38

    y = max(y, box_y + box_h) + 30
    draw.line([margin, y, W - margin, y], fill=(217, 220, 228), width=2)
    y += 18

    col_x = [margin, margin + 170, margin + 720, margin + 820, margin + 1000]
    headers = ["SKU", "DESCRIPTION", "QTY", "UNIT PRICE", "EXTENDED"]
    for i, h_ in enumerate(headers):
        draw.text((col_x[i], y), h_, font=f_label, fill=(91, 98, 117))
    y += 22
    draw.line([margin, y, W - margin, y], fill=(19, 20, 38), width=1)
    y += 14

    subtotal = 0.0
    for li in line_items:
        ext = li["qty"] * li["unit_price"]
        subtotal += ext
        desc = li["description"]
        if len(desc) > 60:
            desc = desc[:57] + "..."
        draw.text((col_x[0], y), li["sku"], font=f_small, fill=(19, 20, 38))
        draw.text((col_x[1], y), desc, font=f_small, fill=(19, 20, 38))
        draw.text((col_x[2], y), str(li["qty"]), font=f_small, fill=(19, 20, 38))
        draw.text((col_x[3], y), f"${li['unit_price']:,.2f}", font=f_small, fill=(19, 20, 38))
        draw.text((col_x[4], y), f"${ext:,.2f}", font=f_small, fill=(19, 20, 38))
        y += 30
        draw.line([margin, y - 4, W - margin, y - 4], fill=(231, 232, 238), width=1)

    y += 14
    draw.line([col_x[3], y, W - margin, y], fill=(19, 20, 38), width=1)
    y += 14
    draw.text((col_x[3], y), "TOTAL", font=f_h, fill=(19, 20, 38))
    draw.text((col_x[4], y), f"${subtotal:,.2f}", font=f_h, fill=(19, 20, 38))

    for i, x in enumerate(range(0, W, 4)):
        if i % 7 == 0:
            draw.point((x, H - 20), fill=(180, 180, 180))
    draw.text((margin, H - 50), "scanned copy — please OCR", font=f_label, fill=(150, 150, 160))

    img.save(out_path, "PNG")
    return out_path


AMBER = colors.HexColor("#B45309")
AMBER_TINT = colors.HexColor("#FEF3C7")
ROSE = colors.HexColor("#B91C1C")
ROSE_TINT = colors.HexColor("#FEE2E2")
EMERALD = colors.HexColor("#047857")
EMERALD_TINT = colors.HexColor("#D1FAE5")

KEYSIGHT_FROM = (
    "Keysight Technologies, Inc.<br/>"
    "1400 Fountaingrove Pkwy, Santa Rosa, CA 95403, USA<br/>"
    "Tax ID: 26-2818193 &nbsp;·&nbsp; DUNS: 04-997-2148"
)
KEYSIGHT_REMIT = (
    "<b>Remit Payment To:</b><br/>"
    "Keysight Technologies, Inc. — Accounts Receivable<br/>"
    "P.O. Box 740393, Los Angeles, CA 90074-0393<br/>"
    "ACH/Wire: Bank of America · Routing 026009593 · Acct 1499-3-77104<br/>"
    "Reference invoice number on all remittances."
)


def _draw_status_watermark(canvas: Canvas, doc, *, label: str, color):
    canvas.saveState()
    canvas.setFont("Helvetica-Bold", 84)
    canvas.setFillColor(color)
    canvas.setFillAlpha(0.10)
    canvas.translate(LETTER[0] / 2, LETTER[1] / 2)
    canvas.rotate(28)
    canvas.drawCentredString(0, 0, label.upper())
    canvas.restoreState()


def _doc_styles():
    base = ParagraphStyle("base", fontName="Helvetica", fontSize=9, leading=12, textColor=INK)
    return {
        "base": base,
        "title": ParagraphStyle("title", parent=base, fontName="Helvetica-Bold", fontSize=20, leading=24),
        "subtitle": ParagraphStyle("subtitle", parent=base, fontSize=8.5, textColor=MUTED, leading=11),
        "kicker": ParagraphStyle("kicker", parent=base, fontSize=8, textColor=MUTED, leading=10),
        "label": ParagraphStyle("label", parent=base, fontSize=7.5, textColor=MUTED, leading=9, spaceAfter=2),
        "value": ParagraphStyle("value", parent=base, fontSize=9.5, leading=12),
        "value_b": ParagraphStyle("value_b", parent=base, fontName="Helvetica-Bold", fontSize=9.5, leading=12),
        "small": ParagraphStyle("small", parent=base, fontSize=8.5, leading=11),
        "small_b": ParagraphStyle("small_b", parent=base, fontName="Helvetica-Bold", fontSize=8.5, leading=11),
        "right": ParagraphStyle("right", parent=base, fontSize=9, leading=12, alignment=TA_RIGHT),
        "right_b": ParagraphStyle("right_b", parent=base, fontName="Helvetica-Bold", fontSize=9, leading=12, alignment=TA_RIGHT),
        "right_lg": ParagraphStyle("right_lg", parent=base, fontName="Helvetica-Bold", fontSize=11, leading=14, alignment=TA_RIGHT),
        "center": ParagraphStyle("center", parent=base, fontSize=9, leading=12, alignment=TA_CENTER),
        "section": ParagraphStyle("section", parent=base, fontName="Helvetica-Bold", fontSize=10, leading=12, textColor=INK, spaceAfter=4),
        "footnote": ParagraphStyle("footnote", parent=base, fontSize=7.5, textColor=MUTED, leading=10),
        "stamp": ParagraphStyle("stamp", parent=base, fontName="Helvetica-Bold", fontSize=9, leading=11, textColor=INK, alignment=TA_CENTER),
        "stamp_sm": ParagraphStyle("stamp_sm", parent=base, fontSize=7.5, leading=9, textColor=MUTED, alignment=TA_CENTER),
    }


def _meta_table(rows: list[tuple[str, str]], s, *, col_widths=(1.3 * inch, 1.85 * inch)):
    body = [[Paragraph(label, s["label"]), Paragraph(value, s["value_b"] if i == 0 else s["value"])] for i, (label, value) in enumerate(rows)]
    t = Table(body, colWidths=list(col_widths))
    t.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), TINT),
                ("BOX", (0, 0), (-1, -1), 0.4, DIVIDER),
                ("INNERGRID", (0, 0), (-1, -1), 0.25, DIVIDER),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]
        )
    )
    return t


def _two_box(left_label: str, left_body: str, right_label: str, right_body: str, s):
    left = [Paragraph(left_label, s["label"]), Paragraph(left_body, s["small"])]
    right = [Paragraph(right_label, s["label"]), Paragraph(right_body, s["small"])]
    t = Table([[left, right]], colWidths=[3.625 * inch, 3.625 * inch])
    t.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ("BOX", (0, 0), (-1, -1), 0.4, DIVIDER),
                ("LINEAFTER", (0, 0), (0, -1), 0.4, DIVIDER),
            ]
        )
    )
    return t


def make_invoice_pdf(
    out_path: Path,
    *,
    invoice_number: str,
    invoice_date,
    due_date,
    customer_name: str,
    customer_addr: str,
    bill_to: str | None,
    ship_to: str | None,
    line_items: Iterable[dict],
    subtotal: float,
    tax: float,
    total: float,
    currency: str = "USD",
    payment_terms: str = "Net 45",
    notes: str | None = None,
    status: str = "issued",
) -> Path:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=LETTER,
        leftMargin=0.55 * inch,
        rightMargin=0.55 * inch,
        topMargin=0.55 * inch,
        bottomMargin=0.55 * inch,
        title=f"Invoice {invoice_number}",
        author="Keysight Technologies, Inc.",
    )
    s = _doc_styles()
    story = []

    header_left = [
        Paragraph("Keysight Technologies, Inc.", s["value_b"]),
        Spacer(1, 2),
        Paragraph(KEYSIGHT_FROM, s["small"]),
    ]
    header_right = [
        Paragraph("INVOICE", s["title"]),
        Paragraph(f"Invoice {invoice_number}", s["subtitle"]),
    ]
    header_table = Table(
        [[header_left, header_right]],
        colWidths=[4.05 * inch, 3.20 * inch],
    )
    header_table.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (0, 0), "TOP"),
                ("VALIGN", (1, 0), (1, 0), "TOP"),
                ("ALIGN", (1, 0), (1, 0), "RIGHT"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ]
        )
    )
    story.append(header_table)
    story.append(Spacer(1, 0.18 * inch))

    inv_date_str = invoice_date.isoformat() if hasattr(invoice_date, "isoformat") else str(invoice_date)
    due_date_str = due_date.isoformat() if hasattr(due_date, "isoformat") else str(due_date)
    meta = _meta_table(
        [
            ("Invoice Number", invoice_number),
            ("Invoice Date", inv_date_str),
            ("Due Date", due_date_str),
            ("Payment Terms", payment_terms),
            ("Currency", currency),
            ("Status", status.upper()),
        ],
        s,
        col_widths=(1.6 * inch, 1.95 * inch),
    )
    customer_block = [
        Paragraph("CUSTOMER", s["label"]),
        Paragraph(customer_name, s["value_b"]),
        Spacer(1, 2),
        Paragraph(customer_addr, s["small"]),
    ]
    cust_meta = Table(
        [[customer_block, meta]],
        colWidths=[3.70 * inch, 3.55 * inch],
    )
    cust_meta.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ]
        )
    )
    story.append(cust_meta)
    story.append(Spacer(1, 0.16 * inch))

    story.append(_two_box("BILL TO", bill_to or customer_name, "SHIP TO", ship_to or customer_name, s))
    story.append(Spacer(1, 0.20 * inch))

    rows = [
        [
            Paragraph("SKU", s["small_b"]),
            Paragraph("Description", s["small_b"]),
            Paragraph("Qty", s["right_b"]),
            Paragraph("Unit Price", s["right_b"]),
            Paragraph("Extended", s["right_b"]),
        ]
    ]
    for li in line_items:
        ext = li["qty"] * li["unit_price"]
        rows.append(
            [
                Paragraph(li["sku"], s["small"]),
                Paragraph(li["description"], s["small"]),
                Paragraph(str(li["qty"]), s["right"]),
                Paragraph(f"${li['unit_price']:,.2f}", s["right"]),
                Paragraph(f"${ext:,.2f}", s["right"]),
            ]
        )

    line_count = len(rows) - 1
    rows.append(["", "", "", "", ""])
    rows.append(["", "", "", Paragraph("Subtotal", s["right"]), Paragraph(f"${subtotal:,.2f}", s["right"])])
    rows.append(["", "", "", Paragraph("Tax", s["right"]), Paragraph(f"${tax:,.2f}", s["right"])])
    rows.append(["", "", "", Paragraph("TOTAL", s["right_lg"]), Paragraph(f"${total:,.2f} {currency}", s["right_lg"])])

    table = Table(
        rows,
        colWidths=[0.95 * inch, 3.05 * inch, 0.55 * inch, 1.30 * inch, 1.40 * inch],
        repeatRows=1,
    )
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), TINT),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("LINEBELOW", (0, 0), (-1, 0), 0.6, INK),
                ("LINEBELOW", (0, 1), (-1, line_count), 0.25, DIVIDER),
                ("LINEABOVE", (3, line_count + 2), (-1, line_count + 2), 0.4, DIVIDER),
                ("LINEABOVE", (3, line_count + 4), (-1, line_count + 4), 0.6, INK),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 0.20 * inch))

    remit_block = Table(
        [[Paragraph(KEYSIGHT_REMIT, s["small"])]],
        colWidths=[7.25 * inch],
    )
    remit_block.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), TINT),
                ("BOX", (0, 0), (-1, -1), 0.4, DIVIDER),
                ("LEFTPADDING", (0, 0), (-1, -1), 12),
                ("RIGHTPADDING", (0, 0), (-1, -1), 12),
                ("TOPPADDING", (0, 0), (-1, -1), 10),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
            ]
        )
    )
    story.append(remit_block)

    if notes:
        story.append(Spacer(1, 0.14 * inch))
        story.append(Paragraph("NOTES", s["label"]))
        story.append(Paragraph(notes, s["small"]))

    story.append(Spacer(1, 0.16 * inch))
    story.append(
        Paragraph(
            "Payment is due by the date shown above. Late payments accrue interest at 1.5% per month or "
            "the maximum permitted by law, whichever is less. Disputed amounts must be reported in writing "
            "within 10 business days of receipt of this invoice. Goods sold subject to Keysight's standard "
            "Terms & Conditions of Sale (keysight.com/find/terms).",
            s["footnote"],
        )
    )

    on_page = None
    if status == "paid":
        on_page = lambda c, d: _draw_status_watermark(c, d, label="PAID", color=EMERALD)
    elif status == "overdue":
        on_page = lambda c, d: _draw_status_watermark(c, d, label="OVERDUE", color=ROSE)

    if on_page is not None:
        doc.build(story, onFirstPage=on_page, onLaterPages=on_page)
    else:
        doc.build(story)
    return out_path


def make_work_order_pdf(
    out_path: Path,
    *,
    wo_number: str,
    customer_name: str,
    asset_serial: str,
    asset_sku: str | None,
    type: str,
    scheduled_date,
    sla_target_date,
    technician: str | None,
    region: str,
    assigned_team: str | None,
    description: str | None,
    standards_referenced: list[str] | None,
    parts_used: list[dict] | None,
    labor_hours: float,
    cost_usd: float,
    signoff_status: str,
    root_cause: str | None,
    service_contract_id: str | None,
) -> Path:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=LETTER,
        leftMargin=0.55 * inch,
        rightMargin=0.55 * inch,
        topMargin=0.55 * inch,
        bottomMargin=0.55 * inch,
        title=f"Work Order {wo_number}",
        author="Keysight Technologies, Inc.",
    )
    s = _doc_styles()
    story = []

    header_left = [
        Paragraph("Keysight Technologies, Inc.", s["value_b"]),
        Spacer(1, 2),
        Paragraph("Field Service & Calibration Operations", s["small"]),
        Paragraph("1400 Fountaingrove Pkwy, Santa Rosa, CA 95403, USA", s["small"]),
    ]
    header_right = [
        Paragraph("KEYSIGHT WORK ORDER", s["title"]),
        Paragraph(f"WO {wo_number}", s["subtitle"]),
    ]
    header_table = Table(
        [[header_left, header_right]],
        colWidths=[4.05 * inch, 3.20 * inch],
    )
    header_table.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ALIGN", (1, 0), (1, 0), "RIGHT"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ]
        )
    )
    story.append(header_table)
    story.append(Spacer(1, 0.18 * inch))

    sched_str = scheduled_date.isoformat()[:10] if hasattr(scheduled_date, "isoformat") else (str(scheduled_date) if scheduled_date else "-")
    sla_str = sla_target_date.isoformat()[:10] if hasattr(sla_target_date, "isoformat") else (str(sla_target_date) if sla_target_date else "-")

    cust_block = [
        Paragraph("CUSTOMER", s["label"]),
        Paragraph(customer_name, s["value_b"]),
        Spacer(1, 4),
        Paragraph("ASSET", s["label"]),
        Paragraph(f"S/N {asset_serial}", s["value_b"]),
        Paragraph(asset_sku or "-", s["small"]),
    ]
    meta_rows = [
        ("WO Number", wo_number),
        ("Service Type", type.upper()),
        ("Scheduled Date", sched_str),
        ("SLA Target", sla_str),
        ("Region", region),
        ("Assigned Team", assigned_team or "-"),
        ("Technician", technician or "Unassigned"),
        ("Service Contract", service_contract_id or "-"),
        ("Sign-off Status", signoff_status.upper()),
    ]
    meta = _meta_table(meta_rows, s, col_widths=(1.55 * inch, 2.0 * inch))

    top = Table(
        [[cust_block, meta]],
        colWidths=[3.70 * inch, 3.55 * inch],
    )
    top.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ]
        )
    )
    story.append(top)
    story.append(Spacer(1, 0.18 * inch))

    story.append(Paragraph("DESCRIPTION OF WORK", s["section"]))
    story.append(Paragraph(description or "Routine field service per scope of contract.", s["small"]))
    story.append(Spacer(1, 0.14 * inch))

    standards = standards_referenced or []
    if standards:
        story.append(Paragraph("STANDARDS & PROCEDURES REFERENCED", s["section"]))
        bullets = "<br/>".join(f"• {std}" for std in standards)
        story.append(Paragraph(bullets, s["small"]))
        story.append(Spacer(1, 0.14 * inch))

    parts = parts_used or []
    story.append(Paragraph("PARTS USED", s["section"]))
    if parts:
        prows = [[
            Paragraph("Part #", s["small_b"]),
            Paragraph("Description", s["small_b"]),
            Paragraph("Qty", s["right_b"]),
            Paragraph("Unit Cost", s["right_b"]),
            Paragraph("Extended", s["right_b"]),
        ]]
        parts_total = 0.0
        for p in parts:
            ext = float(p.get("qty", 1)) * float(p.get("unit_cost", 0.0))
            parts_total += ext
            prows.append([
                Paragraph(str(p.get("part_number", "-")), s["small"]),
                Paragraph(str(p.get("description", "-")), s["small"]),
                Paragraph(str(p.get("qty", 1)), s["right"]),
                Paragraph(f"${float(p.get('unit_cost', 0.0)):,.2f}", s["right"]),
                Paragraph(f"${ext:,.2f}", s["right"]),
            ])
        ptable = Table(prows, colWidths=[1.05 * inch, 3.40 * inch, 0.55 * inch, 1.10 * inch, 1.15 * inch], repeatRows=1)
        ptable.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), TINT),
                    ("LINEBELOW", (0, 0), (-1, 0), 0.6, INK),
                    ("LINEBELOW", (0, 1), (-1, -1), 0.25, DIVIDER),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ]
            )
        )
        story.append(ptable)
    else:
        story.append(Paragraph("None.", s["small"]))
    story.append(Spacer(1, 0.16 * inch))

    cost_rows = [
        [Paragraph("Labor Hours", s["label"]), Paragraph(f"{labor_hours:.1f} h", s["right"])],
        [Paragraph("Parts & Labor Cost", s["label"]), Paragraph(f"${cost_usd:,.2f} USD", s["right_b"])],
    ]
    cost_table = Table(cost_rows, colWidths=[1.6 * inch, 1.95 * inch])
    cost_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), TINT),
                ("BOX", (0, 0), (-1, -1), 0.4, DIVIDER),
                ("INNERGRID", (0, 0), (-1, -1), 0.25, DIVIDER),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]
        )
    )
    cost_wrap = Table([["", cost_table]], colWidths=[3.70 * inch, 3.55 * inch])
    cost_wrap.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ]
        )
    )
    story.append(cost_wrap)
    story.append(Spacer(1, 0.16 * inch))

    if root_cause:
        story.append(Paragraph("ROOT CAUSE / FINDINGS", s["section"]))
        story.append(Paragraph(root_cause, s["small"]))
        story.append(Spacer(1, 0.14 * inch))

    sig_left = [
        Paragraph("TECHNICIAN SIGN-OFF", s["label"]),
        Spacer(1, 28),
        Paragraph("_______________________________________", s["small"]),
        Paragraph(f"{technician or 'Field Technician'} &nbsp;·&nbsp; Date ______________", s["footnote"]),
    ]
    sig_right = [
        Paragraph("CUSTOMER ACCEPTANCE", s["label"]),
        Spacer(1, 28),
        Paragraph("_______________________________________", s["small"]),
        Paragraph("Authorized Customer Representative &nbsp;·&nbsp; Date ______________", s["footnote"]),
    ]
    sig = Table([[sig_left, sig_right]], colWidths=[3.625 * inch, 3.625 * inch])
    sig.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, -1), 10),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 12),
                ("BOX", (0, 0), (-1, -1), 0.4, DIVIDER),
                ("LINEAFTER", (0, 0), (0, -1), 0.4, DIVIDER),
            ]
        )
    )
    story.append(KeepTogether(sig))
    story.append(Spacer(1, 0.10 * inch))
    story.append(
        Paragraph(
            "This work order is performed under Keysight's standard Field Service & Calibration Terms. "
            "Customer signature confirms work was completed to the customer's satisfaction. Retain this "
            "document for audit and ISO 9001 / ISO/IEC 17025 records.",
            s["footnote"],
        )
    )

    doc.build(story)
    return out_path


def make_calibration_cert_pdf(
    out_path: Path,
    *,
    cert_number: str,
    customer_name: str,
    asset_sku: str | None,
    asset_serial: str | None,
    asset_description: str | None,
    traceability: str,
    lab_id: str | None,
    technician: str | None,
    issued_date,
    expires_date,
    out_of_tolerance: bool,
    as_found_summary: str | None,
    as_left_summary: str | None,
    standards_referenced: list[str] | None = None,
) -> Path:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=LETTER,
        leftMargin=0.55 * inch,
        rightMargin=0.55 * inch,
        topMargin=0.55 * inch,
        bottomMargin=0.55 * inch,
        title=f"Calibration Certificate {cert_number}",
        author="Keysight Technologies, Inc.",
    )
    s = _doc_styles()
    story = []

    header_left = [
        Paragraph("Keysight Technologies, Inc.", s["value_b"]),
        Spacer(1, 2),
        Paragraph("Metrology &amp; Calibration Services", s["small"]),
        Paragraph("1400 Fountaingrove Pkwy, Santa Rosa, CA 95403, USA", s["small"]),
    ]
    header_right = [
        Paragraph("CALIBRATION CERTIFICATE", s["title"]),
        Paragraph(f"Certificate {cert_number}", s["subtitle"]),
    ]
    header_table = Table(
        [[header_left, header_right]],
        colWidths=[4.05 * inch, 3.20 * inch],
    )
    header_table.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ALIGN", (1, 0), (1, 0), "RIGHT"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ]
        )
    )
    story.append(header_table)
    story.append(Spacer(1, 0.18 * inch))

    issued_str = issued_date.isoformat()[:10] if hasattr(issued_date, "isoformat") else (str(issued_date) if issued_date else "-")
    expires_str = expires_date.isoformat()[:10] if hasattr(expires_date, "isoformat") else (str(expires_date) if expires_date else "-")

    cust_block = [
        Paragraph("CUSTOMER OF RECORD", s["label"]),
        Paragraph(customer_name, s["value_b"]),
        Spacer(1, 6),
        Paragraph("INSTRUMENT UNDER TEST", s["label"]),
        Paragraph(asset_description or asset_sku or "-", s["value_b"]),
        Paragraph(f"Model: {asset_sku or '—'} &nbsp;·&nbsp; S/N: {asset_serial or '—'}", s["small"]),
    ]
    meta_rows = [
        ("Certificate Number", cert_number),
        ("Issued Date", issued_str),
        ("Expires Date", expires_str),
        ("Calibration Lab", lab_id or "Keysight Metrology"),
        ("Technician", technician or "-"),
        ("Traceability", traceability.replace("_", " ")),
    ]
    meta = _meta_table(meta_rows, s, col_widths=(1.65 * inch, 1.9 * inch))

    top = Table(
        [[cust_block, meta]],
        colWidths=[3.70 * inch, 3.55 * inch],
    )
    top.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ]
        )
    )
    story.append(top)
    story.append(Spacer(1, 0.18 * inch))

    story.append(Paragraph("STATEMENT OF TRACEABILITY", s["section"]))
    story.append(
        Paragraph(
            f"This calibration is traceable to the International System of Units (SI) through the National "
            f"Institute of Standards and Technology (NIST) and equivalent National Metrology Institutes via "
            f"unbroken chain of comparisons under <b>{traceability.replace('_', ' ')}</b>. Measurement "
            f"uncertainties are reported with a coverage factor k=2, providing a confidence level of "
            f"approximately 95%.",
            s["small"],
        )
    )
    story.append(Spacer(1, 0.16 * inch))

    af_label_color = AMBER if out_of_tolerance else EMERALD
    af_tint_color = AMBER_TINT if out_of_tolerance else EMERALD_TINT
    af_status = "OUT OF TOLERANCE" if out_of_tolerance else "IN TOLERANCE"

    af_block = Table(
        [
            [Paragraph("AS-FOUND CONDITION", s["label"])],
            [Paragraph(f"<b>{af_status}</b>", ParagraphStyle("af", parent=s["small_b"], textColor=af_label_color))],
            [Paragraph(as_found_summary or "All measurement points within published spec.", s["small"])],
        ],
        colWidths=[3.55 * inch],
    )
    af_block.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), af_tint_color),
                ("BOX", (0, 0), (-1, -1), 0.5, af_label_color),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )

    al_block = Table(
        [
            [Paragraph("AS-LEFT CONDITION", s["label"])],
            [Paragraph("<b>IN TOLERANCE</b>", ParagraphStyle("al", parent=s["small_b"], textColor=EMERALD))],
            [Paragraph(as_left_summary or "All as-left points within published spec; cert package issued.", s["small"])],
        ],
        colWidths=[3.55 * inch],
    )
    al_block.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), EMERALD_TINT),
                ("BOX", (0, 0), (-1, -1), 0.5, EMERALD),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )

    panels = Table(
        [[af_block, al_block]],
        colWidths=[3.625 * inch, 3.625 * inch],
    )
    panels.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ]
        )
    )
    story.append(panels)
    story.append(Spacer(1, 0.18 * inch))

    standards = standards_referenced or [
        "ANSI/NCSL Z540.3 — Requirements for the Calibration of Measuring and Test Equipment",
        "ISO/IEC 17025:2017 — General requirements for the competence of testing and calibration laboratories",
        "NIST Handbook 150 — NVLAP Procedures and General Requirements",
    ]
    story.append(Paragraph("STANDARDS REFERENCED", s["section"]))
    standards_body = "<br/>".join(f"• {std}" for std in standards)
    standards_box = Table(
        [[Paragraph(standards_body, s["small"])]],
        colWidths=[7.25 * inch],
    )
    standards_box.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), TINT),
                ("BOX", (0, 0), (-1, -1), 0.4, DIVIDER),
                ("LEFTPADDING", (0, 0), (-1, -1), 12),
                ("RIGHTPADDING", (0, 0), (-1, -1), 12),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ]
        )
    )
    story.append(standards_box)
    story.append(Spacer(1, 0.20 * inch))

    accred_block = [
        Paragraph("ACCREDITATION", s["stamp"]),
        Spacer(1, 4),
        Paragraph("ANAB / A2LA accredited<br/>Certificate ANAB AC-2298<br/>ISO/IEC 17025:2017", s["stamp_sm"]),
    ]
    sig_tech = [
        Paragraph("CALIBRATED BY", s["label"]),
        Spacer(1, 22),
        Paragraph("_______________________________", s["small"]),
        Paragraph(f"{technician or 'Calibration Technician'} &nbsp;·&nbsp; {issued_str}", s["footnote"]),
    ]
    sig_qa = [
        Paragraph("QA APPROVED", s["label"]),
        Spacer(1, 22),
        Paragraph("_______________________________", s["small"]),
        Paragraph(f"Lab Manager, {lab_id or 'Keysight Metrology'}", s["footnote"]),
    ]
    footer = Table(
        [[accred_block, sig_tech, sig_qa]],
        colWidths=[2.30 * inch, 2.475 * inch, 2.475 * inch],
    )
    footer.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, -1), 10),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 12),
                ("BOX", (0, 0), (-1, -1), 0.4, DIVIDER),
                ("LINEAFTER", (0, 0), (1, -1), 0.4, DIVIDER),
            ]
        )
    )
    story.append(KeepTogether(footer))
    story.append(Spacer(1, 0.12 * inch))
    story.append(
        Paragraph(
            "This certificate may not be reproduced except in full, without the written approval of the issuing "
            "laboratory. Results relate only to the items calibrated. Calibration interval is determined by the "
            "customer based on use and environmental conditions.",
            s["footnote"],
        )
    )

    doc.build(story)
    return out_path
